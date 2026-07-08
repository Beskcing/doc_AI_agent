#!/usr/bin/env python
"""
GB/T DOCX Format Analyzer
=========================
Analyzes formatting of GB/T (Chinese national standard) DOCX files.
Outputs structured JSON and human-readable Markdown reports.

Usage:
    python analyze_docx.py --input "D:/Qwork/gbt" --output "D:/Qwork/gbt/output" --verbose
    python analyze_docx.py --single "path/to/file.docx" --output "./out"
"""

import argparse
import datetime
import glob
import json
import os
import re
import sys

try:
    import docx
    from docx import Document
except ImportError:
    sys.stderr.write("ERROR: python-docx is not installed. Run: pip install python-docx\n")
    sys.exit(2)

# ---------------------------------------------------------------------------
# Constants & Namespace helpers
# ---------------------------------------------------------------------------

EMU_PER_PT = 12700.0
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

ANALYZER_VERSION = "1.0.0"

# Heading detection regex: matches "1", "3.1", "4.1.1", etc. at start of text
HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)")


def qn(tag):
    """Build a namespace-qualified tag name from 'w:localname'."""
    prefix, local = tag.split(":")
    nsmap = {"w": W_NS}
    return "{%s}%s" % (nsmap[prefix], local)


def emu_to_pt(emu):
    """Convert EMU to points. Returns None if input is None."""
    if emu is None:
        return None
    try:
        return round(int(emu) / EMU_PER_PT, 2)
    except (ValueError, TypeError):
        return None


def half_pt_to_pt(half_pt):
    """Convert half-points to points. Returns None if input is None."""
    if half_pt is None:
        return None
    try:
        return round(int(half_pt) / 2.0, 2)
    except (ValueError, TypeError):
        return None


def safe_text(text, max_len=80):
    """Truncate and clean text for safe display."""
    if not text:
        return ""
    cleaned = text.replace("\n", " ").replace("\r", " ").replace("\t", " ")
    cleaned = cleaned.strip()
    if len(cleaned) > max_len:
        cleaned = cleaned[:max_len] + "..."
    return cleaned


def infer_paper_size(w_emu, h_emu):
    """Infer paper size from EMU dimensions."""
    if not w_emu or not h_emu:
        return "UNKNOWN"
    w_mm = int(w_emu) / 36000.0
    h_mm = int(h_emu) / 36000.0
    lo, hi = min(w_mm, h_mm), max(w_mm, h_mm)
    if abs(lo - 210) < 5 and abs(hi - 297) < 5:
        return "A4"
    if abs(lo - 297) < 5 and abs(hi - 420) < 5:
        return "A3"
    if abs(lo - 215.9) < 5 and abs(hi - 279.4) < 5:
        return "Letter"
    return "CUSTOM (%.0f x %.0f mm)" % (w_mm, h_mm)


def tag_present(rpr, tag_local):
    """Check if a w: tag exists in rPr element."""
    if rpr is None:
        return False
    return rpr.find(qn(tag_local)) is not None


def get_vert_align(rpr):
    """Get vertical alignment (superscript/subscript)."""
    if rpr is None:
        return None
    va = rpr.find(qn("w:vertAlign"))
    if va is not None:
        return va.get(qn("w:val"))
    return None


# ---------------------------------------------------------------------------
# Font extraction (core — reads eastAsia directly from XML)
# ---------------------------------------------------------------------------


def extract_run_fonts(run_element):
    """
    Extract all font attributes from a run's w:rPr/w:rFonts.
    python-docx's run.font.name only returns the ascii font;
    the Chinese font is in w:eastAsia.
    """
    result = {
        "ascii": None,
        "hAnsi": None,
        "eastAsia": None,
        "cs": None,
        "asciiTheme": None,
        "hAnsiTheme": None,
        "eastAsiaTheme": None,
        "cstheme": None,
    }
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        return result
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return result
    for key, attr in [
        ("ascii", "w:ascii"),
        ("hAnsi", "w:hAnsi"),
        ("eastAsia", "w:eastAsia"),
        ("cs", "w:cs"),
        ("asciiTheme", "w:asciiTheme"),
        ("hAnsiTheme", "w:hAnsiTheme"),
        ("eastAsiaTheme", "w:eastAsiaTheme"),
        ("cstheme", "w:cstheme"),
    ]:
        val = rfonts.get(qn(attr))
        if val:
            result[key] = val
    return result


def extract_run_size(run_element):
    """Extract font size in half-points from w:sz and w:szCs."""
    rpr = run_element.find(qn("w:rPr"))
    sz_half = None
    sz_cs_half = None
    if rpr is not None:
        sz_el = rpr.find(qn("w:sz"))
        if sz_el is not None:
            try:
                sz_half = int(sz_el.get(qn("w:val")))
            except (ValueError, TypeError):
                pass
        szcs_el = rpr.find(qn("w:szCs"))
        if szcs_el is not None:
            try:
                sz_cs_half = int(szcs_el.get(qn("w:val")))
            except (ValueError, TypeError):
                pass
    return sz_half, sz_cs_half


# ---------------------------------------------------------------------------
# Run analysis
# ---------------------------------------------------------------------------


def analyze_run(run, max_text=60):
    """Analyze a single run's complete formatting."""
    r_el = run._element
    fonts = extract_run_fonts(r_el)
    sz_half, sz_cs_half = extract_run_size(r_el)

    # python-docx font size (EMU)
    size_emu = None
    try:
        if run.font.size is not None:
            size_emu = int(run.font.size)
    except Exception:
        pass
    size_pt = emu_to_pt(size_emu)
    # Prefer half-pt derived size if available (more precise from XML)
    if sz_half is not None:
        size_pt = half_pt_to_pt(sz_half)

    # Color
    rpr = r_el.find(qn("w:rPr"))
    color_rgb = None
    if rpr is not None:
        color_el = rpr.find(qn("w:color"))
        if color_el is not None:
            color_rgb = color_el.get(qn("w:val"))

    return {
        "text_preview": safe_text(run.text, max_text),
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline if run.underline is not None else tag_present(rpr, "w:u"),
        "strike": tag_present(rpr, "w:strike"),
        "fonts": fonts,
        "size_emu": size_emu,
        "size_pt": size_pt,
        "sz_half_pt": sz_half,
        "sz_cs_half_pt": sz_cs_half,
        "color_rgb": color_rgb,
        "vert_align": get_vert_align(rpr),
    }


# ---------------------------------------------------------------------------
# Heading classification
# ---------------------------------------------------------------------------


def classify_heading(text, first_line_indent_emu):
    """
    GB/T docs don't use Word Heading styles.
    Heading = numbered prefix + zero first-line indent.
    Returns (is_heading, level, number_str).
    """
    stripped = text.strip()
    if not stripped:
        return False, 0, None

    m = HEADING_RE.match(stripped)
    if not m:
        return False, 0, None

    num_part = m.group(1)
    rest = stripped[len(num_part) :]

    # The char after the number must be whitespace, or end of string
    if rest and rest[0] not in (" ", "\t", "\u3000", "\xa0"):
        # No space after number — likely body text like "3.5MPa" or "2008-11-04"
        return False, 0, None

    # Must have zero first-line indent to qualify as heading
    if first_line_indent_emu is not None and int(first_line_indent_emu) > 0:
        return False, 0, None

    level = num_part.count(".") + 1
    return True, level, num_part


# ---------------------------------------------------------------------------
# Paragraph analysis
# ---------------------------------------------------------------------------


def analyze_paragraph(p, idx, max_runs=5):
    """Analyze a paragraph's complete formatting."""
    pf = p.paragraph_format
    text = p.text
    fi_emu = pf.first_line_indent
    fi_emu_int = int(fi_emu) if fi_emu is not None else 0
    is_heading, level, number = classify_heading(text, fi_emu_int)

    alignment = None
    if p.alignment is not None:
        alignment = p.alignment.name if hasattr(p.alignment, "name") else str(p.alignment)

    line_spacing_rule = None
    if pf.line_spacing_rule is not None:
        line_spacing_rule = (
            pf.line_spacing_rule.name if hasattr(pf.line_spacing_rule, "name") else str(pf.line_spacing_rule)
        )

    runs_data = [analyze_run(r) for r in p.runs[:max_runs]]

    return {
        "index": idx,
        "text_preview": safe_text(text, 80),
        "style": p.style.name if p.style else None,
        "alignment": alignment,
        "is_heading": is_heading,
        "heading_level": level,
        "heading_number": number,
        "paragraph_format": {
            "first_line_indent_emu": fi_emu,
            "first_line_indent_pt": emu_to_pt(fi_emu),
            "left_indent_emu": pf.left_indent,
            "left_indent_pt": emu_to_pt(pf.left_indent),
            "right_indent_emu": pf.right_indent,
            "right_indent_pt": emu_to_pt(pf.right_indent),
            "space_before_emu": pf.space_before,
            "space_before_pt": emu_to_pt(pf.space_before),
            "space_after_emu": pf.space_after,
            "space_after_pt": emu_to_pt(pf.space_after),
            "line_spacing": pf.line_spacing,
            "line_spacing_rule": line_spacing_rule,
        },
        "run_count": len(p.runs),
        "runs": runs_data,
    }


# ---------------------------------------------------------------------------
# Table analysis
# ---------------------------------------------------------------------------


def _first_run_from_cells(cells, max_cells=3):
    """Get the first run from the first non-empty cell paragraph."""
    for cell in cells[:max_cells]:
        for p in cell.paragraphs:
            for r in p.runs[:1]:
                return analyze_run(r)
    return None


def _first_bold_from_cells(cells, max_cells=3):
    """Check if runs in header cells are bold."""
    for cell in cells[:max_cells]:
        for p in cell.paragraphs:
            for r in p.runs[:1]:
                return r.bold
    return None


def analyze_table(t, idx):
    """Analyze a table's formatting."""
    style_name = t.style.name if t.style else None
    rows = len(t.rows)
    cols = len(t.columns) if t.columns else (len(t.rows[0].cells) if rows else 0)

    header_info = {}
    if rows > 0:
        first_row = t.rows[0]
        header_cells = [safe_text(c.text, 30) for c in first_row.cells[:12]]
        header_info = {
            "cells": header_cells,
            "is_bold": _first_bold_from_cells(first_row.cells),
            "sample_run": _first_run_from_cells(first_row.cells),
        }

    body_info = {}
    if rows > 1:
        second_row = t.rows[1]
        body_cells = [safe_text(c.text, 30) for c in second_row.cells[:8]]
        body_info = {
            "cells": body_cells,
            "sample_run": _first_run_from_cells(second_row.cells),
        }

    return {
        "index": idx,
        "style": style_name,
        "rows": rows,
        "cols": cols,
        "header_row": header_info,
        "body_sample": body_info,
    }


# ---------------------------------------------------------------------------
# Styles analysis
# ---------------------------------------------------------------------------


def _extract_rpr_info(rpr):
    """Extract font/size info from an rPr element."""
    info = {"fonts": {}, "sz_half_pt": None, "sz_cs_half_pt": None, "bold": None}
    if rpr is None:
        return info
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is not None:
        for attr_name, attr in [
            ("ascii", "w:ascii"),
            ("hAnsi", "w:hAnsi"),
            ("eastAsia", "w:eastAsia"),
            ("cs", "w:cs"),
            ("asciiTheme", "w:asciiTheme"),
            ("eastAsiaTheme", "w:eastAsiaTheme"),
        ]:
            val = rfonts.get(qn(attr))
            if val:
                info["fonts"][attr_name] = val
    sz_el = rpr.find(qn("w:sz"))
    if sz_el is not None:
        try:
            info["sz_half_pt"] = int(sz_el.get(qn("w:val")))
        except (ValueError, TypeError):
            pass
    szcs_el = rpr.find(qn("w:szCs"))
    if szcs_el is not None:
        try:
            info["sz_cs_half_pt"] = int(szcs_el.get(qn("w:val")))
        except (ValueError, TypeError):
            pass
    b_el = rpr.find(qn("w:b"))
    if b_el is not None:
        val = b_el.get(qn("w:val"))
        info["bold"] = False if val == "0" or val == "false" else True
    return info


def _extract_ppr_info(ppr):
    """Extract paragraph spacing info from a pPr element."""
    info = {"spacing": {}}
    if ppr is None:
        return info
    spacing = ppr.find(qn("w:spacing"))
    if spacing is not None:
        for attr in ["w:before", "w:after", "w:line", "w:lineRule"]:
            val = spacing.get(qn(attr))
            if val:
                key = attr.split(":")[-1]
                info["spacing"][key] = val
    return info


def analyze_doc_defaults(doc):
    """Analyze w:docDefaults from styles part."""
    result = {"rPrDefault": {}, "pPrDefault": {}}
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return result
    rpr_def = doc_defaults.find(qn("w:rPrDefault") + "/" + qn("w:rPr"))
    if rpr_def is not None:
        result["rPrDefault"] = _extract_rpr_info(rpr_def)
    ppr_def = doc_defaults.find(qn("w:pPrDefault") + "/" + qn("w:pPr"))
    if ppr_def is not None:
        result["pPrDefault"] = _extract_ppr_info(ppr_def)
    return result


def analyze_styles(doc):
    """Analyze all defined styles."""
    styles_elem = doc.styles.element
    all_styles = []
    normal_style = {}

    for style_elem in styles_elem.findall(qn("w:style")):
        name_elem = style_elem.find(qn("w:name"))
        if name_elem is None:
            continue
        name = name_elem.get(qn("w:val"))
        style_type = style_elem.get(qn("w:type"))
        style_id = style_elem.get(qn("w:styleId"))
        default_attr = style_elem.get(qn("w:default"))

        # Build using paths relative to style_elem
        rpr = style_elem.find(qn("w:rPr"))
        ppr = style_elem.find(qn("w:pPr"))

        style_entry = {
            "name": name,
            "type": style_type.upper() if style_type else None,
            "styleId": style_id,
            "is_default": default_attr == "1",
            "rPr": _extract_rpr_info(rpr) if rpr is not None else {},
            "pPr": _extract_ppr_info(ppr) if ppr is not None else {},
        }

        if name == "Normal":
            normal_style = style_entry
        all_styles.append(
            {
                "name": name,
                "type": style_type.upper() if style_type else None,
                "builtin": style_elem.get(qn("w:custom")) != "1",
                "is_default": default_attr == "1",
            }
        )

    return {"doc_defaults": analyze_doc_defaults(doc), "normal_style": normal_style, "all_styles": all_styles}


# ---------------------------------------------------------------------------
# Sections / Page setup analysis
# ---------------------------------------------------------------------------


def analyze_sections(doc):
    """Analyze page setup for all sections."""
    sections = []
    for i, s in enumerate(doc.sections):
        w_emu = s.page_width
        h_emu = s.page_height
        orient = str(s.orientation).split(".")[-1] if s.orientation else None
        actual_orient = "LANDSCAPE" if w_emu and h_emu and int(w_emu) > int(h_emu) else "PORTRAIT"

        sections.append(
            {
                "index": i,
                "page_width_emu": w_emu,
                "page_height_emu": h_emu,
                "page_width_pt": emu_to_pt(w_emu),
                "page_height_pt": emu_to_pt(h_emu),
                "orientation_declared": orient,
                "orientation_actual": actual_orient,
                "margin_left_emu": s.left_margin,
                "margin_right_emu": s.right_margin,
                "margin_top_emu": s.top_margin,
                "margin_bottom_emu": s.bottom_margin,
                "margin_left_pt": emu_to_pt(s.left_margin),
                "margin_right_pt": emu_to_pt(s.right_margin),
                "margin_top_pt": emu_to_pt(s.top_margin),
                "margin_bottom_pt": emu_to_pt(s.bottom_margin),
                "header_distance_emu": s.header_distance,
                "footer_distance_emu": s.footer_distance,
                "gutter_emu": s.gutter,
                "paper_size_guess": infer_paper_size(w_emu, h_emu),
            }
        )
    return sections


# ---------------------------------------------------------------------------
# Format pattern detection (heuristic)
# ---------------------------------------------------------------------------


def detect_format_patterns(paragraphs):
    """Detect cover block, preface block, and body block patterns."""
    patterns = {"cover_block": [], "preface_block": None, "body_block": None}

    for p in paragraphs:
        text = p["text_preview"]
        runs = p["runs"]
        if not runs:
            continue
        first_run = runs[0]
        size_pt = first_run.get("size_pt")
        east = first_run["fonts"].get("eastAsia")
        latin = first_run["fonts"].get("ascii")
        bold = first_run.get("bold")

        # Cover block: >= 15pt and bold (within first 20 paragraphs)
        if size_pt and size_pt >= 15 and bold and p["index"] < 20:
            patterns["cover_block"].append(
                {
                    "para_index": p["index"],
                    "text": text,
                    "size_pt": size_pt,
                    "bold": bold,
                    "east_asian_font": east,
                    "latin_font": latin,
                }
            )

        # Preface heading: contains both "前" and "言"
        if "\u524d" in text and "\u8a00" in text and p["index"] < 30 and patterns["preface_block"] is None:
            patterns["preface_block"] = {
                "para_index": p["index"],
                "size_pt": size_pt,
                "east_asian_font": east,
            }

        # Body block: 10-11pt with first-line indent
        pf = p["paragraph_format"]
        fi_emu = pf.get("first_line_indent_emu")
        if size_pt and 10 <= size_pt <= 11 and fi_emu and int(fi_emu) > 0 and patterns["body_block"] is None:
            patterns["body_block"] = {
                "para_index": p["index"],
                "size_pt": size_pt,
                "east_asian_font": east,
                "latin_font": latin,
                "alignment": p["alignment"],
                "first_line_indent_pt": pf.get("first_line_indent_pt"),
                "line_spacing": pf.get("line_spacing"),
            }

    return patterns


# ---------------------------------------------------------------------------
# Font aggregation
# ---------------------------------------------------------------------------


def aggregate_fonts(paragraphs, tables):
    """Aggregate font usage and size distribution across document."""
    east_fonts = {}
    latin_fonts = {}
    sizes = {}

    def collect(run):
        f = run.get("fonts", {})
        ea = f.get("eastAsia")
        la = f.get("ascii")
        if ea:
            east_fonts[ea] = east_fonts.get(ea, 0) + 1
        if la:
            latin_fonts[la] = latin_fonts.get(la, 0) + 1
        sp = run.get("size_pt")
        if sp:
            key = "%.1fpt" % sp
            sizes[key] = sizes.get(key, 0) + 1

    for p in paragraphs:
        for r in p["runs"]:
            collect(r)
    for t in tables:
        for block_key in ("header_row", "body_sample"):
            sr = t.get(block_key, {}).get("sample_run")
            if sr:
                collect(sr)

    return {
        "east_asian_fonts_used": dict(sorted(east_fonts.items(), key=lambda x: -x[1])),
        "latin_fonts_used": dict(sorted(latin_fonts.items(), key=lambda x: -x[1])),
        "font_size_distribution": dict(sorted(sizes.items(), key=lambda x: -x[1])),
    }


# ---------------------------------------------------------------------------
# Document-level analysis
# ---------------------------------------------------------------------------


def analyze_document(path, max_runs=5):
    """Analyze a single DOCX file and return complete format data."""
    doc = Document(path)
    fname = os.path.basename(path)
    file_size = os.path.getsize(path)

    # Paragraphs
    paragraphs = []
    headings_outline = []
    for i, p in enumerate(doc.paragraphs):
        p_data = analyze_paragraph(p, i, max_runs)
        paragraphs.append(p_data)
        if p_data["is_heading"]:
            headings_outline.append(
                {
                    "para_index": i,
                    "level": p_data["heading_level"],
                    "number": p_data["heading_number"],
                    "text": p_data["text_preview"],
                }
            )

    # Tables
    tables = [analyze_table(t, i) for i, t in enumerate(doc.tables)]

    # Styles
    styles_data = analyze_styles(doc)

    # Sections
    sections = analyze_sections(doc)

    # Fonts summary
    fonts_summary = aggregate_fonts(paragraphs, tables)

    # Format patterns
    format_patterns = detect_format_patterns(paragraphs)

    # Image count
    image_count = 0
    try:
        image_count = len(doc.inline_shapes)
    except Exception:
        pass

    return {
        "metadata": {
            "file_name": fname,
            "file_path": os.path.abspath(path),
            "file_size_bytes": file_size,
            "analyzed_at": datetime.datetime.now().isoformat(),
            "analyzer_version": ANALYZER_VERSION,
        },
        "document_summary": {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "image_count": image_count,
            "style_count": len(styles_data.get("all_styles", [])),
            "heading_count": len(headings_outline),
        },
        "page_setup": {"sections": sections},
        "styles": styles_data,
        "fonts_summary": fonts_summary,
        "paragraphs": paragraphs,
        "headings_outline": headings_outline,
        "tables": tables,
        "format_patterns": format_patterns,
    }


# ---------------------------------------------------------------------------
# Markdown generation
# ---------------------------------------------------------------------------


def generate_markdown(data):
    """Generate a Markdown report for a single document."""
    md = []
    meta = data["metadata"]
    ds = data["document_summary"]

    md.append("# Format Analysis Report: %s\n" % meta["file_name"])
    md.append(
        "> Analyzed: %s | Analyzer v%s | Size: %d bytes\n"
        % (meta["analyzed_at"], meta["analyzer_version"], meta["file_size_bytes"])
    )
    md.append(
        "> Paragraphs: %d | Tables: %d | Sections: %d | Headings: %d | Images: %d\n"
        % (ds["paragraph_count"], ds["table_count"], ds["section_count"], ds["heading_count"], ds["image_count"])
    )

    # 1. Page Setup
    md.append("\n## 1. Page Setup\n")
    md.append("| Sec | Paper | Orientation | W(pt) | H(pt) | Margin L | Margin R | Margin T | Margin B |")
    md.append("|-----|-------|-------------|-------|-------|----------|----------|----------|----------|")
    for s in data["page_setup"]["sections"]:
        md.append(
            "| %d | %s | %s | %s | %s | %s | %s | %s | %s |"
            % (
                s["index"],
                s["paper_size_guess"],
                s["orientation_actual"],
                s["page_width_pt"],
                s["page_height_pt"],
                s["margin_left_pt"],
                s["margin_right_pt"],
                s["margin_top_pt"],
                s["margin_bottom_pt"],
            )
        )

    # 2. Document Defaults & Normal Style
    md.append("\n## 2. Document Defaults & Normal Style\n")
    dd = data["styles"]["doc_defaults"]
    if dd.get("rPrDefault"):
        rdef = dd["rPrDefault"]
        fonts = rdef.get("fonts", {})
        md.append(
            "- **Doc Defaults Fonts**: ascii=%s, eastAsia=%s"
            % (
                fonts.get("ascii") or fonts.get("asciiTheme") or "N/A",
                fonts.get("eastAsia") or fonts.get("eastAsiaTheme") or "N/A",
            )
        )
        if rdef.get("sz_half_pt"):
            md.append("- **Doc Defaults Size**: %s half-pt (%.1fpt)" % (rdef["sz_half_pt"], rdef["sz_half_pt"] / 2.0))
    ns = data["styles"].get("normal_style", {})
    if ns:
        nfonts = ns.get("rPr", {}).get("fonts", {})
        md.append(
            "- **Normal Style Fonts**: ascii=%s, eastAsia=%s"
            % (
                nfonts.get("ascii") or nfonts.get("asciiTheme") or "N/A",
                nfonts.get("eastAsia") or nfonts.get("eastAsiaTheme") or "N/A",
            )
        )
        if ns.get("rPr", {}).get("sz_half_pt"):
            sz = ns["rPr"]["sz_half_pt"]
            md.append("- **Normal Style Size**: %s half-pt (%.1fpt)" % (sz, sz / 2.0))
    md.append("- **Total Styles Defined**: %d" % ds["style_count"])

    # 3. Font Usage Summary
    md.append("\n## 3. Font Usage Summary\n")
    fs = data["fonts_summary"]
    md.append("### East-Asian Fonts (Chinese)\n")
    md.append("| Font | Count |")
    md.append("|------|-------|")
    for font, count in fs.get("east_asian_fonts_used", {}).items():
        md.append("| %s | %d |" % (font, count))
    if not fs.get("east_asian_fonts_used"):
        md.append("| (none detected) | 0 |")

    md.append("\n### Latin Fonts (English/Numbers)\n")
    md.append("| Font | Count |")
    md.append("|------|-------|")
    for font, count in fs.get("latin_fonts_used", {}).items():
        md.append("| %s | %d |" % (font, count))
    if not fs.get("latin_fonts_used"):
        md.append("| (none detected) | 0 |")

    md.append("\n### Font Size Distribution\n")
    md.append("| Size | Count |")
    md.append("|------|-------|")
    for size, count in fs.get("font_size_distribution", {}).items():
        md.append("| %s | %d |" % (size, count))

    # 4. Cover Block Format
    md.append("\n## 4. Cover Block Format\n")
    cover = data["format_patterns"]["cover_block"]
    if cover:
        md.append("| Para | Text | Size | Bold | CJK Font | Latin Font |")
        md.append("|------|------|------|------|----------|------------|")
        for c in cover:
            md.append(
                "| %d | %s | %spt | %s | %s | %s |"
                % (
                    c["para_index"],
                    c["text"],
                    c["size_pt"],
                    c["bold"],
                    c["east_asian_font"] or "N/A",
                    c["latin_font"] or "N/A",
                )
            )
    else:
        md.append("(No cover block detected)")

    # 5. Body Block Format
    md.append("\n## 5. Body Block Format\n")
    body = data["format_patterns"]["body_block"]
    if body:
        md.append("- **CJK Font**: %s" % (body.get("east_asian_font") or "N/A"))
        md.append("- **Latin Font**: %s" % (body.get("latin_font") or "N/A"))
        md.append("- **Size**: %spt" % body.get("size_pt"))
        md.append("- **Alignment**: %s" % body.get("alignment"))
        md.append("- **First-line Indent**: %spt" % body.get("first_line_indent_pt"))
        md.append("- **Line Spacing**: %s" % body.get("line_spacing"))
    else:
        md.append("(No body block detected)")

    # 6. Heading Hierarchy
    md.append("\n## 6. Heading Hierarchy\n")
    md.append("> Note: This document does not use Word Heading styles. Headings are")
    md.append("> identified by numbered prefix + zero first-line indent.\n")
    outline = data["headings_outline"]
    if outline:
        md.append("| Level | Number | Text |")
        md.append("|-------|--------|------|")
        for h in outline[:50]:
            md.append("| %d | %s | %s |" % (h["level"], h["number"], h["text"]))
        if len(outline) > 50:
            md.append("| ... | ... | (%d more headings) |" % (len(outline) - 50))
    else:
        md.append("(No headings detected)")

    # 7. Tables
    md.append("\n## 7. Tables\n")
    tables = data["tables"]
    if tables:
        md.append("| # | Style | Rows x Cols | Header Bold | Header CJK Font | Header Size |")
        md.append("|---|-------|-------------|-------------|-----------------|-------------|")
        for t in tables:
            hr = t.get("header_row", {})
            sr = hr.get("sample_run", {})
            hf = sr.get("fonts", {}).get("eastAsia") if sr else None
            hs = sr.get("size_pt") if sr else None
            md.append(
                "| %d | %s | %dx%d | %s | %s | %s |"
                % (
                    t["index"],
                    t["style"] or "N/A",
                    t["rows"],
                    t["cols"],
                    hr.get("is_bold"),
                    hf or "N/A",
                    ("%spt" % hs) if hs else "N/A",
                )
            )
    else:
        md.append("(No tables)")

    # 8. Key Findings
    md.append("\n## 8. Key Findings\n")
    findings = []
    if cover:
        sizes = set(str(c["size_pt"]) for c in cover)
        findings.append("Cover block uses %spt bold formatting with varying CJK fonts" % "/".join(sizes))
    if body:
        findings.append(
            "Body text uses %spt %s + %s, %s aligned, %.1fpt first-line indent"
            % (
                body.get("size_pt"),
                body.get("east_asian_font") or "?",
                body.get("latin_font") or "?",
                body.get("alignment"),
                body.get("first_line_indent_pt") or 0,
            )
        )
    if outline:
        levels = set(h["level"] for h in outline)
        findings.append(
            "Headings detected at %d level(s): %s (inferred from numbering, not Word styles)"
            % (len(levels), ", ".join("L%d" % l for l in sorted(levels)))
        )
    if tables:
        findings.append(
            "%d table(s) using '%s' style with bold headers" % (len(tables), tables[0].get("style") or "default")
        )
    for f in findings:
        md.append("%d. %s" % (findings.index(f) + 1, f))
    if not findings:
        md.append("(No significant findings)")

    return "\n".join(md) + "\n"


# ---------------------------------------------------------------------------
# Summary report
# ---------------------------------------------------------------------------


def generate_summary(all_data, errors):
    """Generate a cross-document comparison summary."""
    md = []
    md.append("# GB/T Document Format Summary Report\n")
    md.append(
        "> Files analyzed: %d | Errors: %d | Generated: %s\n"
        % (len(all_data), len(errors), datetime.datetime.now().isoformat())
    )

    # Comparison table
    md.append("\n## Cross-File Comparison\n")
    md.append("| File | Paper | Orient | Body Size | CJK Font | Latin Font | Align | Indent | Tables |")
    md.append("|------|-------|--------|-----------|----------|------------|-------|--------|--------|")
    for d in all_data:
        fname = d["metadata"]["file_name"]
        sec0 = d["page_setup"]["sections"][0] if d["page_setup"]["sections"] else {}
        paper = sec0.get("paper_size_guess", "?")
        orient = sec0.get("orientation_actual", "?")
        body = d["format_patterns"].get("body_block") or {}
        bsize = ("%spt" % body["size_pt"]) if body.get("size_pt") else "?"
        cjk = body.get("east_asian_font") or "?"
        latin = body.get("latin_font") or "?"
        align = body.get("alignment") or "?"
        indent = body.get("first_line_indent_pt")
        indent_s = ("%.1fpt" % indent) if indent else "?"
        tcount = d["document_summary"]["table_count"]
        md.append(
            "| %s | %s | %s | %s | %s | %s | %s | %s | %d |"
            % (fname, paper, orient, bsize, cjk, latin, align, indent_s, tcount)
        )

    # Common patterns
    md.append("\n## Common Format Patterns\n")
    body_fonts = {}
    body_sizes = {}
    for d in all_data:
        body = d["format_patterns"].get("body_block") or {}
        ea = body.get("east_asian_font")
        sz = body.get("size_pt")
        if ea:
            body_fonts[ea] = body_fonts.get(ea, 0) + 1
        if sz:
            body_sizes[sz] = body_sizes.get(sz, 0) + 1
    if body_fonts:
        md.append(
            "- **Body CJK Font(s)**: %s"
            % ", ".join("%s (%d files)" % (k, v) for k, v in sorted(body_fonts.items(), key=lambda x: -x[1]))
        )
    if body_sizes:
        md.append(
            "- **Body Size(s)**: %s"
            % ", ".join("%spt (%d files)" % (k, v) for k, v in sorted(body_sizes.items(), key=lambda x: -x[1]))
        )

    cover_found = sum(1 for d in all_data if d["format_patterns"].get("cover_block"))
    md.append("- **Cover Block Detected**: %d/%d files" % (cover_found, len(all_data)))
    md.append("- **Heading Style**: Numbered prefix (1, 3.1, 4.1.1) + zero indent (no Word Heading styles)")

    # Differences
    md.append("\n## Differences\n")
    paper_sizes = set()
    orientations = set()
    for d in all_data:
        for s in d["page_setup"]["sections"]:
            paper_sizes.add(s["paper_size_guess"])
            orientations.add(s["orientation_actual"])
    if len(paper_sizes) > 1:
        md.append("- **Paper sizes vary**: %s" % ", ".join(paper_sizes))
    else:
        md.append("- **Paper size uniform**: %s" % (list(paper_sizes)[0] if paper_sizes else "N/A"))
    if len(orientations) > 1:
        md.append("- **Orientations vary**: %s" % ", ".join(orientations))

    # Errors
    if errors:
        md.append("\n## Errors\n")
        for e in errors:
            md.append("- **%s**: %s" % (e["file"], e["error"]))

    return "\n".join(md) + "\n"


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(description="Analyze DOCX formatting of GB/T standard documents.")
    parser.add_argument("--input", default=".", help="Folder containing DOCX files (searched recursively)")
    parser.add_argument("--output", default="./output", help="Output directory for JSON and Markdown reports")
    parser.add_argument("--single", default=None, help="Analyze a single DOCX file instead of a folder")
    parser.add_argument("--max-runs", type=int, default=5, help="Max runs to detail per paragraph")
    parser.add_argument("--verbose", action="store_true", help="Print progress to stdout")
    args = parser.parse_args()

    # Collect files
    if args.single:
        files = [args.single]
    else:
        pattern = os.path.join(args.input, "**", "*.docx")
        files = sorted(glob.glob(pattern, recursive=True))
        files = [f for f in files if not os.path.basename(f).startswith("~$")]

    if not files:
        sys.stderr.write("ERROR: No DOCX files found in: %s\n" % args.input)
        sys.exit(1)

    if args.verbose:
        print("Found %d DOCX file(s)" % len(files))

    # Create output directories
    json_dir = os.path.join(args.output, "json")
    md_dir = os.path.join(args.output, "markdown")
    os.makedirs(json_dir, exist_ok=True)
    os.makedirs(md_dir, exist_ok=True)

    all_data = []
    errors = []

    for fpath in files:
        fname = os.path.basename(fpath)
        stem = os.path.splitext(fname)[0]
        if args.verbose:
            print("Analyzing [%d/%d]: %s" % (len(all_data) + len(errors) + 1, len(files), fname))
        try:
            data = analyze_document(fpath, args.max_runs)
            all_data.append(data)

            # Write JSON
            jp = os.path.join(json_dir, stem + ".json")
            with open(jp, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            # Write Markdown
            md_text = generate_markdown(data)
            mp = os.path.join(md_dir, stem + ".md")
            with open(mp, "w", encoding="utf-8") as f:
                f.write(md_text)

            if args.verbose:
                print("  -> JSON: %s" % jp)
                print("  -> MD:   %s" % mp)
        except Exception as e:
            err = {"file": fname, "error": str(e), "type": type(e).__name__}
            errors.append(err)
            sys.stderr.write("ERROR analyzing %s: %s\n" % (fname, e))

    # Generate summary
    summary = generate_summary(all_data, errors)
    sp = os.path.join(args.output, "summary.md")
    with open(sp, "w", encoding="utf-8") as f:
        f.write(summary)

    print("Done. %d analyzed, %d errors. Summary: %s" % (len(all_data), len(errors), sp))
    if errors:
        sys.exit(1)


if __name__ == "__main__":
    main()
