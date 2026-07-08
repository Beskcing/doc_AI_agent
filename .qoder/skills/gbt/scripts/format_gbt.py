#!/usr/bin/env python
"""
GB/T DOCX Format Adjuster
=========================
Applies standard GB/T formatting to a DOCX file based on format
specifications discovered through analysis of reference documents.

Usage:
    python format_gbt.py "input.docx" --output "output.docx"
"""

import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Emu, Mm, Pt
except ImportError:
    sys.stderr.write("ERROR: python-docx is not installed. Run: pip install python-docx\n")
    sys.exit(2)


# ===========================================================================
# XML helpers
# ===========================================================================


def get_or_add_rPr(run):
    """Get or create the rPr element on a run."""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)
    return rPr


def get_or_add_pPr(paragraph):
    """Get or create the pPr element on a paragraph."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)
    return pPr


def set_child(el, tag, val, attrs=None):
    """Set or update a child element with given tag and value."""
    child = el.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        el.append(child)
    if val is not None:
        if attrs:
            for k, v in attrs.items():
                child.set(qn(k), str(v))
        else:
            child.set(qn("w:val"), str(val))
    return child


def remove_child(el, tag):
    """Remove child element if it exists."""
    child = el.find(qn(tag))
    if child is not None:
        el.remove(child)


# ===========================================================================
# Run-level formatting
# ===========================================================================


def set_run_font(rPr, east_asia, latin, size_pt, bold=None, bold_cs=None):
    """Set font properties on a run's rPr element."""
    # Font names
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    if east_asia:
        rFonts.set(qn("w:eastAsia"), east_asia)
    if latin:
        rFonts.set(qn("w:ascii"), latin)
        rFonts.set(qn("w:hAnsi"), latin)
        rFonts.set(qn("w:cs"), latin)

    # Remove theme references to force explicit fonts
    for attr in ("w:asciiTheme", "w:eastAsiaTheme", "w:hAnsiTheme", "w:cstheme"):
        try:
            del rFonts.attrib[qn(attr)]
        except KeyError:
            pass

    # Font size (half-points)
    if size_pt is not None:
        half_pt = int(round(size_pt * 2))
        set_child(rPr, "w:sz", half_pt)
        set_child(rPr, "w:szCs", half_pt)

    # Bold
    if bold is not None:
        if bold:
            set_child(rPr, "w:b", None, {"w:val": "1"})
        else:
            remove_child(rPr, "w:b")

    if bold_cs is not None:
        if bold_cs:
            set_child(rPr, "w:bCs", None, {"w:val": "1"})
        else:
            remove_child(rPr, "w:bCs")


def format_run(run, east_asia, latin, size_pt, bold=None, bold_cs=None):
    """Format a single run with font properties."""
    rPr = get_or_add_rPr(run)
    set_run_font(rPr, east_asia, latin, size_pt, bold, bold_cs)


# ===========================================================================
# Paragraph-level formatting
# ===========================================================================

# Alignment value to XML string mapping
ALIGN_TO_XML = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "dist",
}


def set_para_format(
    paragraph,
    alignment=None,
    first_line_indent_pt=None,
    left_indent_pt=None,
    space_before_pt=None,
    space_after_pt=None,
    line_spacing=None,
):
    """Set paragraph formatting properties."""
    pPr = get_or_add_pPr(paragraph)

    # Alignment
    if alignment is not None:
        set_child(pPr, "w:jc", ALIGN_TO_XML.get(alignment, "left"))

    # Indentation
    ind = pPr.find(qn("w:ind"))
    if first_line_indent_pt is not None or left_indent_pt is not None:
        if ind is None:
            ind = OxmlElement("w:ind")
            # Insert before any existing spacing
            spacing = pPr.find(qn("w:spacing"))
            if spacing is not None:
                pPr.insert(list(pPr).index(spacing), ind)
            else:
                pPr.append(ind)
        if first_line_indent_pt is not None:
            ind.set(qn("w:firstLine"), str(int(first_line_indent_pt * 20)))
        if left_indent_pt is not None:
            ind.set(qn("w:left"), str(int(left_indent_pt * 20)))
    else:
        if ind is not None:
            # Remove indent attributes but keep the element
            for attr in ("w:firstLine", "w:left", "w:right", "w:hanging"):
                try:
                    del ind.attrib[qn(attr)]
                except KeyError:
                    pass

    # Spacing
    spacing = pPr.find(qn("w:spacing"))
    if space_before_pt is not None or space_after_pt is not None or line_spacing is not None:
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        if space_before_pt is not None:
            spacing.set(qn("w:before"), str(int(space_before_pt * 20)))
        if space_after_pt is not None:
            spacing.set(qn("w:after"), str(int(space_after_pt * 20)))
        if line_spacing is not None:
            spacing.set(qn("w:line"), str(int(line_spacing * 240)))
            spacing.set(qn("w:lineRule"), "auto")


# ===========================================================================
# Style-level formatting
# ===========================================================================


def format_normal_style(doc):
    """Modify the Normal style to GB/T standard."""
    style = doc.styles["Normal"]

    # Modify rPr
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.append(rPr)

    set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None, bold_cs=None)

    # Remove bold from style if present
    b = rPr.find(qn("w:b"))
    bCs = rPr.find(qn("w:bCs"))
    if b is not None:
        rPr.remove(b)
    if bCs is not None:
        rPr.remove(bCs)

    # Modify pPr
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style.element.append(pPr)

    # Remove spacing after
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        # Set after to 0
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")  # single line spacing
        spacing.set(qn("w:lineRule"), "auto")
    else:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)


def format_doc_defaults(doc):
    """Set document defaults for fonts."""
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_elem.insert(0, doc_defaults)

    # rPrDefault
    rPr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rPr_default is None:
        rPr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0 if doc_defaults[0].tag != qn("w:pPrDefault") else 1, rPr_default)

    rPr = rPr_default.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPr_default.append(rPr)

    set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)


# ===========================================================================
# Page setup
# ===========================================================================


def format_page_setup(doc):
    """Set page to A4 with standard GB/T margins."""
    for section in doc.sections:
        # A4: 210mm x 297mm
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT

        # Standard margins
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(21)


# ===========================================================================
# Paragraph classification
# ===========================================================================


def is_numbered_heading(text):
    """Check if text starts with a numbered heading prefix.

    Matches patterns like:
    - '1 范围' (main chapters)
    - '3.1 分析天平' (sub-chapters)
    - 'A.1 磷酸盐缓冲溶液' (appendix sections)
    - 'B.2.1 成分' (appendix sub-sections)
    """
    # Standard numbered headings: 1, 1.1, 2.3.4, etc.
    if re.match(r"^\d+(\.\d+)*\s+\S", text):
        return True

    # Appendix numbered headings: A.1, B.2.1, etc.
    if re.match(r"^[A-Z]\.\d+(\.\d+)*\s+\S", text):
        return True

    return False


def is_method_heading(text):
    """Check if text is a method heading like '第一法 密度瓶法'."""
    return bool(re.match(r"^第[一二三四五六七八九十百]+法\s", text))


def is_appendix_heading(text):
    """Check if text is an appendix heading.

    Matches patterns like:
    - '附录 A'
    - '附录A'
    - '附录 B 培养基和试剂'
    """
    return bool(re.match(r"^附录\s*[A-Z]", text))


def find_cover_start(paragraphs):
    """Find the paragraph index where the cover page begins.

    Scans for '中华人民共和国国家标准' which reliably marks the cover start.
    Returns the index, or 0 if not found (assumes cover starts at P[0]).
    """
    for i, p in enumerate(paragraphs[:30]):
        t = p.text.strip()
        if "中华人民共和国国家标准" in t:
            return i
    return 0


def _is_standard_name_line(text, index, paragraphs):
    """Check if text looks like a repeated standard name (e.g. '橄榄油、油橄榄果渣油').

    Standard name lines are typically short, contain Chinese product names,
    and appear before '1  范围' heading.
    """
    # Check if next non-empty paragraph starts with "1  " (first chapter)
    for j in range(index + 1, min(index + 5, len(paragraphs))):
        next_text = paragraphs[j].text.strip()
        if next_text:
            # Only match real chapter headings like "1  范围" (number + 2+ spaces + Chinese title)
            # Exclude variable definitions like "1 000" or "1.5 mL"
            if re.match(r"^\d+\s{2,}[\u4e00-\u9fff]", next_text):
                return True
            break
    return False


def _find_preface_index(paragraphs, cover_start):
    """Dynamically find the index of '前言'/'前 言' after the cover block.

    Cover block has variable length (3-6 paragraphs depending on whether
    the document has '代替' line, English title, etc.). This function
    scans forward from cover_start+3 to locate the preface title.

    Returns the index of '前言', or -1 if not found within 10 paragraphs.
    """
    for i in range(cover_start + 3, min(cover_start + 12, len(paragraphs))):
        t = paragraphs[i].text.strip()
        if t in ("前言", "前 言"):
            return i
    return -1


def _has_reached_chapter_1(index, paragraphs):
    """Check if we've reached or passed the first chapter heading '1  范围'.

    Scans backward from current index to see if we've already encountered
    a numbered heading starting with '1 ' (like '1  范围').
    """
    for i in range(index - 1, max(0, index - 50), -1):  # Scan backward up to 50 paragraphs
        t = paragraphs[i].text.strip()
        if not t:
            continue
        if re.match(r"^1\s+\S", t):
            return True
        # Don't stop at other headings; keep searching for '1 '
    return False


def classify_paragraph(text, index, paragraphs=None, cover_start=0):
    """Classify a paragraph by its content and position.

    After date merge, cover has 5 items starting at cover_start.
    Preface starts at cover_start + 5.

    Returns one of: 'cover', 'preface_title', 'preface_body',
    'body_title', 'method_heading', 'appendix_heading',
    'heading', 'body', 'formula', 'table_caption', 'empty'
    """
    text = text.strip()

    if not text:
        return "empty"

    # Pre-cover content (e.g. amendment notices)
    if index < cover_start:
        # Detect titles: GB/T standard number lines or Chinese numeral sections
        if re.match(r"^GB[/T ]*\s*\d", text) or re.match(r"^[一二三四五六七八九十]+[、．.]", text):
            return "pre_cover_title"
        return "pre_cover"

    # Cover block: paragraphs from cover_start up to (but not including) preface title
    preface_idx = _find_preface_index(paragraphs, cover_start)
    if preface_idx >= 0 and index < preface_idx:
        return "cover"
    if preface_idx < 0 and index < cover_start + 5:
        return "cover"

    # Preface title
    if index == preface_idx:
        return "preface_title"

    # Preface body: paragraphs after preface title until standard name or first heading
    if preface_idx >= 0 and index > preface_idx:
        # Check if this looks like a standard name line (short, no punctuation ending)
        # and is followed by "1  ..." heading or is the standard name
        # IMPORTANT: Exclude actual numbered headings like "1  范围"
        if len(text) < 30 and not text.endswith("。") and not text.endswith("；"):
            # Only treat as body_title if it's NOT a numbered heading
            if not is_numbered_heading(text) and not is_method_heading(text) and not is_appendix_heading(text):
                if re.match(r"^1\s+\S", text) or _is_standard_name_line(text, index, paragraphs):
                    return "body_title"

        # Preface body continues until we hit body_title or first numbered heading
        if (
            not is_numbered_heading(text)
            and not is_method_heading(text)
            and not is_appendix_heading(text)
            and not text.startswith("1  ")
            and not re.match(r"^表\s+[A-Z]?\.?\d+", text)
            and not re.match(r"^图\s*\d+", text)
        ):
            # Check if we're still in preface region (before first chapter)
            if not _has_reached_chapter_1(index, paragraphs):
                return "preface_body"

    # Method headings like "第一法 密度瓶法"
    if is_method_heading(text):
        return "method_heading"

    # Appendix headings
    if is_appendix_heading(text):
        return "appendix_heading"

    # Table captions (MUST be before numbered headings check)
    if re.match(r"^表\s+[A-Z]?\.?\d+", text):
        return "table_caption"

    # Figure captions (MUST be before numbered headings check)
    if re.match(r"^图\s*\d+", text):
        return "figure_caption"

    # Numbered headings
    if is_numbered_heading(text):
        return "heading"

    # Formula lines (LaTeX or equation references)
    if text.startswith("$$") or re.match(r"^[A-Z] = ", text):
        return "formula"

    # Everything else is body text
    return "body"


# ===========================================================================
# Format application
# ===========================================================================


def apply_cover_format(paragraph, text, para_idx):
    """Apply cover block formatting: 16pt bold, left aligned.

    Uses content-based detection to determine paragraph type within cover:
    - '中华人民共和国国家标准' → 宋体
    - GB/T/GB number line → TNR
    - Date line (YYYY-XX-XX) → 黑体
    - Chinese title (everything else before preface) → 黑体
    - Publishing institution → 宋体
    """
    text_stripped = text.strip()
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        if "国家标准" in text_stripped:
            # "中华人民共和国国家标准"
            set_run_font(rPr, east_asia="宋体", latin="宋体", size_pt=16.0, bold=True, bold_cs=True)
        elif re.match(r"^GB[/T\s]*\d", text_stripped):
            # Standard number line (GB/T 11856.2-2023, GB 8538-2022, etc.)
            set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=16.0, bold=True, bold_cs=True)
        elif re.search(r"\d{4}[-/]\d{2}[-/]\d{2}", text_stripped):
            # Date line (2022-06-30发布 2022-12-30实施)
            set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
        elif "发布" in text_stripped or "实施" in text_stripped:
            # Alternative date format
            set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
        elif re.match(r"^(代替|替代)", text_stripped):
            # "代替" line
            set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
        elif re.match(r"^[A-Z][a-z]+", text_stripped) and len(text_stripped) > 20:
            # English title line
            set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=16.0, bold=True, bold_cs=True)
        elif para_idx == 0 or "国家标准" in text_stripped:
            # First line or contains "国家标准"
            set_run_font(rPr, east_asia="宋体", latin="宋体", size_pt=16.0, bold=True, bold_cs=True)
        else:
            # Default: Chinese title → 黑体
            set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)

    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_preface_title_format(paragraph):
    """Format preface title: 16pt 黑体, centered."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_body_title_format(paragraph):
    """Format body title: 16pt 黑体, centered."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_pre_cover_title_format(paragraph):
    """Format pre-cover title: 16pt 宋体, centered."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=16.0, bold=True, bold_cs=True)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_pre_cover_format(paragraph):
    """Format pre-cover body: 10.5pt 宋体+TNR, JUSTIFY, 21pt indent."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=21.0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_heading_format(paragraph):
    """Format numbered heading: 10.5pt not bold, justified, no indent."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False, bold_cs=False)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_method_heading_format(paragraph):
    """Format method heading: 10.5pt not bold, centered."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False, bold_cs=False)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_appendix_heading_format(paragraph):
    """Format appendix heading: not bold, justified (两端对齐).

    Appendix headings like '附录 A  培养基和试剂' should be justified,
    matching the GB/T standard requirement.
    """
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False, bold_cs=False)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_body_format(paragraph):
    """Format body text: 10.5pt, JUSTIFY, first-line indent 21pt."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=21.0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def normalize_body_spaces(paragraph):
    """Normalize spaces in body text: collapse multi-spaces, strip leading/trailing.

    Preserves the two-space separator in numbered headings (already handled
    by fix_heading_spacing). Only applied to body paragraphs.
    """
    text = paragraph.text
    if not text:
        return

    # Strip leading/trailing, collapse 2+ spaces to single
    normalized = " ".join(text.split())

    if normalized == text:
        return  # No change needed

    runs = paragraph.runs
    if runs:
        runs[0].text = normalized
        for r in runs[1:]:
            r.text = ""


def fix_heading_spacing(paragraph):
    """Ensure two spaces between heading number and text.

    E.g. '1 范围' -> '1  范围', '3.1 分析天平' -> '3.1  分析天平'.
    """
    text = paragraph.text.strip()
    match = re.match(r"^(\d+(?:\.\d+)*)(\s+)(.*)", text)
    if not match:
        return

    num = match.group(1)
    rest = match.group(3)
    current_spaces = match.group(2)

    if current_spaces == "  ":
        return  # Already has two spaces

    new_text = num + "  " + rest
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""


def apply_table_caption_format(paragraph):
    """Format table caption: centered, not bold."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False, bold_cs=False)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_figure_caption_format(paragraph):
    """Format figure caption: centered, not bold."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False, bold_cs=False)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


def apply_formula_format(paragraph):
    """Format formula: centered like body."""
    for run in paragraph.runs:
        rPr = get_or_add_rPr(run)
        set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
    set_para_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
        space_before_pt=0,
        space_after_pt=0,
        line_spacing=1.0,
    )


# ===========================================================================
# Table formatting
# ===========================================================================


def format_tables(doc):
    """Format all tables with GB/T standards."""
    for table in doc.tables:
        # Set table style for borders
        try:
            table.style = doc.styles["Table Grid"]
        except KeyError:
            pass

        # Center the table
        tbl = table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Set table alignment to center
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "center")

        # Format header row (first row) - NOT bold
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        rPr = get_or_add_rPr(run)
                        set_run_font(
                            rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False, bold_cs=False
                        )


def format_images(doc):
    """Center all images/pictures in the document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for paragraph in doc.paragraphs:
        # Check if paragraph contains an image
        if paragraph._element.xpath(".//pic:pic"):
            # Center the paragraph containing the image
            pPr = paragraph._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                paragraph._element.insert(0, pPr)

            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            jc.set(qn("w:val"), "center")


# ===========================================================================
# Main processing
# ===========================================================================


def merge_date_paragraphs(doc):
    """Merge consecutive date paragraphs like '2016-08-31 发布' + '2017-03-01 实施'
    into a single paragraph: '2016-08-31发布 2017-03-01实施'.

    Scans dynamically rather than relying on fixed indices.
    """
    paragraphs = doc.paragraphs
    if len(paragraphs) < 5:
        return

    # Scan for adjacent paragraphs with 发布 and 实施
    for i in range(len(paragraphs) - 1):
        p1_text = paragraphs[i].text.strip()
        p2_text = paragraphs[i + 1].text.strip()

        if p1_text.endswith("发布") and p2_text.endswith("实施"):
            date_pub = p1_text.replace("发布", "").strip()
            date_impl = p2_text.replace("实施", "").strip()
            merged = f"{date_pub}发布 {date_impl}实施"

            # Update first paragraph with merged text
            p1 = paragraphs[i]
            for run in p1.runs:
                run.text = ""
            if p1.runs:
                p1.runs[0].text = merged
            else:
                p1.add_run(merged)

            # Remove second paragraph from the document
            p2_elem = paragraphs[i + 1]._element
            p2_elem.getparent().remove(p2_elem)

            print(f"    Merged dates -> '{merged}'")
            return  # Only merge first occurrence


def merge_split_headings(doc):
    """Merge split heading paragraphs into one line.

    Handles cases where a numbered heading (like '3.1.1' or '附录 A') is on one line
    and its descriptive text is on the next line. Merges them with two spaces.

    Examples:
    - '3.1.1' + '白兰地 brandy' -> '3.1.1  白兰地 brandy'
    - '附录 A' + '培养基和试剂' -> '附录 A  培养基和试剂'
    """
    paragraphs = doc.paragraphs
    if len(paragraphs) < 2:
        return

    merged_count = 0
    i = 0
    while i < len(paragraphs) - 1:
        p1_text = paragraphs[i].text.strip()
        p2_text = paragraphs[i + 1].text.strip()

        should_merge = False

        # Case 1: Appendix heading like '附录 A'
        if re.match(r"^附录\s*[A-Z]$", p1_text) and p2_text:
            if (
                not p2_text.startswith("附录")
                and not p2_text.startswith("前言")
                and not p2_text.startswith("参考文献")
                and not re.match(r"^\d+\s+", p2_text)
            ):
                should_merge = True

        # Case 2: Numbered heading like '3.1.1', 'A.1', 'B.2.1'
        elif re.match(r"^(\d+(\.\d+)*|[A-Z]\.\d+(\.\d+)*)$", p1_text) and p2_text:
            # Check if next paragraph is descriptive text (not another heading)
            if (
                not re.match(r"^(\d+(\.\d+)*|[A-Z]\.\d+)", p2_text)
                and not p2_text.startswith("附录")
                and not p2_text.startswith("前言")
                and not p2_text.startswith("参考文献")
            ):
                should_merge = True

        if should_merge:
            # Merge with two spaces separator
            merged = f"{p1_text}  {p2_text}"

            # Update first paragraph with merged text
            p1 = paragraphs[i]
            for run in p1.runs:
                run.text = ""
            if p1.runs:
                p1.runs[0].text = merged
            else:
                p1.add_run(merged)

            # Remove second paragraph from the document
            p2_elem = paragraphs[i + 1]._element
            p2_elem.getparent().remove(p2_elem)

            merged_count += 1
            i += 1  # Skip the merged paragraph
            continue

        i += 1

    if merged_count > 0:
        print(f"    Merged {merged_count} split headings")


def remove_toc(doc):
    """Detect and remove TOC (目次/目录) section between cover and preface.

    Finds paragraphs containing '目次' or '目录', then removes all
    paragraphs from that point until '前言' or the first numbered heading.
    """
    paragraphs = doc.paragraphs
    toc_start = -1

    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if toc_start < 0:
            if t in ("目  次", "目次", "目  录", "目录"):
                toc_start = i
        elif t.startswith("前") and "言" in t:
            # Found end of TOC - remove all TOC paragraphs from XML body
            toc_end = i
            body = doc.element.body
            for j in range(toc_start, toc_end):
                p_elem = paragraphs[j]._element
                body.remove(p_elem)
            return toc_end - toc_start  # number of paras removed

    return 0


def process_document(input_path, output_path):
    """Process a DOCX file, applying GB/T formatting."""
    print(f"Opening: {input_path}")
    doc = Document(input_path)

    # 1. Page setup
    print("  Setting page setup (A4, standard margins)...")
    format_page_setup(doc)

    # 2. Document defaults
    print("  Setting document defaults (宋体+TNR, 10.5pt)...")
    format_doc_defaults(doc)

    # 3. Normal style
    print("  Setting Normal style...")
    format_normal_style(doc)

    # 3.5 Merge date paragraphs (e.g. "2016-08-31 发布" + "2017-03-01 实施")
    print("  Merging date paragraphs...")
    merge_date_paragraphs(doc)

    # 3.6 Remove TOC (目次/目录) if present
    print("  Removing TOC (if present)...")
    removed = remove_toc(doc)
    if removed:
        print(f"    Removed {removed} TOC paragraphs")

    # 3.7 Merge split headings (e.g. '3.1.1' + '白兰地 brandy' -> '3.1.1  白兰地 brandy')
    print("  Merging split headings...")
    merge_split_headings(doc)

    # 3.8 Find actual cover start (may be offset by amendment content)
    cover_start = find_cover_start(doc.paragraphs)
    if cover_start > 0:
        print(f"  Cover starts at P[{cover_start}] (pre-cover content detected)")

    # 4. Process each paragraph
    print("  Processing paragraphs...")
    stats = {
        "pre_cover_title": 0,
        "pre_cover": 0,
        "cover": 0,
        "preface_title": 0,
        "preface_body": 0,
        "body_title": 0,
        "method_heading": 0,
        "appendix_heading": 0,
        "heading": 0,
        "body": 0,
        "formula": 0,
        "table_caption": 0,
        "figure_caption": 0,
        "empty": 0,
        "total": 0,
    }

    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        cls = classify_paragraph(text, idx, doc.paragraphs, cover_start)
        stats[cls] = stats.get(cls, 0) + 1
        stats["total"] += 1

        # Apply formatting based on classification
        if cls == "empty":
            # Clear spacing for empty paragraphs (often formula image placeholders)
            set_para_format(para, space_before_pt=0, space_after_pt=0)
        elif cls == "cover":
            apply_cover_format(para, text, idx - cover_start)
        elif cls == "pre_cover_title":
            apply_pre_cover_title_format(para)
        elif cls == "pre_cover":
            apply_pre_cover_format(para)
            # normalize_body_spaces(para)  # Disabled: preserves original spacing for formulas/tables
        elif cls == "preface_title":
            apply_preface_title_format(para)
        elif cls == "preface_body":
            apply_body_format(para)
            # normalize_body_spaces(para)  # Disabled: preserves original spacing for formulas/tables
        elif cls == "body_title":
            apply_body_title_format(para)
        elif cls == "method_heading":
            apply_method_heading_format(para)
        elif cls == "appendix_heading":
            apply_appendix_heading_format(para)
        elif cls == "heading":
            fix_heading_spacing(para)
            apply_heading_format(para)
        elif cls == "formula":
            apply_formula_format(para)
        elif cls == "table_caption":
            apply_table_caption_format(para)
        elif cls == "figure_caption":
            apply_figure_caption_format(para)
        elif cls == "body":
            apply_body_format(para)
            # normalize_body_spaces(para)  # Disabled: preserves original spacing for formulas/tables

    # 5. Format tables
    print("  Formatting tables...")
    format_tables(doc)

    # 6. Format images (center them)
    print("  Formatting images...")
    format_images(doc)

    # 7. Save
    print(f"  Saving to: {output_path}")
    doc.save(output_path)

    # Print stats
    print()
    print("  Paragraph classification:")
    for cls in sorted(stats.keys()):
        if cls != "total":
            print(f"    {cls}: {stats[cls]}")
    print(f"    total: {stats['total']}")

    print()
    print("Done!")


def main():
    parser = argparse.ArgumentParser(description="Apply GB/T standard formatting to a DOCX file")
    parser.add_argument("input", help="Input DOCX file path")
    parser.add_argument("--output", "-o", help="Output DOCX file path")
    args = parser.parse_args()

    input_path = args.input
    if not os.path.exists(input_path):
        sys.stderr.write(f"ERROR: File not found: {input_path}\n")
        sys.exit(1)

    if args.output:
        output_path = args.output
    else:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_formatted{ext}"

    process_document(input_path, output_path)


if __name__ == "__main__":
    main()
