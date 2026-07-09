"""全面分析排版结果 DOCX 是否符合 GB/T 1.1 规范"""

import re

from docx import Document
from docx.oxml.ns import qn

DOC_PATH = r"c:\Users\Spzx\Downloads\GB_T14294-2008_排版结果.docx"

doc = Document(DOC_PATH)

print("=" * 80)
print("  GB/T 1.1 排版结果全面分析")
print("=" * 80)

# =============================================================================
# 1. 页面设置
# =============================================================================
print("\n[1] 页面设置")
print("-" * 40)
for si, section in enumerate(doc.sections):
    w_mm = section.page_width / 360000 * 25.4  # EMU to mm
    h_mm = section.page_height / 360000 * 25.4
    t = section.top_margin / 360000 * 25.4
    b = section.bottom_margin / 360000 * 25.4
    l = section.left_margin / 360000 * 25.4
    r = section.right_margin / 360000 * 25.4
    is_a4 = abs(w_mm - 210) < 2 and abs(h_mm - 297) < 2
    print(f"  Section {si}: {w_mm:.0f}×{h_mm:.0f}mm {'[OK] A4' if is_a4 else '[FAIL] 非A4'}")
    print(f"    上边距={t:.0f}mm, 下边距={b:.0f}mm, 左边距={l:.0f}mm, 右边距={r:.0f}mm")

# =============================================================================
# 2. docDefaults
# =============================================================================
print("\n[2] 文档默认字体 (docDefaults)")
print("-" * 40)
styles_elem = doc.styles.element
docDefaults = styles_elem.find(qn("w:docDefaults"))
if docDefaults is not None:
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is not None:
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            rFonts = rPr.find(qn("w:rFonts"))
            ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else "?"
            asc = rFonts.get(qn("w:ascii")) if rFonts is not None else "?"
            sz_val = sz.get(qn("w:val")) if sz is not None else "?"
            half_pt = int(sz_val)
            pt_val = half_pt / 2 if sz_val != "?" else "?"
            ea_ok = ea == "宋体"
            asc_ok = asc == "Times New Roman"
            sz_ok = pt_val == 10.5
            print(f"  rFonts: eastAsia={ea} {'[OK]' if ea_ok else '[FAIL]应为宋体'}")
            print(f"  rFonts: ascii={asc} {'[OK]' if asc_ok else '[FAIL]应为Times New Roman'}")
            print(f"  sz: {half_pt} 半磅 = {pt_val}pt {'[OK]' if sz_ok else '[FAIL]应为10.5pt'}")
        else:
            print("  [FAIL] 无 rPr 元素")
    else:
        print("  [FAIL] 无 rPrDefault")
else:
    print("  [FAIL] 无 docDefaults 元素")

# =============================================================================
# 3. 逐段落详细分析
# =============================================================================
print("\n[3] 段落详细分析 (按结构分组)")
print("-" * 40)

# 先找关键位置
cover_start = -1
preface_idx = -1
for i, p in enumerate(doc.paragraphs[:50]):
    t = p.text.strip()
    if "中华人民共和国国家标准" in t and cover_start < 0:
        cover_start = i
    if t in ("前言", "前 言") and preface_idx < 0:
        preface_idx = i

print(f"  cover_start=P[{cover_start}], preface_idx=P[{preface_idx}]")

# 分类段落并统计
issues = []
stats = {}

# 已找到的 section 标记
in_preface = False
in_body = False
in_appendix = False
passed_first_heading = False

for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    pPr = para._element.find(qn("w:pPr"))
    jc_val = "?"
    indent_val = "?"
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            first = ind.get(qn("w:firstLine"))
            if first:
                indent_val = f"{int(first)/20:.0f}pt"
            else:
                indent_val = "0"

    # 获取字体信息
    rprs = []
    for r in para.runs[:1]:
        rpr = r._element.find(qn("w:rPr"))
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            rFonts = rpr.find(qn("w:rFonts"))
            b = rpr.find(qn("w:b"))
            ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else None
            asc = rFonts.get(qn("w:ascii")) if rFonts is not None else None
            sz_val = int(sz.get(qn("w:val"))) if sz is not None else None
            sz_pt = sz_val / 2 if sz_val else None
            bold = "B" if b is not None else ""
            rprs.append(f"{ea}/{asc} {sz_pt}pt {bold}")

    font_info = ", ".join(rprs) if rprs else "NO_RUN"
    short_text = text[:80].replace("\n", " ")

    # --- 判断段落角色并检查格式 ---
    role = None
    expected = {}

    # 封面之前 (index < cover_start)
    if cover_start >= 0 and idx < cover_start:
        if re.match(r"^GB[/T ]*\s*\d", text) or re.match(r"^[一二三四五六七八九十]+[、．.]", text):
            role = "pre_cover_title"
            expected = {"font": "宋体/TNR", "size": "16pt", "bold": True, "align": "center", "indent": "0pt"}
        else:
            role = "pre_cover"
            expected = {"font": "宋体/TNR", "size": "10.5pt", "bold": None, "align": "both", "indent": "21pt"}

    # 封面区域
    elif cover_start >= 0 and preface_idx >= 0 and idx < preface_idx:
        role = "cover"
        expected = None  # 封面格式多变，不检查

    elif preface_idx < 0 and cover_start >= 0 and idx < cover_start + 5:
        role = "cover"
        expected = None

    # 前言标题
    elif idx == preface_idx:
        role = "preface_title"
        expected = {"font": "黑体/黑体", "size": "16pt", "bold": True, "align": "center", "indent": "0pt"}

    # 前言正文或正文
    elif preface_idx >= 0 and idx > preface_idx:
        t = text
        # 检查是否是标准名称 (body_title)
        if len(t) < 30 and not t.endswith("。") and not t.endswith("；"):
            if (
                not re.match(r"^\d+(\.\d+)*\s+\S", t)
                and not re.match(r"^第[一二三四五六七八九十百]+法\s", t)
                and not re.match(r"^附录\s*[A-Z]", t)
            ):
                if re.match(r"^1\s+\S", t):
                    role = "body_title"
                    expected = {"font": "黑体/黑体", "size": "16pt", "bold": True, "align": "center", "indent": "0pt"}
                    passed_first_heading = True

        if role is None:
            # 检查是否是编号标题 (heading)
            if re.match(r"^\d+(\.\d+)*\s+\S", t):
                role = "heading"
                expected = {"font": "宋体/TNR", "size": "10.5pt", "bold": False, "align": "both", "indent": "0pt"}
                if re.match(r"^1\s+\S", t):
                    passed_first_heading = True
            # 附录标题
            elif re.match(r"^附录\s*[A-Z]", t):
                role = "appendix_heading"
                expected = {"font": "宋体/TNR", "size": "10.5pt", "bold": False, "align": "both", "indent": "0pt"}
                in_appendix = True
            # 表格标题
            elif re.match(r"^表\s+[A-Z]?\.?\d+", t):
                role = "table_caption"
                expected = {"font": "宋体/TNR", "size": "10.5pt", "bold": False, "align": "center", "indent": "0pt"}
            # 图片标题
            elif re.match(r"^图\s*\d+", t):
                role = "figure_caption"
                expected = {"font": "宋体/TNR", "size": "10.5pt", "bold": False, "align": "center", "indent": "0pt"}
            # 正文
            else:
                role = "body"
                expected = {"font": "宋体/TNR", "size": "10.5pt", "bold": None, "align": "both", "indent": "21pt"}

    # 统计
    if role:
        stats[role] = stats.get(role, 0) + 1
    else:
        stats["other"] = stats.get("other", 0) + 1

    # 检查格式是否符合预期
    has_issue = False
    checks = []
    if expected:
        # 获取实际字体
        actual_ea = None
        actual_asc = None
        actual_sz = None
        actual_bold = False
        for r in para.runs[:1]:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is not None:
                sz_el = rpr.find(qn("w:sz"))
                rFonts_el = rpr.find(qn("w:rFonts"))
                b_el = rpr.find(qn("w:b"))
                actual_ea = rFonts_el.get(qn("w:eastAsia")) if rFonts_el is not None else None
                actual_asc = rFonts_el.get(qn("w:ascii")) if rFonts_el is not None else None
                actual_sz = int(sz_el.get(qn("w:val"))) / 2 if sz_el is not None else None
                actual_bold = b_el is not None

        # 检查尺寸
        exp_sz = expected.get("size")
        if exp_sz:
            exp_num = float(exp_sz.replace("pt", ""))
            if actual_sz is not None and abs(actual_sz - exp_num) < 0.1:
                checks.append(f"[OK] sz={actual_sz}pt")
            else:
                checks.append(f"[FAIL] sz={actual_sz}pt (应为{exp_sz})")

        # 检查粗体
        exp_bold = expected.get("bold")
        if exp_bold is not None:
            if exp_bold == actual_bold:
                checks.append(f"[OK] bold={actual_bold}")
            else:
                checks.append(f"[FAIL] bold={actual_bold} (应为{exp_bold})")

        # 检查对齐
        exp_align = expected.get("align")
        if exp_align:
            actual_jc = jc_val
            if actual_jc == exp_align:
                checks.append(f"[OK] align={actual_jc}")
            else:
                checks.append(f"[FAIL] align={actual_jc} (应为{exp_align})")

        # 检查缩进
        exp_indent = expected.get("indent")
        if exp_indent:
            actual_indent_str = indent_val
            if actual_indent_str == exp_indent:
                checks.append(f"[OK] indent={actual_indent_str}")
            else:
                checks.append(f"[FAIL] indent={actual_indent_str} (应为{exp_indent})")

        # 检查字体
        exp_font = expected.get("font")
        if exp_font:
            ea_asc = f"{actual_ea}/{actual_asc}" if actual_ea and actual_asc else "?"
            font_ok = False
            if exp_font == "宋体/TNR":
                font_ok = actual_ea == "宋体" and actual_asc == "Times New Roman"
            elif exp_font == "黑体/黑体":
                font_ok = actual_ea == "黑体" and actual_asc == "黑体"
            if font_ok:
                checks.append(f"[OK] font={ea_asc}")
            else:
                checks.append(f"[FAIL] font={ea_asc} (应为{exp_font})")

        has_issue = any(c.startswith("[FAIL]") for c in checks)
        if has_issue:
            issues.append(
                {
                    "idx": idx,
                    "role": role,
                    "text": short_text,
                    "checks": checks,
                    "font_info": font_info,
                    "jc": jc_val,
                    "indent": indent_val,
                }
            )

    # 打印关键段落
    is_key = any(
        [
            text.startswith("中华人民共和国"),
            text.startswith("GB/T"),
            text in ("前言", "前 言"),
            re.match(r"^1\s+\S", text),
            text == "组合式空调机组",
            text.startswith("Central-station"),
            text.startswith("附录"),
        ]
    )
    if is_key:
        has_issue_flag = has_issue
        issue_mark = " [WARN]" if has_issue_flag else ""
        print(f"\n  P[{idx:3d}] [{role or '??':16s}] {font_info}{issue_mark}")
        print(f"         align={jc_val} indent={indent_val}")
        print(f'        "{short_text}"')
        if checks:
            print(f"         {' '.join(checks)}")

# =============================================================================
# 4. 统计汇总
# =============================================================================
print("\n\n[4] 分类统计")
print("-" * 40)
for role in [
    "pre_cover_title",
    "pre_cover",
    "cover",
    "preface_title",
    "preface_body",
    "body_title",
    "method_heading",
    "appendix_heading",
    "heading",
    "body",
    "table_caption",
    "figure_caption",
    "formula",
    "other",
]:
    cnt = stats.get(role, 0)
    if cnt > 0:
        print(f"  {role:20s}: {cnt:4d}")

print(f"\n  {'TOTAL':20s}: {sum(stats.values()):4d}")

# =============================================================================
# 5. 问题汇总
# =============================================================================
print(f"\n\n[5] 问题汇总 ({len(issues)} 个问题段落)")
print("-" * 40)
if issues:
    for iss in issues:
        print(f"\n  P[{iss['idx']:3d}] [{iss['role']:16s}]")
        print(f"         text: \"{iss['text'][:60]}\"")
        print(f"         actual: font={iss['font_info']} align={iss['jc']} indent={iss['indent']}")
        for c in iss["checks"]:
            print(f"         {c}")
else:
    print("  [OK] 所有检查段落格式符合预期！")

# =============================================================================
# 6. 表格检查
# =============================================================================
print(f"\n\n[6] 表格检查 ({len(doc.tables)} 个表格)")
print("-" * 40)
table_issues = []
for ti, table in enumerate(doc.tables):
    tblPr = table._element.find(qn("w:tblPr"))
    style_val = "?"
    jc_val = "?"
    if tblPr is not None:
        jc = tblPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            style_val = tblStyle.get(qn("w:val"))

    # 表头字体
    header_font = "?"
    if len(table.rows) > 0:
        cells = table.rows[0].cells
        if cells:
            for para in cells[0].paragraphs[:1]:
                for r in para.runs[:1]:
                    rpr = r._element.find(qn("w:rPr"))
                    if rpr is not None:
                        sz = rpr.find(qn("w:sz"))
                        rFonts = rpr.find(qn("w:rFonts"))
                        ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else "?"
                        sz_val = int(sz.get(qn("w:val"))) / 2 if sz is not None else "?"
                        header_font = f"{ea} {sz_val}pt"

    issues_t = []
    if style_val != "TableGrid" and style_val != "Table Grid":
        issues_t.append(f"[FAIL] style={style_val} (应为Table Grid)")
    else:
        issues_t.append(f"[OK] style={style_val}")
    if jc_val != "center":
        issues_t.append(f"[FAIL] table align={jc_val} (应为center)")
    else:
        issues_t.append(f"[OK] table align={jc_val}")
    if "宋体" not in header_font:
        issues_t.append(f"[FAIL] header font={header_font} (应为宋体)")
    else:
        issues_t.append(f"[OK] header font={header_font}")

    header_text = ""
    if len(table.rows) > 0 and table.rows[0].cells:
        header_text = table.rows[0].cells[0].paragraphs[0].text[:50] if table.rows[0].cells[0].paragraphs else ""

    has_bad = any(c.startswith("[FAIL]") for c in issues_t)
    print(f'  Table {ti}: rows={len(table.rows)} cols={len(table.columns)} header="{header_text}"')
    for c in issues_t:
        print(f"    {c}")
    if has_bad:
        table_issues.append({"idx": ti, "header": header_text, "checks": issues_t})

# =============================================================================
# 7. 图片检查
# =============================================================================
print("\n\n[7] 图片检查")
print("-" * 40)
img_count = 0
img_issues = []
for idx, para in enumerate(doc.paragraphs):
    pics = para._element.findall(".//" + qn("pic:pic"))
    if pics:
        img_count += 1
        pPr = para._element.find(qn("w:pPr"))
        jc_val = "?"
        if pPr is not None:
            jc = pPr.find(qn("w:jc"))
            if jc is not None:
                jc_val = jc.get(qn("w:val"))
        if jc_val != "center":
            img_issues.append(f"  Image in P[{idx}]: align={jc_val} [FAIL] 应为center")
print(f"  图片总数: {img_count}")
if img_issues:
    for iss in img_issues:
        print(f"  {iss}")
else:
    print("  [OK] 所有图片已居中")

# =============================================================================
# 8. 特殊检查: 前言后标准名称
# =============================================================================
print("\n\n[8] 特殊检查: 前言后标准名称 (body_title)")
print("-" * 40)
if preface_idx >= 0:
    for idx in range(preface_idx + 1, min(preface_idx + 15, len(doc.paragraphs))):
        t = doc.paragraphs[idx].text.strip()
        if t and len(t) < 40:
            # 检查这个段落
            para = doc.paragraphs[idx]
            pPr = para._element.find(qn("w:pPr"))
            jc_val = "?"
            if pPr is not None:
                jc = pPr.find(qn("w:jc"))
                if jc is not None:
                    jc_val = jc.get(qn("w:val"))
            rprs = []
            for r in para.runs[:1]:
                rpr = r._element.find(qn("w:rPr"))
                if rpr is not None:
                    sz = rpr.find(qn("w:sz"))
                    rFonts = rpr.find(qn("w:rFonts"))
                    b = rpr.find(qn("w:b"))
                    ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else "?"
                    sz_val = int(sz.get(qn("w:val"))) / 2 if sz is not None else "?"
                    bold = "B" if b is not None else ""
                    rprs.append(f"{ea} {sz_val}pt {bold}")
            font = ", ".join(rprs) if rprs else "NO_RUN"

            is_body_title = len(t) < 30 and not t.endswith("。") and not t.endswith("；")
            is_chapter_1 = re.match(r"^1\s+\S", t)
            print(f'  P[{idx:3d}]: align={jc_val} font={font} text="{t[:60]}"')
            print(f"         is_body_title_candidate={is_body_title} is_chapter1={bool(is_chapter_1)}")

# =============================================================================
# 9. 总体评估
# =============================================================================
print("\n\n[9] 总体评估")
print("-" * 40)
total_issues = len(issues) + len(table_issues) + len(img_issues)
if total_issues == 0:
    print("  [OK] 文档格式完全符合 GB/T 1.1 规范！")
else:
    print(f"  [WARN]  发现 {total_issues} 个格式问题:")
    print(f"     段落问题: {len(issues)}")
    print(f"     表格问题: {len(table_issues)}")
    print(f"     图片问题: {len(img_issues)}")
