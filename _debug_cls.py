"""Debug classification for paragraph 28"""

import re

from docx import Document


def is_numbered_heading(text):
    if re.match(r"^\d+(\.\d+)*\s+\S", text):
        return True
    if re.match(r"^[A-Z]\.\d+(\.\d+)*\s+\S", text):
        return True
    return False


def is_standard_name_line(text, index, paragraphs):
    for j in range(index + 1, min(index + 5, len(paragraphs))):
        next_text = paragraphs[j].text.strip()
        if next_text:
            match = re.match(r"^\d+\s{2,}[\u4e00-\u9fff]", next_text)
            print(f"    paragraphs[{j}] = {repr(next_text[:30])}  match={bool(match)}")
            if match:
                return True
            break
    return False


doc = Document("data/output/e91a5193-1962-4bbe-809e-68034b1c7ec3/formatted_test.docx")

# Check around index 28
for idx in range(25, 35):
    text = doc.paragraphs[idx].text.strip()
    pPr = doc.paragraphs[idx]._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    jc_val = "?"
    if pPr is not None:
        jc = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc")
        if jc is not None:
            jc_val = jc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    print(f"[{idx:3d}] jc={jc_val:7s}  text={repr(text[:60])}")

    if idx == 28:
        print(f'  len={len(text)} endswith_period={text.endswith("。")} endswith_semicolon={text.endswith("；")}')
        print(f"  is_numbered={is_numbered_heading(text)}")
        print(f"  is_std_name_line={is_standard_name_line(text, idx, doc.paragraphs)}")
        print(f'  regex_match_1: {re.match(r"^1\s+\S", text)}')
