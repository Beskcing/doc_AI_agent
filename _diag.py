"""Diagnose DOCX paragraph structure"""

from docx import Document
from docx.oxml.ns import qn

doc = Document("data/output/e91a5193-1962-4bbe-809e-68034b1c7ec3/full.docx")

print("=== Original full.docx (first 30 paragraphs) ===")
for i, p in enumerate(doc.paragraphs[:30]):
    style = p.style.name if p.style else "None"
    text = p.text[:80].replace("\n", " ")
    rprs = []
    for r in p.runs[:2]:
        rpr = r._element.find(qn("w:rPr"))
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            rFonts = rpr.find(qn("w:rFonts"))
            ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else None
            latin = rFonts.get(qn("w:ascii")) if rFonts is not None else None
            sz_val = sz.get(qn("w:val")) if sz is not None else "?"
            rprs.append(f"font={ea}/{latin} sz={sz_val}")
    jc_val = "?"
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
    print(f"[{i:3d}] style={style[:15]:15s} jc={jc_val:8s}  text={text}")
    if rprs:
        print(f"      runs: {rprs[0]}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Check images in first few paragraphs
img_count = 0
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall(".//" + qn("w:drawing"))
    if drawings:
        img_count += 1
        if img_count <= 5:
            print(f"  Image in paragraph [{i}]: {p.text[:60]}")
print(f"Total images: {img_count}")
