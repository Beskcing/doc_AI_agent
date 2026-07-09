"""Diagnose formatted DOCX paragraph structure and run formatting"""

from docx import Document
from docx.oxml.ns import qn

print("=== formatted_test.docx (first 25 paragraphs) ===")
doc = Document("data/output/e91a5193-1962-4bbe-809e-68034b1c7ec3/formatted_test.docx")

for i, p in enumerate(doc.paragraphs[:25]):
    text = p.text[:80].replace("\n", " ")
    pPr = p._element.find(qn("w:pPr"))
    jc_val = "?"
    indent_val = "?"
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            first = ind.get(qn("w:firstLine"))
            indent_val = first if first else "0"
    rprs = []
    for r in p.runs[:2]:
        rpr = r._element.find(qn("w:rPr"))
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            rFonts = rpr.find(qn("w:rFonts"))
            b = rpr.find(qn("w:b"))
            ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else None
            ascii_f = rFonts.get(qn("w:ascii")) if rFonts is not None else None
            sz_val = sz.get(qn("w:val")) if sz is not None else "?"
            bold = "B" if b is not None else ""
            rprs.append(f"{ea}/{ascii_f} sz={sz_val}{bold}")
    print(f"[{i:3d}] jc={jc_val:7s} indent={indent_val:5s}  {text}")
    if rprs:
        print(f'      runs: {", ".join(rprs[:2])}')

# Check docDefaults
print("\n=== docDefaults ===")
styles_elem = doc.styles.element
docDefaults = styles_elem.find(qn("w:docDefaults"))
if docDefaults is not None:
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is not None:
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is not None:
            print("  docDefaults rPr: found")
            for child in rPr:
                print(f"    {child.tag}: {dict(child.attrib)}")

# Check page setup
print("\n=== Page Setup ===")
for i, section in enumerate(doc.sections):
    print(f"Section {i}: width={section.page_width}, height={section.page_height}")
    print(f"  margins: top={section.top_margin}, bottom={section.bottom_margin}")
    print(f"           left={section.left_margin}, right={section.right_margin}")

# Check table formatting
print(f"\n=== Tables ({len(doc.tables)}) ===")
for ti, table in enumerate(doc.tables[:3]):
    tblPr = table._element.find(qn("w:tblPr"))
    jc_val = "?"
    style_val = "?"
    if tblPr is not None:
        jc = tblPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            style_val = tblStyle.get(qn("w:val"))
    print(f"Table {ti}: style={style_val}, jc={jc_val}, rows={len(table.rows)}")
    # Check header font
    if len(table.rows) > 0:
        for ci, cell in enumerate(table.rows[0].cells[:2]):
            for pi, para in enumerate(cell.paragraphs[:1]):
                text = para.text[:40]
                rprs = []
                for r in para.runs[:1]:
                    rpr = r._element.find(qn("w:rPr"))
                    if rpr is not None:
                        sz = rpr.find(qn("w:sz"))
                        rFonts = rpr.find(qn("w:rFonts"))
                        ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else None
                        sz_val = sz.get(qn("w:val")) if sz is not None else "?"
                        rprs.append(f"{ea} sz={sz_val}")
                print(f'  Cell[{ci}]: {text} ({", ".join(rprs)})')
