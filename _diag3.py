"""Full diagnostic of formatted output"""

from docx import Document
from docx.oxml.ns import qn

doc = Document("data/output/e91a5193-1962-4bbe-809e-68034b1c7ec3/formatted_test.docx")

print("=== All paragraphs (headings & body) ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text[:100].replace("\n", " ")
    if not text.strip():
        continue

    pPr = p._element.find(qn("w:pPr"))
    jc_val = "?"
    indent_val = "?"
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            first = ind.get(qn("w:firstLine"))
            indent_val = first if first else "0"

    # Check run font
    rprs = []
    for r in p.runs[:1]:
        rpr = r._element.find(qn("w:rPr"))
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            rFonts = rpr.find(qn("w:rFonts"))
            b = rpr.find(qn("w:b"))
            ea = rFonts.get(qn("w:eastAsia")) if rFonts is not None else None
            sz_val = sz.get(qn("w:val")) if sz is not None else "?"
            bold = "B" if b is not None else " "
            rprs.append(f"{ea} sz={sz_val} {bold}")

    font_info = ", ".join(rprs) if rprs else "NO_RUN"

    # Highlight suspicious patterns
    flag = ""
    if jc_val in ("left", "?") and indent_val in ("0", "?"):
        # Could be heading/title without proper formatting
        t = text.strip()
        if len(t) < 40 and not t.endswith("。"):
            # Short text with left/0 - might be a missed heading
            pass

    # Only show paragraphs that look like titles, headings, or are suspicious
    t = text.strip()
    is_heading = any(
        [
            t.startswith("附录"),
            t.startswith("前"),
            "国家标准" in t,
            len(t) < 25 and not t.endswith("。"),
        ]
    )

    if is_heading or len(t) < 30:
        print(f"[{i:3d}] jc={jc_val:7s} indent={indent_val:5s} {font_info} | {text}")

print(f"\n=== Tables ({len(doc.tables)}) ===")
for ti, table in enumerate(doc.tables):
    tblPr = table._element.find(qn("w:tblPr"))
    style_val = "?"
    jc_val = "?"
    if tblPr is not None:
        jc = tblPr.find(qn("w:jc"))
        if jc is not None:
            jc_val = jc.get(qn("w:val"))
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            style_val = tblStyle.get(qn("w:val"))

    # Check header
    header_text = ""
    if len(table.rows) > 0:
        for cell in table.rows[0].cells[:1]:
            header_text = cell.paragraphs[0].text[:60] if cell.paragraphs else ""

    # Check if Table Grid exists
    has_grid = "Table Grid" in [s.name for s in doc.styles]

    print(f'  Table {ti}: style={style_val}, jc={jc_val}, has_TableGrid={has_grid}, header="{header_text}"')
