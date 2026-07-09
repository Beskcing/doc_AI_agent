"""Verify GbtDocxFormatter output after fix"""

import re

from docx import Document
from docx.oxml.ns import qn

doc = Document("data/output/gbt14294_formatted.docx")

# Correct page size conversion: EMU / 36000 = mm
s = doc.sections[0]
w = s.page_width / 36000
h = s.page_height / 36000
t = s.top_margin / 36000
b = s.bottom_margin / 36000
l = s.left_margin / 36000
r = s.right_margin / 36000
print(f"[Page] {w:.0f}x{h:.0f}mm A4={abs(w-210)<1 and abs(h-297)<1}")
print(f"[Margins] top={t:.0f}mm bottom={b:.0f}mm left={l:.0f}mm right={r:.0f}mm")

# Key paragraphs
print()
ok_count = 0
fail_count = 0
checks = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue

    pPr = p._element.find(qn("w:pPr"))
    jc = "?"
    indent_pt = "?"
    if pPr is not None:
        jc_el = pPr.find(qn("w:jc"))
        if jc_el is not None:
            jc = jc_el.get(qn("w:val"))
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            first = ind.get(qn("w:firstLine"))
            if first is not None and first != "0":
                indent_pt = f"{int(first)/20:.0f}pt"
            else:
                indent_pt = "0pt"

    rprs = []
    for r in p.runs[:1]:
        rp = r._element.find(qn("w:rPr"))
        if rp is not None:
            sz_el = rp.find(qn("w:sz"))
            rf = rp.find(qn("w:rFonts"))
            b_el = rp.find(qn("w:b"))
            ea = rf.get(qn("w:eastAsia")) if rf is not None else "?"
            asc = rf.get(qn("w:ascii")) if rf is not None else "?"
            sz_pt = int(sz_el.get(qn("w:val"))) / 2 if sz_el is not None else "?"
            bold = "B" if b_el is not None else "_"
            rprs.append(f"{ea}/{asc} {sz_pt}pt {bold}")
    font = ", ".join(rprs) if rprs else "no_run"

    is_key = any(
        [
            text == "前言",
            text.startswith("中华人民共和国"),
            text.startswith("GB/T"),
            text == "组合式空调机组",
            text.startswith("Central-station"),
            re.match(r"^1\s+\S", text),
            re.match(r"^附录\s*[A-Z]", text),
            text.startswith("2  "),
            text.startswith("3  "),
        ]
    )
    if is_key:
        label = f"P[{i:3d}]"
        print(f"{label} align={jc:6s} indent={indent_pt:5s} {font:35s} | {text[:70]}")

# Table check
print(f"\n[Tables] {len(doc.tables)} tables")
for ti, table in enumerate(doc.tables[:2]):
    tblPr = table._element.find(qn("w:tblPr"))
    style = "?"
    tbl_jc = "?"
    if tblPr is not None:
        s_el = tblPr.find(qn("w:tblStyle"))
        if s_el is not None:
            style = s_el.get(qn("w:val"))
        j_el = tblPr.find(qn("w:jc"))
        if j_el is not None:
            tbl_jc = j_el.get(qn("w:val"))
    hdr = table.rows[0].cells[0].paragraphs[0].text[:30] if table.rows else ""
    ok_style = style in ("TableGrid", "Table Grid")
    ok_align = tbl_jc == "center"
    print(
        f"  T{ti}: style={style} {'[OK]' if ok_style else '[FAIL]'}, jc={tbl_jc} {'[OK]' if ok_align else '[FAIL]'}, header=\"{hdr}\""
    )

# Image check
imgs = sum(1 for p in doc.paragraphs if p._element.findall(".//" + qn("pic:pic")))
print(f"\n[Images] {imgs} images in document")

# Key format checks
print("\n=== Key Format Verification ===")
# Check body_title (P[21] "组合式空调机组")
p21 = doc.paragraphs[21]
p21_jc = p21._element.find(qn("w:pPr")).find(qn("w:jc")).get(qn("w:val"))
p21_rpr = p21.runs[0]._element.find(qn("w:rPr"))
p21_ea = p21_rpr.find(qn("w:rFonts")).get(qn("w:eastAsia"))
p21_sz = int(p21_rpr.find(qn("w:sz")).get(qn("w:val"))) / 2
p21_b = p21_rpr.find(qn("w:b")) is not None
print(f"body_title '组合式空调机组': center={p21_jc=='center'} font={p21_ea} sz={p21_sz}pt bold={p21_b}")
all_ok = p21_jc == "center" and p21_ea == "黑体" and p21_sz == 16.0 and p21_b
print(f"  => {'[OK] ALL CORRECT!' if all_ok else '[FAIL]'}")

# Check heading "1 范围"
p22 = doc.paragraphs[22]
p22_jc = p22._element.find(qn("w:pPr")).find(qn("w:jc")).get(qn("w:val"))
p22_rpr = p22.runs[0]._element.find(qn("w:rPr"))
p22_ea = p22_rpr.find(qn("w:rFonts")).get(qn("w:eastAsia"))
p22_sz = int(p22_rpr.find(qn("w:sz")).get(qn("w:val"))) / 2
p22_b = p22_rpr.find(qn("w:b")) is not None
p22_ind = p22._element.find(qn("w:pPr")).find(qn("w:ind"))
p22_first = p22_ind.get(qn("w:firstLine")) if p22_ind is not None else "0"
print(f"heading '1 范围': align={p22_jc} font={p22_ea} sz={p22_sz}pt bold={p22_b} indent={p22_first}")
all_ok = p22_jc == "both" and p22_ea == "宋体" and p22_sz == 10.5 and not p22_b and p22_first in ("0", None)
print(f"  => {'[OK] ALL CORRECT!' if all_ok else '[FAIL]'}")
