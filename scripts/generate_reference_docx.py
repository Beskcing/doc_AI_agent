"""生成 Pandoc reference.docx 模板

为 Pandoc Markdown→DOCX 转换提供标准样式定义，改善转换质量。
生成的 reference.docx 包含国标排版所需的段落样式、字体、字号定义。

使用方法:
    python scripts/generate_reference_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def generate_reference_docx(output_path: str = "configs/reference.docx") -> None:
    """生成 Pandoc reference.docx"""

    doc = Document()

    # 设置默认字体（正文样式）
    style = doc.styles["Normal"]
    font = style.font
    font.name = "仿宋_GB2312"
    font.size = Pt(16)  # 三号字
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 标题 1（一级标题）
    h1 = doc.styles["Heading 1"]
    h1.font.name = "黑体"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # 标题 2（二级标题）
    h2 = doc.styles["Heading 2"]
    h2.font.name = "黑体"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)

    # 标题 3（三级标题）
    h3 = doc.styles["Heading 3"]
    h3.font.name = "黑体"
    h3.font.size = Pt(16)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(3)
    h3.paragraph_format.space_after = Pt(3)

    # 标题 4
    h4 = doc.styles["Heading 4"]
    h4.font.name = "黑体"
    h4.font.size = Pt(16)
    h4.font.bold = True
    h4.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 标题 5
    h5 = doc.styles["Heading 5"]
    h5.font.name = "黑体"
    h5.font.size = Pt(16)
    h5.font.bold = True
    h5.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 标题 6
    h6 = doc.styles["Heading 6"]
    h6.font.name = "黑体"
    h6.font.size = Pt(16)
    h6.font.bold = True
    h6.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # Block Text（引用块）
    try:
        bt = doc.styles["Block Text"]
        bt.font.name = "楷体_GB2312"
        bt.font.size = Pt(16)
        bt.element.rPr.rFonts.set(qn("w:eastAsia"), "楷体_GB2312")
    except Exception:
        pass

    # 列表样式
    for list_style_name in ["List Bullet", "List Number", "List Paragraph"]:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = "仿宋_GB2312"
            ls.font.size = Pt(16)
            ls.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        except Exception:
            pass

    # 添加示例内容（Pandoc 需要这些段落来映射样式）
    doc.add_paragraph("Pandoc 样式参考文档", style="Normal")
    doc.add_heading("一级标题示例", level=1)
    doc.add_paragraph("正文内容示例。", style="Normal")
    doc.add_heading("二级标题示例", level=2)
    doc.add_paragraph("正文内容示例。", style="Normal")
    doc.add_heading("三级标题示例", level=3)
    doc.add_paragraph("正文内容示例。", style="Normal")

    # 保存
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"reference.docx 已生成: {output} ({output.stat().st_size} bytes)")


if __name__ == "__main__":
    generate_reference_docx()
