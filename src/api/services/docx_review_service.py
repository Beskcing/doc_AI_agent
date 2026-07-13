"""DOCX 排版后审查服务

排版完成后对最终 DOCX 进行全文质量审查：
- quick_review: 规则化快速检查（自动执行，不调用 LLM）
- deep_review: LLM 深度审查（用户手动触发，按章节分块审查）

审查维度：OCR 残留错误、语义错误、文字错误、结构完整性、格式一致性
"""

from __future__ import annotations

import json
import re
from collections.abc import Callable
from pathlib import Path
from typing import Any

from src.config import AppConfig
from src.db.crud import TaskCRUD, TaskReviewCRUD
from src.db.session import get_db_session
from src.tools.docx_text_extractor import DocxText, DocxTextExtractor
from src.utils.logger import get_logger
from src.utils.text_diff import DiffResult, compute_diff, get_changed_text

logger = get_logger(__name__)

# 快速审查：规则化检查的常见 OCR 错误模式
# 注：异常 Unicode 检测已扩充合法字符范围，涵盖国标技术文档常见符号：
#   拉丁扩展(U+00B0-00FF)、希腊字母(U+0370-03FF)、通用标点(U+2000-206F)
#   上下标(U+2070-209F)、类字母符号(U+2100-214F)、箭头(U+2190-21FF)、数学运算符(U+2200-22FF)
_QUICK_CHECK_PATTERNS: list[tuple[str, str, str, str]] = [
    # (正则模式, 错误类型, 位置提示, 修正建议)
    (r"[a-zA-Z]{2,}[，。；：]", "format", "英文单词后出现中文标点", "将中文标点替换为英文标点"),
    (r"\d+[。，]\d+", "text", "数字间出现中文标点", "将中文标点替换为小数点或英文逗号"),
    (
        r"[^\x00-\x7f\u00b0-\u00ff\u0370-\u03ff\u2000-\u206f\u2070-\u209f\u2100-\u214f\u2190-\u21ff\u2200-\u22ff\u3000-\u303f\u4e00-\u9fff\uff00-\uffef]",
        "ocr",
        "包含异常 Unicode 字符（疑似乱码）",
        "检查是否为 OCR 识别错误",
    ),
    (r"\b[01OIlZSB]{2,}\b", "ocr", "疑似数字/字母混淆", "检查是否为 0/O、1/l/I 混淆"),
    (r"([\u4e00-\u9fff])([A-Za-z0-9])", "format", "中文与英文/数字间缺少空格", "建议在中文与英文/数字之间添加空格"),
    (
        r"\$(?:\$|[^$]*\$)"  # $$...$$ 或 $...$ 分隔符残留
        r"|\\left|\\right|\\frac|\\sqrt|\\sum|\\int|\\prod|\\lim"
        r"|\\alpha|\\beta|\\gamma|\\delta|\\epsilon|\\lambda|\\mu|\\sigma|\\omega"
        r"|\\Delta|\\Gamma|\\Theta|\\Lambda|\\Xi|\\Pi|\\Sigma|\\Upsilon|\\Phi|\\Psi|\\Omega"
        r"|\\begin|\\end"
        r"|\\text|\\mathrm|\\mathbf|\\mathit|\\mathcal|\\mathbb|\\mathfrak"
        r"|\\pm|\\cdot|\\times|\\div|\\leq|\\geq|\\neq|\\approx|\\infty"
        r"|\\equiv|\\propto|\\sim|\\simeq|\\cong|\\asymp|\\doteq"
        r"|\\partial|\\nabla|\\forall|\\exists|\\in|\\subset|\\cup|\\cap"
        r"|\\notin|\\ni|\\subseteq|\\supseteq|\\supset|\\setminus|\\emptyset"
        r"|\\rightarrow|\\Rightarrow|\\langle|\\rangle"
        r"|\\to|\\mapsto|\\iff|\\implies|\\impliedby"
        r"|\\leftarrow|\\longleftarrow|\\longrightarrow|\\longleftrightarrow"
        r"|\\uparrow|\\downarrow|\\updownarrow"
        r"|\\hookleftarrow|\\hookrightarrow|\\leadsto"
        r"|\\sin|\\cos|\\tan|\\log|\\ln|\\exp|\\max|\\min|\\det|\\gcd"
        r"|\\limsup|\\liminf|\\arg|\\dim|\\hom|\\ker"
        r"|\\overline|\\underline|\\hat|\\tilde|\\bar|\\dot|\\vec"
        r"|\\widehat|\\widetilde|\\check|\\breve|\\acute|\\grave|\\ddot"
        r"|\\ldots|\\cdots|\\ddots|\\vdots"
        r"|\\big|\\Big|\\bigg|\\Bigg"
        r"|\\neg|\\land|\\lor|\\oplus|\\otimes|\\ominus"
        r"|\\circ|\\bullet|\\diamond|\\star|\\ast|\\dagger|\\ddagger"
        r"|\\parallel|\\perp|\\mid|\\gg|\\ll|\\prec|\\succ|\\preceq|\\succeq"
        r"|\\iint|\\iiint|\\oint"
        r"|\\displaystyle|\\textstyle|\\operatorname|\\stackrel|\\binom"
        r"|\\colon|\\smile|\\frown|\\models|\\bowtie"
        r"|\\not\\b|\\triangle|\\angle|\\measuredangle|\\square|\\Box"
        r"|\\aleph|\\hbar|\\imath|\\jmath|\\ell|\\wp|\\Re|\\Im"
        r"|\\prime|\\backslash|\\surd|\\nabla|\\clubsuit|\\diamondsuit|\\heartsuit|\\spadesuit",
        "latex",
        "残留 LaTeX 公式语法",
        "将 LaTeX 语法转换为对应 Unicode 符号或删除无法转换的公式",
    ),
]


class DocxReviewService:
    """DOCX 排版后审查服务"""

    # 深度审查每块最大字符数
    DEEP_REVIEW_CHUNK_SIZE: int = 8000

    # 标记版 HTML 缓存: {task_id: {html, issues, summary}}
    _html_cache: dict[str, dict] = {}

    def __init__(
        self,
        config: AppConfig,
        get_llm_client: Callable[[], Any],
        review_prompt_template: str = "",
    ):
        """初始化审查服务

        Args:
            config: 应用配置
            get_llm_client: LLM 客户端工厂函数（懒加载）
            review_prompt_template: 审查提示词模板（含 {review_text} 等占位符）
        """
        self._config = config
        self._get_llm_client = get_llm_client
        self._review_prompt_template = review_prompt_template
        self._extractor = DocxTextExtractor()

    def quick_review(self, task_id: str) -> dict | None:
        """快速审查：规则化检查（不调用 LLM）

        在排版管线完成后自动执行。仅做模式匹配检查，速度快、免费。

        Args:
            task_id: 任务 ID

        Returns:
            审查结果 dict，含 issues 和 summary；失败返回 None
        """
        logger.info("任务 %s: 开始快速审查", task_id)

        try:
            # 获取任务信息
            with get_db_session() as db:
                task = TaskCRUD.get(db, task_id)
                if not task:
                    logger.warning("任务 %s: 不存在", task_id)
                    return None
                docx_path = task.result_path
                user_id = task.user_id

            if not docx_path or not Path(docx_path).exists():
                logger.warning("任务 %s: DOCX 文件不存在: %s", task_id, docx_path)
                return None

            # 创建审查记录
            with get_db_session() as db:
                review = TaskReviewCRUD.create(
                    db,
                    task_id=task_id,
                    user_id=user_id,
                    review_type="quick",
                    status="running",
                )
                review_id = review.id

            # 提取 DOCX 文本
            docx_text: DocxText = self._extractor.extract(docx_path)
            full_text = docx_text.get_full_text()

            # 获取原始 cleaned_markdown 做增量对比（无原文时跳过，避免全量标记为 added）
            cleaned_md = task.cleaned_markdown_preview or ""
            if cleaned_md.strip():
                diff_result: DiffResult = compute_diff(cleaned_md, full_text)
                diff_summary = {
                    "total_added": diff_result.total_added,
                    "total_removed": diff_result.total_removed,
                    "total_modified": diff_result.total_modified,
                    "total_unchanged": diff_result.total_unchanged,
                    "has_changes": diff_result.has_changes,
                }
            else:
                diff_summary = None

            # 执行规则检查
            issues: list[dict] = self._run_rule_checks(full_text, docx_text)

            # 构建结果
            summary = self._build_summary(issues)
            summary["diff"] = diff_summary

            result = {"issues": issues, "summary": summary}

            # 保存结果
            with get_db_session() as db:
                TaskReviewCRUD.update_issues(db, review_id, result, status="completed")

            logger.info(
                "任务 %s: 快速审查完成, %d 个问题",
                task_id,
                summary["total_issues"],
            )
            # 审查结果变化，清除 HTML 缓存
            self.invalidate_html_cache(task_id)
            return result

        except Exception as e:
            logger.error("任务 %s: 快速审查失败: %s", task_id, e)
            # 尝试标记失败
            try:
                with get_db_session() as db:
                    if "review_id" in dir():
                        TaskReviewCRUD.mark_failed(db, review_id, str(e))
            except Exception:
                pass
            return None

    def deep_review(
        self,
        task_id: str,
        progress_callback: Callable[[int, int, int], None] | None = None,
    ) -> dict | None:
        """深度审查：LLM 分块审查

        用户手动触发。按一级标题分块，每块独立 LLM 审查后汇总。

        Args:
            task_id: 任务 ID
            progress_callback: 进度回调 (progress, current_chunk, total_chunks)

        Returns:
            审查结果 dict；失败返回 None
        """
        logger.info("任务 %s: 开始深度审查", task_id)

        try:
            # 获取任务信息
            with get_db_session() as db:
                task = TaskCRUD.get(db, task_id)
                if not task:
                    return None
                docx_path = task.result_path
                user_id = task.user_id
                cleaned_md = task.cleaned_markdown_preview or ""
                doc_type = (task.config or {}).get("document_type", "通用文档")
                standard = task.standard or "GB/T 9704"

            if not docx_path or not Path(docx_path).exists():
                logger.warning("任务 %s: DOCX 文件不存在", task_id)
                return None

            # 创建审查记录
            with get_db_session() as db:
                review = TaskReviewCRUD.create(
                    db,
                    task_id=task_id,
                    user_id=user_id,
                    review_type="deep",
                    status="running",
                )
                review_id = review.id

            # 提取 DOCX 文本
            docx_text: DocxText = self._extractor.extract(docx_path)

            # 增量对比：只审查有变化的部分
            diff_result: DiffResult = compute_diff(cleaned_md, docx_text.get_full_text())
            if diff_result.has_changes:
                review_text = get_changed_text(diff_result, side="docx")
                review_scope = "增量审查（仅审查排版过程中发生变化的部分）"
            else:
                review_text = docx_text.get_full_text()
                review_scope = "全文审查"

            if not review_text.strip():
                # 无文本可审查
                empty_result = {"issues": [], "summary": self._build_summary([])}
                with get_db_session() as db:
                    TaskReviewCRUD.update_issues(db, review_id, empty_result)
                return empty_result

            # 按章节分块
            chunks = self._split_into_chunks(docx_text, max_chars=self.DEEP_REVIEW_CHUNK_SIZE)
            total_chunks = len(chunks)

            logger.info("任务 %s: 深度审查分 %d 块", task_id, total_chunks)

            if progress_callback:
                progress_callback(0, 0, total_chunks)

            # LLM 逐块审查
            all_issues: list[dict] = []
            llm = self._get_llm_client()

            for i, chunk_text in enumerate(chunks):
                logger.info("任务 %s: 审查第 %d/%d 块 (%d 字符)", task_id, i + 1, total_chunks, len(chunk_text))

                try:
                    chunk_issues = self._review_chunk(
                        llm,
                        chunk_text,
                        doc_type,
                        standard,
                        review_scope,
                    )
                    all_issues.extend(chunk_issues)
                except Exception as e:
                    logger.warning("任务 %s: 第 %d 块审查失败: %s", task_id, i + 1, e)
                    # 单块失败不中断整体流程

                if progress_callback:
                    progress_callback(
                        int((i + 1) / total_chunks * 100),
                        i + 1,
                        total_chunks,
                    )

                # 更新进度到 DB
                try:
                    with get_db_session() as db:
                        TaskReviewCRUD.update_progress(
                            db,
                            review_id,
                            progress=int((i + 1) / total_chunks * 100),
                            current_chunk=i + 1,
                            total_chunks=total_chunks,
                        )
                except Exception:
                    pass

            # 去重和汇总
            all_issues = self._deduplicate_issues(all_issues)
            summary = self._build_summary(all_issues)

            result = {"issues": all_issues, "summary": summary}

            # 保存结果
            with get_db_session() as db:
                TaskReviewCRUD.update_issues(db, review_id, result, status="completed")

            logger.info(
                "任务 %s: 深度审查完成, %d 个问题 (OCR:%d 语义:%d 文字:%d 结构:%d 格式:%d)",
                task_id,
                summary["total_issues"],
                summary["ocr_errors"],
                summary["semantic_errors"],
                summary["text_errors"],
                summary["structure_issues"],
                summary["format_issues"],
            )
            # 审查结果变化，清除 HTML 缓存
            self.invalidate_html_cache(task_id)
            return result

        except Exception as e:
            logger.error("任务 %s: 深度审查失败: %s", task_id, e)
            try:
                with get_db_session() as db:
                    if "review_id" in dir():
                        TaskReviewCRUD.mark_failed(db, review_id, str(e))
            except Exception:
                pass
            return None

    # ==================== 规则检查 ====================

    def _run_rule_checks(self, text: str, docx_text: DocxText) -> list[dict]:
        """执行规则化快速检查"""
        issues: list[dict] = []

        for pattern, issue_type, location_hint, suggestion in _QUICK_CHECK_PATTERNS:
            for match in re.finditer(pattern, text):
                # 找到匹配所在的段落
                para_idx = self._find_paragraph_index(docx_text, match.start())
                issues.append(
                    {
                        "type": issue_type,
                        "severity": "medium" if issue_type == "latex" else "low",
                        "location": f"第{para_idx + 1}段",
                        "original": match.group(0),
                        "suggested": suggestion,
                        "reason": location_hint,
                    }
                )

        return issues

    def _find_paragraph_index(self, docx_text: DocxText, char_pos: int) -> int:
        """根据字符位置查找段落索引"""
        current_pos = 0
        for para in docx_text.paragraphs:
            para_len = len(para.text) + 1  # +1 for newline
            if current_pos + para_len > char_pos:
                return para.index
            current_pos += para_len
        return len(docx_text.paragraphs) - 1

    # ==================== LLM 审查 ====================

    def _split_into_chunks(self, docx_text: DocxText, max_chars: int = 8000) -> list[str]:
        """按一级标题将文档分块

        优先在标题处分割，每块不超过 max_chars。
        如果单段超过 max_chars，强制在 max_chars 处截断。
        """
        chunks: list[str] = []
        current: list[str] = []
        current_len = 0

        for para in docx_text.paragraphs:
            para_text = para.text
            para_len = len(para_text) + 1

            # 一级标题始终开启新块
            if para.level == 1 and current_len > 0:
                chunks.append("\n".join(current))
                current = []
                current_len = 0

            # 如果当前段加入后超长，先保存当前块
            if current_len + para_len > max_chars and current:
                chunks.append("\n".join(current))
                current = []
                current_len = 0

            current.append(para_text)
            current_len += para_len

        if current:
            chunks.append("\n".join(current))

        return chunks

    def _review_chunk(
        self,
        llm: Any,
        chunk_text: str,
        doc_type: str,
        standard: str,
        scope: str,
    ) -> list[dict]:
        """审查单个文本块（调用 LLM）"""
        if not llm or not self._review_prompt_template:
            logger.warning("LLM 或提示词不可用，跳过审查")
            return []

        prompt = self._review_prompt_template
        prompt = prompt.replace("{review_text}", chunk_text)
        prompt = prompt.replace("{document_type}", doc_type)
        prompt = prompt.replace("{detected_standard}", standard)
        prompt = prompt.replace("{review_scope}", scope)

        response = llm.invoke(prompt).content

        # 解析 LLM JSON 响应
        try:
            # 尝试直接解析
            data = json.loads(response)
        except json.JSONDecodeError:
            # 尝试提取 JSON（可能有 Markdown 包裹）
            json_match = re.search(r"\{[\s\S]*\}", response)
            if json_match:
                try:
                    data = json.loads(json_match.group(0))
                except json.JSONDecodeError:
                    logger.warning("LLM 审查响应 JSON 解析失败")
                    return []
            else:
                logger.warning("LLM 审查响应无 JSON")
                return []

        return data.get("issues", [])

    # ==================== 工具方法 ====================

    @staticmethod
    def _build_summary(issues: list[dict]) -> dict:
        """根据 issues 列表构建汇总统计"""
        ocr = sum(1 for i in issues if i.get("type") == "ocr")
        semantic = sum(1 for i in issues if i.get("type") == "semantic")
        text = sum(1 for i in issues if i.get("type") == "text")
        structure = sum(1 for i in issues if i.get("type") == "structure")
        fmt = sum(1 for i in issues if i.get("type") == "format")
        latex = sum(1 for i in issues if i.get("type") == "latex")
        total = len(issues)

        if total == 0:
            quality = "good"
        elif sum(1 for i in issues if i.get("severity") == "high") > 0:
            quality = "poor"
        elif total > 10:
            quality = "poor"
        elif total > 3:
            quality = "fair"
        else:
            quality = "good"

        return {
            "total_issues": total,
            "ocr_errors": ocr,
            "semantic_errors": semantic,
            "text_errors": text,
            "structure_issues": structure,
            "format_issues": fmt,
            "latex_residue": latex,
            "overall_quality": quality,
        }

    @staticmethod
    def _deduplicate_issues(issues: list[dict]) -> list[dict]:
        """去除重复的 issue（按 original 文本去重）"""
        seen: set[str] = set()
        unique: list[dict] = []
        for issue in issues:
            key = issue.get("original", "")
            if key and key not in seen:
                seen.add(key)
                unique.append(issue)
        return unique

    # ==================== 标记 DOCX 生成 ====================

    def generate_marked_docx(self, task_id: str) -> dict | None:
        """生成标记版 DOCX（黄色高亮 + 批注）

        Returns:
            dict: {marked_docx_path, total_issues, highlighted, commented, errors}
            None: 失败
        """
        logger.info("任务 %s: 生成标记版 DOCX", task_id)

        try:
            # 获取任务信息
            with get_db_session() as db:
                task = TaskCRUD.get(db, task_id)
                if not task or not task.result_path:
                    return None
                docx_path = task.result_path

            if not docx_path or not Path(docx_path).exists():
                logger.warning("任务 %s: DOCX 不存在", task_id)
                return None

            # 获取审查 issues
            issues = self._get_review_issues(task_id)
            if not issues:
                logger.info("任务 %s: 无审查 issues，跳过标记", task_id)
                return {
                    "marked_docx_path": docx_path,
                    "total_issues": 0,
                    "highlighted": 0,
                    "commented": 0,
                    "errors": 0,
                }

            # 生成标记版
            from src.tools.docx_review_marker import DocxReviewMarker

            output_dir = Path(docx_path).parent
            marked_path = output_dir / "review_marked.docx"

            marker = DocxReviewMarker()
            stats = marker.mark_issues(docx_path, issues, str(marked_path))
            stats["marked_docx_path"] = str(marked_path)

            logger.info(
                "任务 %s: 标记版 DOCX 已生成 → %s (%s)",
                task_id,
                marked_path.name,
                stats,
            )
            return stats

        except Exception as e:
            logger.error("任务 %s: 生成标记版 DOCX 失败: %s", task_id, e)
            return None

    # ==================== HTML 预览 ====================

    def generate_marked_html(self, task_id: str) -> dict | None:
        """生成标记版 HTML 预览

        Returns:
            dict: {html, issues, summary}
            None: 失败
        """
        # 检查缓存
        if task_id in self._html_cache:
            logger.info("任务 %s: 使用缓存的标记版 HTML", task_id)
            return self._html_cache[task_id]

        logger.info("任务 %s: 生成标记版 HTML 预览", task_id)

        try:
            # 获取任务信息
            with get_db_session() as db:
                task = TaskCRUD.get(db, task_id)
                if not task or not task.result_path:
                    return None
                docx_path = task.result_path

            if not docx_path or not Path(docx_path).exists():
                return None

            # 获取审查 issues
            issues = self._get_review_issues(task_id)
            summary = self._build_summary(issues)

            # 提取 DOCX 文本
            docx_text: DocxText = self._extractor.extract(docx_path)

            # 构建带标记的 HTML
            html = self._build_marked_html(docx_text, issues)

            result = {"html": html, "issues": issues, "summary": summary}

            # 写入缓存
            self._html_cache[task_id] = result

            return result

        except Exception as e:
            logger.error("任务 %s: 生成 HTML 预览失败: %s", task_id, e)
            return None

    def invalidate_html_cache(self, task_id: str) -> None:
        """清除指定任务的 HTML 缓存

        在审查数据变化时调用（修正、重新审查等）
        """
        if task_id in self._html_cache:
            del self._html_cache[task_id]
            logger.info("任务 %s: 已清除标记版 HTML 缓存", task_id)

    def _build_marked_html(self, docx_text: DocxText, issues: list[dict]) -> str:
        """构建带标记的 HTML 内容"""
        # 按段落索引建立 issues 索引
        issues_by_para: dict[int, list[dict]] = {}
        for i, issue in enumerate(issues):
            location = issue.get("location", "")
            para_idx = self._parse_location_index(location)
            if para_idx >= 0:
                if para_idx not in issues_by_para:
                    issues_by_para[para_idx] = []
                issues_by_para[para_idx].append({**issue, "_idx": i})

        type_labels = {
            "ocr": "OCR错误",
            "semantic": "语义错误",
            "text": "文字错误",
            "structure": "结构问题",
            "format": "格式问题",
            "latex": "LaTeX残留",
        }

        paragraphs_html: list[str] = []
        for para in docx_text.paragraphs:
            text = para.text
            if not text.strip():
                paragraphs_html.append("<div class='para empty'></div>")
                continue

            # 转义 HTML
            escaped = self._html_escape(text)

            if para.index in issues_by_para:
                para_issues = issues_by_para[para.index]
                # 对每个 issue 进行标记
                marked = escaped
                for issue in para_issues:
                    original = issue.get("original", "")
                    if original and original in marked:
                        i_type = issue.get("type", "unknown")
                        i_severity = issue.get("severity", "low")
                        i_reason = self._html_escape(issue.get("reason", ""))
                        i_suggested = self._html_escape(issue.get("suggested", ""))
                        i_idx = issue.get("_idx", 0)

                        severity_text = "严重" if i_severity == "high" else "中等" if i_severity == "medium" else "轻微"
                        tooltip = f"[{type_labels.get(i_type, i_type)}] [{severity_text}]\n原因: {i_reason}"
                        if i_suggested:
                            tooltip += f"\n建议: {i_suggested}"

                        mark_tag = (
                            f"<mark class='review-issue' "
                            f"data-issue-idx='{i_idx}' "
                            f"data-type='{i_type}' "
                            f"data-severity='{i_severity}' "
                            f"data-tooltip='{self._html_escape(tooltip)}'"
                            f">{original}</mark>"
                        )
                        marked = marked.replace(original, mark_tag, 1)

                cls = "para heading" if para.type == "heading" else "para body"
                level_attr = f" data-level='{para.level}'" if para.level > 0 else ""
                paragraphs_html.append(f"<div class='{cls}'{level_attr}>{marked}</div>")
            else:
                cls = "para heading" if para.type == "heading" else "para body"
                level_attr = f" data-level='{para.level}'" if para.level > 0 else ""
                paragraphs_html.append(f"<div class='{cls}'{level_attr}>{escaped}</div>")

        # 构建完整 HTML
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
  font-family: '宋体', SimSun, 'Times New Roman', serif;
  font-size: 14px;
  line-height: 1.8;
  color: #333;
  padding: 40px 60px;
  background: #fff;
}}
.para {{ margin-bottom: 4px; }}
.para.empty {{ height: 12px; }}
.heading {{ font-weight: bold; margin-top: 12px; }}
.heading[data-level="1"] {{ font-size: 18px; text-align: center; margin-top: 20px; }}
.heading[data-level="2"] {{ font-size: 15px; margin-top: 16px; }}
mark.review-issue {{
  background-color: #fff3b0;
  padding: 1px 3px;
  border-radius: 2px;
  cursor: pointer;
  position: relative;
  border-bottom: 2px dashed #e6a817;
  transition: background-color 0.2s;
}}
mark.review-issue:hover {{
  background-color: #ffe58f;
  border-bottom-color: #d48806;
}}
mark.review-issue[data-severity="high"] {{
  border-bottom-color: #f5222d;
  background-color: #fff1f0;
}}
mark.review-issue[data-severity="high"]:hover {{
  background-color: #ffccc7;
}}
mark.review-issue::after {{
  content: attr(data-tooltip);
  display: none;
  position: absolute;
  left: 0;
  bottom: 120%;
  background: #fff;
  border: 1px solid #d9d9d9;
  border-radius: 6px;
  padding: 8px 12px;
  font-size: 12px;
  line-height: 1.6;
  white-space: pre-line;
  z-index: 1000;
  box-shadow: 0 3px 12px rgba(0,0,0,0.15);
  font-weight: normal;
  color: #333;
  max-width: 400px;
}}
mark.review-issue:hover::after {{
  display: block;
}}
</style>
</head>
<body>
{"".join(paragraphs_html)}
<script>
// 点击 issue 高亮通知父窗口
document.querySelectorAll('mark.review-issue').forEach(function(el) {{
  el.addEventListener('click', function() {{
    var idx = this.getAttribute('data-issue-idx');
    window.parent.postMessage({{ type: 'issue-click', issueIndex: parseInt(idx) }}, '*');
  }});
}});
</script>
</body>
</html>"""
        return html

    @staticmethod
    def _html_escape(text: str) -> str:
        """HTML 转义"""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    # ==================== AI 修正 ====================

    def fix_single_issue(
        self,
        task_id: str,
        issue_index: int,
        fix_text: str | None = None,
        mode: str = "ai",
    ) -> dict | None:
        """修正单条审查 issue

        Args:
            task_id: 任务 ID
            issue_index: issues 列表中的索引
            fix_text: 手动修正文本（mode=manual 时使用）
            mode: "ai" 或 "manual"

        Returns:
            dict: {success, original, fixed_text, issue_index}
        """
        logger.info("任务 %s: 修正 issue #%d (mode=%s)", task_id, issue_index, mode)

        try:
            # 获取 issues
            issues = self._get_review_issues(task_id)
            # 按 _idx 查找目标 issue（_idx 指向原始未过滤列表中的位置）
            target: dict | None = None
            for issue in issues:
                if issue.get("_idx") == issue_index:
                    target = issue
                    break
            if target is None:
                return {"success": False, "error": f"Issue 索引 {issue_index} 未找到"}

            issue = target

            # 获取 DOCX 路径
            with get_db_session() as db:
                task = TaskCRUD.get(db, task_id)
                if not task or not task.result_path:
                    return {"success": False, "error": "任务或 DOCX 不存在"}
                docx_path = task.result_path

            original = issue.get("original", "")
            suggested = issue.get("suggested", "")

            if mode == "ai":
                # LLM 修正
                fixed_text = self._ai_fix_text(task_id, issue)
                if not fixed_text:
                    # AI 修正失败，使用建议兜底
                    fixed_text = suggested if suggested else original
            elif mode == "manual" and fix_text:
                fixed_text = fix_text
            else:
                return {"success": False, "error": "缺少修正文本"}

            if not fixed_text or fixed_text == original:
                return {
                    "success": False,
                    "error": "修正文本为空或与原文相同",
                    "original": original,
                }

            # 在 DOCX 中替换文本
            replaced = self._replace_text_in_docx(docx_path, original, fixed_text)

            if replaced:
                # 移除该 issue（标记为已修正）
                self._mark_issue_fixed(task_id, issue_index)

            return {
                "success": True,
                "original": original,
                "fixed_text": fixed_text,
                "issue_index": issue_index,
            }

        except Exception as e:
            logger.error("任务 %s: 修正 issue 失败: %s", task_id, e)
            return {"success": False, "error": str(e)}

    def batch_fix_issues(
        self,
        task_id: str,
        auto_fix_low: bool = True,
        issue_indices: list[int] | None = None,
    ) -> dict | None:
        """批量修正审查 issues

        Args:
            task_id: 任务 ID
            auto_fix_low: 是否自动修正 low 级别
            issue_indices: 指定修正的索引列表，为空则全部

        Returns:
            dict: {fixed, failed, pending, details}
        """
        logger.info(
            "任务 %s: 批量修正 (auto_low=%s, indices=%s)",
            task_id,
            auto_fix_low,
            issue_indices,
        )

        issues = self._get_review_issues(task_id)
        if not issues:
            return {"fixed": 0, "failed": 0, "pending": [], "details": []}

        # 确定需要修正的 issues（收集 _idx）
        # auto_fix_low=True: 仅自动修正 low 级别，medium/high 放入 pending 待确认
        # auto_fix_low=False: 修正全部指定 issues
        targets: list[int] = []
        pending: list[dict] = []
        for i, issue in enumerate(issues):
            issue_idx = issue.get("_idx", i)
            if issue_indices is not None and issue_idx not in issue_indices:
                continue
            severity = issue.get("severity", "low")
            if auto_fix_low:
                if severity == "low":
                    targets.append(issue_idx)
                else:
                    # high/medium 放入待确认列表
                    pending.append(
                        {
                            "issue_index": issue_idx,
                            "type": issue.get("type", ""),
                            "severity": severity,
                            "original": (issue.get("original", "") or "")[:80],
                            "location": issue.get("location", ""),
                            "reason": issue.get("reason", ""),
                        }
                    )
            else:
                targets.append(issue_idx)

        if not targets and not pending:
            return {"fixed": 0, "failed": 0, "pending": [], "details": []}

        details: list[dict] = []
        fixed = 0
        failed = 0

        for idx in targets:
            result = self.fix_single_issue(task_id, idx, mode="ai")
            if result is None:
                details.append({"issue_index": idx, "success": False, "error": "修正服务异常"})
                failed += 1
            elif result.get("success"):
                details.append(result)
                fixed += 1
            else:
                details.append(result)
                failed += 1

        logger.info("任务 %s: 批量修正完成, fixed=%d, failed=%d, pending=%d", task_id, fixed, failed, len(pending))
        return {"fixed": fixed, "failed": failed, "pending": pending, "details": details}

    def _ai_fix_text(self, task_id: str, issue: dict) -> str | None:
        """使用 LLM 生成修正文本"""
        llm = self._get_llm_client()
        if not llm:
            return None

        original = issue.get("original", "")
        suggested = issue.get("suggested", "")
        reason = issue.get("reason", "")
        itype = issue.get("type", "unknown")

        # 格式类问题的简单修正：建议通常已经是正确的替换文本
        if itype == "format" and suggested and suggested != original:
            return suggested

        # LaTeX 残留需要 LLM 智能转换（不能直接返回通用提示语）
        if itype == "latex":
            prompt = f"""你是一个文档审查修正助手。以下文本中包含 LaTeX 公式语法残留，需要清理为纯文本。

**原文（含 LaTeX 残留）:**
{original}

**问题原因:** {reason}

处理规则：
1. 将单个 LaTeX 命令替换为对应的 Unicode 符号（如 \\times → ×, \\alpha → α, \\rightarrow → →, \\leq → ≤）
2. 如果是 $$...$$ 或 $...$ 包裹的简短数学表达式，提取其中可转文字的部分
3. 如果是复杂的多行公式或无法转换为纯文本的公式，直接删除整段公式
4. 保留所有非 LaTeX 的正常文本内容

请直接输出清理后的文本（只输出结果，不要任何解释或标记）:"""
            try:
                response = llm.invoke(prompt).content
                fixed = response.strip().strip('"').strip("'")
                if fixed.startswith("```"):
                    fixed = fixed.split("\n", 1)[-1].rsplit("\n```", 1)[0]
                return fixed if fixed and fixed != original else original
            except Exception as e:
                logger.warning("AI LaTeX 修正失败: %s", e)
                return original

        # 通用 LLM 修正
        prompt = f"""你是一个文档审查修正助手。请根据审查意见修正以下文本。

**原文:**
{original}

**问题类型:** {itype}
**问题原因:** {reason}
**修正建议:** {suggested}

请直接输出修正后的文本（只输出修正后的文本，不要任何解释）。如果原文本身没问题，直接输出原文。"""

        try:
            response = llm.invoke(prompt).content
            fixed = response.strip().strip('"').strip("'")
            # 清理可能的代码块
            if fixed.startswith("```"):
                fixed = fixed.split("\n", 1)[-1].rsplit("\n```", 1)[0]
            return fixed if fixed else suggested
        except Exception as e:
            logger.warning("AI 修正失败: %s，使用建议兜底", e)
            return suggested if suggested else original

    @staticmethod
    def _replace_text_in_docx(docx_path: str, original: str, replacement: str) -> bool:
        """在 DOCX 中替换文本（简单替换，作用于所有段落和表格）"""
        from docx import Document
        from docx.oxml.ns import qn

        if not original or original == replacement:
            return False

        try:
            doc = Document(docx_path)
            replaced = False

            # 遍历段落
            for para in doc.paragraphs:
                if original in para.text:
                    for run in para.runs:
                        if original in run.text:
                            run.text = run.text.replace(original, replacement, 1)
                            replaced = True
                            break
                    # 如果跨 run，用段落级别替换
                    if not replaced and original in para.text:
                        inline = para._element
                        text = "".join((t.text or "") for t in inline.iter(qn("w:t")))
                        new_text = text.replace(original, replacement, 1)
                        if new_text != text:
                            # 清除所有 run 并重新设置
                            for r in para.runs:
                                r.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            replaced = True

            if replaced:
                doc.save(docx_path)
            return replaced

        except Exception as e:
            logger.warning("DOCX 文本替换失败: %s", e)
            return False

    def _mark_issue_fixed(self, task_id: str, issue_index: int) -> None:
        """标记 issue 为已修正（从审查结果中移除或标记）"""
        try:
            with get_db_session() as db:
                review = TaskReviewCRUD.get_by_task(db, task_id, review_type="quick")
                if not review or not review.issues:
                    return

                issues_data = review.issues.copy()
                issue_list = issues_data.get("issues", [])
                if issue_index < len(issue_list):
                    # 标记为已修正
                    issue_list[issue_index]["_fixed"] = True
                    issues_data["issues"] = issue_list
                    review.issues = issues_data
                    db.commit()
        except Exception as e:
            logger.warning("标记 issue 已修正失败: %s", e)

    # ==================== 辅助方法 ====================

    def _get_review_issues(self, task_id: str) -> list[dict]:
        """获取任务的最新审查 issues（优先 quick review）

        返回的每个 issue 含 _idx 字段，指向原始 issues 列表中的位置，
        用于前端修正操作时定位。
        """
        try:
            with get_db_session() as db:
                review = TaskReviewCRUD.get_by_task(db, task_id, review_type="quick")
                if not review or not review.issues:
                    review = TaskReviewCRUD.get_by_task(db, task_id)
                if review and review.issues:
                    issues = review.issues.get("issues", [])
                    # 过滤已修正的，但保留原始索引
                    result = []
                    for idx, issue in enumerate(issues):
                        if not issue.get("_fixed"):
                            result.append({**issue, "_idx": idx})
                    return result
            return []
        except Exception as e:
            logger.warning("获取审查 issues 失败: %s", e)
            return []

    @staticmethod
    def _parse_location_index(location: str) -> int:
        """解析 location 字段获取段落索引（0-based）"""
        import re

        match = re.search(r"(\d+)", location)
        if match:
            return int(match.group(1)) - 1
        return -1
