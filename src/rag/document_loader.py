"""规范文档加载器

加载国标排版规范文档，支持多种格式，并为每个 Document 附加元数据。
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from src.utils.file_utils import find_files, get_file_extension, read_text_file
from src.utils.logger import get_logger

logger = get_logger(__name__)


class StandardsDocumentLoader:
    """国标排版规范文档加载器

    支持格式: .md, .txt, .pdf（通过 MinerU）, .docx（通过 MarkItDown）
    """

    def __init__(self, raw_docs_dir: str | Path):
        """初始化加载器

        Args:
            raw_docs_dir: 原始规范文档目录
        """
        self.raw_docs_dir = Path(raw_docs_dir)
        if not self.raw_docs_dir.exists():
            logger.warning("规范文档目录不存在: %s", self.raw_docs_dir)

    def load_all(self) -> list[Document]:
        """加载目录下所有规范文档

        Returns:
            Document 列表，每个 Document 包含元数据
        """
        if not self.raw_docs_dir.exists():
            logger.warning("目录不存在，返回空列表: %s", self.raw_docs_dir)
            return []

        documents: list[Document] = []

        # 遍历支持的文件格式
        for ext in ["*.md", "*.txt", "*.docx", "*.pdf"]:
            files = find_files(self.raw_docs_dir, ext)
            for file_path in files:
                try:
                    docs = self._load_single(file_path)
                    documents.extend(docs)
                    logger.info("已加载: %s (%d 个片段)", file_path.name, len(docs))
                except Exception as e:
                    logger.error("加载文件失败 %s: %s", file_path, e)

        logger.info("共加载 %d 个文档片段", len(documents))
        return documents

    def _load_single(self, file_path: Path) -> list[Document]:
        """加载单个文档，按章节初步分段

        Args:
            file_path: 文件路径

        Returns:
            Document 列表
        """
        ext = get_file_extension(file_path)

        if ext in ("md", "txt"):
            return self._load_text(file_path)
        elif ext == "docx":
            return self._load_docx(file_path)
        elif ext == "pdf":
            return self._load_pdf(file_path)
        else:
            logger.warning("不支持的文件格式: %s", file_path)
            return []

    def _load_text(self, file_path: Path) -> list[Document]:
        """加载 Markdown/TXT 文件，按章节分段

        Args:
            file_path: 文件路径

        Returns:
            Document 列表
        """
        content = read_text_file(file_path)
        file_name = file_path.name
        standard_id = self._extract_standard_id(file_name)

        # 按一级/二级标题分段
        sections = self._split_by_headings(content)

        documents = []
        for idx, (heading, section_content) in enumerate(sections):
            doc = Document(
                page_content=section_content.strip(),
                metadata={
                    "source": file_name,
                    "source_path": str(file_path),
                    "section_title": heading or "(引言)",
                    "section_index": idx,
                    "standard_id": standard_id,
                    "format": "markdown",
                },
            )
            if doc.page_content:
                documents.append(doc)

        return documents

    def _load_docx(self, file_path: Path) -> list[Document]:
        """加载 DOCX 文件（通过 python-docx 提取文本）

        Args:
            file_path: 文件路径

        Returns:
            Document 列表
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx 未安装")
            return []

        doc = DocxDocument(str(file_path))
        content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        file_name = file_path.name
        standard_id = self._extract_standard_id(file_name)

        return [
            Document(
                page_content=content,
                metadata={
                    "source": file_name,
                    "source_path": str(file_path),
                    "section_title": "(完整文档)",
                    "standard_id": standard_id,
                    "format": "docx",
                },
            )
        ]

    def _load_pdf(self, file_path: Path) -> list[Document]:
        """加载 PDF 文件（通过 MinerU 或 PyPDF）

        Args:
            file_path: 文件路径

        Returns:
            Document 列表
        """
        # 尝试使用 MinerU
        try:
            from src.tools.mineru_parser import MinerUParser
            parser = MinerUParser()
            if parser.check_installation():
                result = parser.parse_pdf(file_path)
                return [
                    Document(
                        page_content=result.raw_markdown,
                        metadata={
                            "source": file_path.name,
                            "source_path": str(file_path),
                            "section_title": "(MinerU 解析)",
                            "standard_id": self._extract_standard_id(file_path.name),
                            "format": "pdf_mineru",
                        },
                    )
                ]
        except Exception as e:
            logger.warning("MinerU 解析 PDF 失败: %s", e)

        # 降级: 返回空
        logger.warning("PDF 文件无法加载（MinerU 未安装）: %s", file_path)
        return []

    def _split_by_headings(self, content: str) -> list[tuple[str, str]]:
        """按标题分段

        Args:
            content: Markdown 文本

        Returns:
            [(标题, 内容), ...] 列表
        """
        import re

        heading_pattern = re.compile(r"^(#{1,3})\s+(.+)$", re.MULTILINE)
        matches = list(heading_pattern.finditer(content))

        if not matches:
            return [("", content)]

        sections = []
        # 第一个标题之前的内容
        if matches[0].start() > 0:
            sections.append(("", content[: matches[0].start()]))

        for i, match in enumerate(matches):
            heading = match.group(2).strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            section_content = content[start:end]
            sections.append((heading, section_content))

        return sections

    def _extract_standard_id(self, file_name: str) -> str:
        """从文件名中提取标准编号

        例如: gbt_9704_2012.md -> GB/T 9704-2012

        Args:
            file_name: 文件名

        Returns:
            标准编号字符串
        """
        import re

        # 匹配 gbt_NNNN_YYYY 格式
        match = re.search(r"gbt[_-](\d+)[_-](\d{4})", file_name, re.IGNORECASE)
        if match:
            return f"GB/T {match.group(1)}-{match.group(2)}"

        # 匹配 gb_NNNN 格式
        match = re.search(r"gb[_-](\d+)", file_name, re.IGNORECASE)
        if match:
            return f"GB {match.group(1)}"

        return file_name
