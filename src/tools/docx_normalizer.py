"""DOCX 内容规整工具

对 MinerU 输出的原始 DOCX 文件进行结构修复，包括：
- 日期行合并（"发布"/"实施"分两行 → 合并为一行）
- TOC 删除（移除"目次"→"前言"之间的目录段落）
- 拆分标题合并（"3.1.1" + "白兰地" → "3.1.1  白兰地"）
- 标题双空格修正（"1 范围" → "1  范围"）

纯内容重组，不涉及字体/样式（由 DocxStyler 负责）。
从 .qoder/skills/gbt/scripts/format_gbt.py 抽取并重构。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxNormalizer:
    """DOCX 内容规整器

    对 DOCX 段落进行结构级重组，修复 MinerU 解析产生的常见问题。
    所有操作通过 python-docx 的 Document 对象和 XML 层级完成。
    """

    def __init__(self) -> None:
        self.changes: list[str] = []

    def normalize(self, docx_path: str, output_path: str) -> str:
        """规整 DOCX 内容

        Args:
            docx_path: 输入 DOCX 文件路径
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        self.changes = []
        input_path = Path(docx_path)

        if not input_path.exists():
            raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")

        logger.info("开始 DOCX 内容规整: %s", docx_path)
        doc = Document(str(input_path))

        # 1. 合并日期行
        self._merge_date_lines(doc)

        # 2. 删除 TOC
        self._remove_toc(doc)

        # 3. 合并拆分标题
        self._merge_split_headings(doc)

        # 4. 修正标题双空格
        self._fix_heading_spaces(doc)

        doc.save(str(output_path))
        logger.info(
            "DOCX 内容规整完成, %d 处更改: %s",
            len(self.changes),
            self.changes[:5] if self.changes else "无",
        )
        return str(output_path)

    # ==================== 日期行合并 ====================

    def _merge_date_lines(self, doc: Document) -> None:
        """合并相邻的"发布"/"实施"日期段落为一行

        将:
            "2016-08-31发布"
            "2017-03-01实施"
        合并为:
            "2016-08-31发布 2017-03-01实施"

        扫描所有段落，找到相邻的"发布"+"实施"对，合并后从文档中删除第二个段落。
        """
        paragraphs = doc.paragraphs
        if len(paragraphs) < 2:
            return

        for i in range(len(paragraphs) - 1):
            p1_text = paragraphs[i].text.strip()
            p2_text = paragraphs[i + 1].text.strip()

            if p1_text.endswith("发布") and p2_text.endswith("实施"):
                date_pub = p1_text.replace("发布", "").strip()
                date_impl = p2_text.replace("实施", "").strip()
                merged = f"{date_pub}发布 {date_impl}实施"

                # 更新第一个段落
                p1 = paragraphs[i]
                for run in p1.runs:
                    run.text = ""
                if p1.runs:
                    p1.runs[0].text = merged
                else:
                    p1.add_run(merged)

                # 从 XML body 中删除第二个段落
                p2_elem = paragraphs[i + 1]._element
                p2_elem.getparent().remove(p2_elem)

                self.changes.append(f"合并日期行: '{merged}'")
                return  # 只合并第一处

    # ==================== TOC 删除 ====================

    def _remove_toc(self, doc: Document) -> None:
        """检测并删除 TOC（目次/目录）段落

        定位"目次"或"目录"段落，删除从此处到"前言"之间的所有段落。
        支持多种 TOC 标题变体：目次、目  次、目录、目  录。
        """
        paragraphs = doc.paragraphs
        toc_start = -1

        for i, p in enumerate(paragraphs):
            t = p.text.strip()
            if toc_start < 0:
                if t in ("目  次", "目次", "目  录", "目录"):
                    toc_start = i
                    logger.debug("检测到 TOC 起始: P[%d] '%s'", i, t)
            elif t.startswith("前") and "言" in t:
                # 找到 TOC 结束位置，删除范围内的所有段落
                toc_end = i
                body = doc.element.body
                for j in range(toc_start, toc_end):
                    p_elem = paragraphs[j]._element
                    body.remove(p_elem)
                removed = toc_end - toc_start
                self.changes.append(f"删除 TOC: {removed} 个段落 (P[{toc_start}]-P[{toc_end - 1}])")
                logger.debug("TOC 已删除: P[%d]-P[%d], 共 %d 段", toc_start, toc_end - 1, removed)
                return

    # ==================== 拆分标题合并 ====================

    def _merge_split_headings(self, doc: Document) -> None:
        """合并拆分到两行的标题

        处理两种情况：
        1. 附录标题拆分: "附录 A" + "培养基和试剂" → "附录 A  培养基和试剂"
        2. 编号标题拆分: "3.1.1" + "白兰地 brandy" → "3.1.1  白兰地 brandy"

        遍历所有段落，检测纯编号行后面紧跟描述文本的情况，
        合并为一行并用双空格分隔，然后从文档中删除第二行。
        """
        paragraphs = doc.paragraphs
        if len(paragraphs) < 2:
            return

        merged_count = 0
        i = 0
        while i < len(paragraphs) - 1:
            p1_text = paragraphs[i].text.strip()
            p2_text = paragraphs[i + 1].text.strip()

            if not p2_text:
                i += 1
                continue

            should_merge = False

            # 情形 1: 附录标题拆分 "附录 A" + "培养基和试剂"
            if re.match(r"^附录\s*[A-Z]$", p1_text):
                if (
                    not p2_text.startswith("附录")
                    and not p2_text.startswith("前言")
                    and not p2_text.startswith("参考文献")
                    and not re.match(r"^\d+\s+", p2_text)
                ):
                    should_merge = True

            # 情形 2: 编号标题拆分 "3.1.1" + "白兰地 brandy"
            elif re.match(r"^(\d+(\.\d+)*|[A-Z]\.\d+(\.\d+)*)$", p1_text):
                if (
                    not re.match(r"^(\d+(\.\d+)*|[A-Z]\.\d+)", p2_text)
                    and not p2_text.startswith("附录")
                    and not p2_text.startswith("前言")
                    and not p2_text.startswith("参考文献")
                ):
                    should_merge = True

            if should_merge:
                merged = f"{p1_text}  {p2_text}"

                # 更新第一个段落
                p1 = paragraphs[i]
                for run in p1.runs:
                    run.text = ""
                if p1.runs:
                    p1.runs[0].text = merged
                else:
                    p1.add_run(merged)

                # 从 XML body 中删除第二个段落
                p2_elem = paragraphs[i + 1]._element
                p2_elem.getparent().remove(p2_elem)

                merged_count += 1
                i += 1  # 跳过已合并的第二个段落
                continue

            i += 1

        if merged_count > 0:
            self.changes.append(f"合并拆分标题: {merged_count} 处")

    # ==================== 标题双空格修正 ====================

    def _fix_heading_spaces(self, doc: Document) -> None:
        """修正编号标题后空格为双空格

        国标要求编号与标题文字之间保留两个空格。
        例如: "1 范围" → "1  范围", "3.1 术语" → "3.1  术语"

        仅修正只有一个空格的情况（已有双空格或更多空格的不处理）。
        """
        fixed_count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            match = re.match(r"^(\d+(?:\.\d+)*)(\s+)(.*)", text)
            if not match:
                continue

            current_spaces = match.group(2)
            if current_spaces == "  ":
                continue  # 已是双空格

            num = match.group(1)
            rest = match.group(3)
            new_text = num + "  " + rest

            runs = paragraph.runs
            if runs:
                runs[0].text = new_text
                for r in runs[1:]:
                    r.text = ""
                fixed_count += 1

        if fixed_count > 0:
            self.changes.append(f"修正标题双空格: {fixed_count} 处")
