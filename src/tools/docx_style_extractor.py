"""DOCX 样式提取工具

从用户上传的 Word 模板文件中提取排版格式，输出 StyleConfig 兼容的 dict。
使用 python-docx 读取文档属性，纯工具函数，无 LLM 参与。

增强功能：
- 读取 Word 样式定义表（doc.styles）作为基线
- 正确提取东亚字体（eastAsia）和西文字体（ascii）
- 多段落采样，取最常见的样式
- 完整表格边框/对齐/背景色/单元格边距/表头重复
- 列表样式、脚注样式、图表标题样式、页眉页脚样式
- 段落格式：左右缩进、keep_together、widow_control、行距类型
"""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

from src.tools.content_pattern_matcher import (
    APPENDIX_CLAUSE_PATTERNS,
    APPENDIX_TITLE_PATTERN,
    HEADING_PATTERNS,
    SPECIAL_HEADING_PATTERNS,
    TABLE_CAPTION_PATTERN,
    classify_heading_level_by_content,
)
from src.utils.logger import get_logger

logger = get_logger(__name__)

# EMU 转换常量
EMU_PER_CM = 360000  # 1cm = 360000 EMU
EMU_PER_PT = 12700  # 1pt = 12700 EMU

# 对齐方式反向映射
ALIGNMENT_REVERSE_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

# 标题样式名映射
HEADING_STYLE_MAP = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
    "heading 4": 4, "heading 5": 5, "heading 6": 6,
    "Title": 1,  # Title 样式视为一级标题
}

# 样式定义表中的样式名映射
STYLE_DEFINITION_MAP = {
    "Normal": "body",
    "Heading 1": "heading_1",
    "Heading 2": "heading_2",
    "Heading 3": "heading_3",
    "Heading 4": "heading_4",
    "Heading 5": "heading_5",
    "Heading 6": "heading_6",
    "List Paragraph": "list",
    "Footnote Text": "footnote",
    "Caption": "caption",
    "Header": "header_footer",
    "Footer": "header_footer",
}

# 常见国标纸张尺寸 (EMU)
PAPER_SIZES = {
    (Cm(21.0), Cm(29.7)): "A4",
    (Cm(29.7), Cm(42.0)): "A3",
    (Cm(17.6), Cm(25.0)): "B5",
    (Cm(21.59), Cm(27.94)): "Letter",
}

# 表格边框名称
TABLE_BORDER_NAMES = ["top", "bottom", "left", "right", "insideH", "insideV"]

# 表格对齐方式映射
TABLE_ALIGN_MAP = {
    "left": "left",
    "center": "center",
    "right": "right",
}

# 单元格垂直对齐方式映射
CELL_VALIGN_MAP = {
    "top": "top",
    "center": "center",
    "bottom": "bottom",
}


def _emu_to_cm(emu: int | None) -> float:
    """将 EMU 转换为厘米"""
    if emu is None:
        return 0.0
    return round(emu / EMU_PER_CM, 2)


def _emu_to_pt(emu: int | None) -> float:
    """将 EMU 转换为磅"""
    if emu is None:
        return 0.0
    return round(emu / EMU_PER_PT, 1)


def _rgb_to_hex(rgb) -> str:
    """将 RGBColor 转换为十六进制字符串"""
    if rgb is None:
        return "#000000"
    return f"#{str(rgb)}"


def _twips_to_pt(twips: int | None) -> float:
    """将 twips 转换为磅（1 pt = 20 twips）"""
    if twips is None:
        return 0.0
    return round(int(twips) / 20, 1)


class DocxStyleExtractor:
    """从 DOCX 文件提取排版样式配置"""

    def extract(self, docx_path: str | Path) -> dict:
        """从 DOCX 提取完整的样式配置

        Args:
            docx_path: DOCX 文件路径

        Returns:
            StyleConfig 兼容的字典
        """
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")

        logger.info("开始提取 DOCX 样式: %s", docx_path)
        doc = Document(str(docx_path))

        # 第一步：读取样式定义表作为基线
        style_defs = self._extract_style_definitions(doc)

        result = {
            "page_layout": self._extract_page_layout(doc),
            "cover_style": self._extract_cover_style(doc),
            "preface_style": self._extract_preface_style(doc),
            "appendix_title_style": self._extract_appendix_title_style(doc),
            "appendix_clause_style": self._extract_appendix_clause_style(doc),
            "heading_styles": self._extract_heading_styles(doc, style_defs),
            "body_style": self._extract_body_style(doc, style_defs),
            "table_style": self._extract_table_style(doc),
            "table_caption_style": self._extract_table_caption_style(doc),
            "list_style": self._extract_list_style(doc, style_defs),
            "footnote_style": self._extract_footnote_style(doc, style_defs),
            "caption_style": self._extract_table_caption_style(doc),
            "header_footer_style": self._extract_header_footer_style(doc, style_defs),
            "rag_sources": [],
        }

        logger.info(
            "DOCX 样式提取完成: 封面=%s, 前言=%s, %d 标题样式, 表格=%s, 列表=%s, 脚注=%s, 页眉页脚=%s",
            result["cover_style"] is not None,
            result["preface_style"] is not None,
            len(result["heading_styles"]),
            result["table_style"] is not None,
            result["list_style"] is not None,
            result["footnote_style"] is not None,
            result["header_footer_style"] is not None,
        )
        return result

    # ==================== 样式定义表 ====================

    def _extract_style_definitions(self, doc: Document) -> dict:
        """从 doc.styles 读取样式定义表

        Returns:
            {
                "body": {...},       # Normal 样式
                "heading_1": {...},  # Heading 1 样式
                ...
                "list": {...},       # List Paragraph 样式
                "footnote": {...},   # Footnote Text 样式
                "caption": {...},    # Caption 样式
                "header_footer": {...},  # Header/Footer 样式
            }
        """
        defs: dict[str, dict] = {}

        for style in doc.styles:
            try:
                style_name = style.name
                if style_name not in STYLE_DEFINITION_MAP:
                    continue

                key = STYLE_DEFINITION_MAP[style_name]
                if key in defs:
                    continue  # header_footer 可能重复（Header + Footer）

                font_config = self._extract_font_from_style(style)
                pf = style.paragraph_format if hasattr(style, "paragraph_format") else None

                style_def = {
                    "font": font_config,
                    "line_spacing": 1.5,
                    "line_spacing_pt": None,
                    "line_spacing_rule": "multiple",
                    "space_before_pt": 0,
                    "space_after_pt": 0,
                    "alignment": "left",
                    "first_line_indent_chars": 0,
                    "left_indent_cm": 0,
                    "right_indent_cm": 0,
                    "keep_together": False,
                    "keep_with_next": False,
                    "widow_control": True,
                }

                if pf:
                    if pf.line_spacing is not None:
                        ls = pf.line_spacing
                        if isinstance(ls, float):
                            style_def["line_spacing"] = float(ls)
                            style_def["line_spacing_rule"] = "multiple"
                        else:
                            # Pt 对象或 int (EMU)
                            try:
                                style_def["line_spacing_pt"] = _emu_to_pt(ls)
                                style_def["line_spacing_rule"] = "exact"
                            except Exception:
                                style_def["line_spacing"] = float(ls)
                    if pf.space_before is not None:
                        style_def["space_before_pt"] = _emu_to_pt(pf.space_before)
                    if pf.space_after is not None:
                        style_def["space_after_pt"] = _emu_to_pt(pf.space_after)
                    if pf.alignment is not None:
                        style_def["alignment"] = ALIGNMENT_REVERSE_MAP.get(
                            pf.alignment, "left"
                        )
                    if pf.first_line_indent is not None:
                        font_size = font_config.get("size_pt", 16)
                        indent_pt = _emu_to_pt(pf.first_line_indent)
                        if font_size > 0:
                            style_def["first_line_indent_chars"] = round(
                                indent_pt / font_size, 1
                            )
                    if pf.left_indent is not None:
                        style_def["left_indent_cm"] = _emu_to_cm(pf.left_indent)
                    if pf.right_indent is not None:
                        style_def["right_indent_cm"] = _emu_to_cm(pf.right_indent)
                    if pf.keep_together is not None:
                        style_def["keep_together"] = pf.keep_together
                    if pf.keep_with_next is not None:
                        style_def["keep_with_next"] = pf.keep_with_next
                    if pf.widow_control is not None:
                        style_def["widow_control"] = pf.widow_control

                defs[key] = style_def
            except Exception as e:
                logger.debug("读取样式定义 '%s' 失败: %s", style_name, e)

        logger.info("从样式定义表提取了 %d 个样式定义", len(defs))
        return defs

    def _extract_font_from_style(self, style) -> dict:
        """从样式定义对象提取字体配置"""
        font_config = {
            "family": "仿宋_GB2312",
            "east_asia_family": None,
            "size_pt": 16,
            "bold": False,
            "italic": False,
            "underline": False,
            "strikethrough": False,
            "color_hex": "#000000",
        }

        try:
            font = style.font
            if font:
                if font.name:
                    font_config["family"] = font.name
                if font.size:
                    font_config["size_pt"] = _emu_to_pt(font.size)
                if font.bold is not None:
                    font_config["bold"] = font.bold
                if font.italic is not None:
                    font_config["italic"] = font.italic
                if font.underline is not None:
                    font_config["underline"] = bool(font.underline)
                if font.color and font.color.rgb:
                    font_config["color_hex"] = _rgb_to_hex(font.color.rgb)
        except Exception:
            pass

        # 从 XML 元素提取东亚字体
        try:
            element = style.element
            rpr = element.find(qn("w:rPr"))
            if rpr is not None:
                r_fonts = rpr.find(qn("w:rFonts"))
                if r_fonts is not None:
                    east_asia = r_fonts.get(qn("w:eastAsia"))
                    if east_asia:
                        font_config["east_asia_family"] = east_asia
                    ascii_font = r_fonts.get(qn("w:ascii"))
                    if ascii_font:
                        font_config["family"] = ascii_font
        except Exception:
            pass

        return font_config

    # ==================== 页面布局 ====================

    def _extract_page_layout(self, doc: Document) -> dict:
        """提取页面布局"""
        layout = {
            "paper_size": "A4",
            "margin_top_cm": 3.7,
            "margin_bottom_cm": 3.5,
            "margin_left_cm": 2.8,
            "margin_right_cm": 2.6,
            "header_distance_cm": 1.5,
            "footer_distance_cm": 1.75,
            "orientation": "portrait",
            "gutter_cm": 0,
            "page_number_format": None,
        }

        if not doc.sections:
            return layout

        section = doc.sections[0]

        # 纸张大小
        try:
            width = section.page_width
            height = section.page_height
            if width and height:
                key = (width, height)
                if key in PAPER_SIZES:
                    layout["paper_size"] = PAPER_SIZES[key]
                else:
                    key_rev = (height, width)
                    if key_rev in PAPER_SIZES:
                        layout["paper_size"] = PAPER_SIZES[key_rev]

                if height < width:
                    layout["orientation"] = "landscape"
        except Exception:
            pass

        # 页边距
        try:
            if section.top_margin:
                layout["margin_top_cm"] = _emu_to_cm(section.top_margin)
            if section.bottom_margin:
                layout["margin_bottom_cm"] = _emu_to_cm(section.bottom_margin)
            if section.left_margin:
                layout["margin_left_cm"] = _emu_to_cm(section.left_margin)
            if section.right_margin:
                layout["margin_right_cm"] = _emu_to_cm(section.right_margin)
            if section.header_distance:
                layout["header_distance_cm"] = _emu_to_cm(section.header_distance)
            if section.footer_distance:
                layout["footer_distance_cm"] = _emu_to_cm(section.footer_distance)
            if section.gutter:
                layout["gutter_cm"] = _emu_to_cm(section.gutter)
        except Exception:
            pass

        # 页码格式
        try:
            sect_pr = section._sectPr
            # 查找页码字段
            footer_refs = sect_pr.findall(qn("w:footerReference"))
            for fref in footer_refs:
                # 页码格式可能在 footer 的 fldChar 中
                pass
            # 从 pgNumType 提取页码格式
            pg_num_type = sect_pr.find(qn("w:pgNumType"))
            if pg_num_type is not None:
                fmt = pg_num_type.get(qn("w:fmt"))
                if fmt:
                    layout["page_number_format"] = fmt
        except Exception:
            pass

        return layout

    # ==================== 内容模式识别 ====================

    # 正则模式已提取到共享模块 content_pattern_matcher，此处引用以保持类内方法兼容
    _HEADING_PATTERNS = HEADING_PATTERNS
    _SPECIAL_HEADING_PATTERNS = SPECIAL_HEADING_PATTERNS
    _APPENDIX_TITLE_PATTERN = APPENDIX_TITLE_PATTERN
    _APPENDIX_CLAUSE_PATTERNS = APPENDIX_CLAUSE_PATTERNS
    _TABLE_CAPTION_PATTERN = TABLE_CAPTION_PATTERN

    def _classify_heading_level_by_content(self, text: str) -> int | None:
        """基于内容模式识别标题级别（委托共享模块）

        用于文档不使用标准 Heading 样式的后备方案。
        识别国标文档常见的条款编号格式。
        注意：附录标题和附录内条款不在此方法中识别，
        它们有独立的样式（appendix_title_style / appendix_clause_style）。

        Args:
            text: 段落文本（已 strip）

        Returns:
            标题级别 (1-5)，非标题返回 None
        """
        return classify_heading_level_by_content(text)

    def _is_cover_or_preface(self, paragraph) -> str | None:
        """判断段落是否属于封面或前言区域

        Returns:
            'cover' / 'preface' / None
        """
        text = paragraph.text.strip()
        if not text:
            return None

        # 封面特征：大字号(>14pt)、居中、加粗
        try:
            font_config = self._extract_font_from_paragraph(paragraph)
            size_pt = font_config.get("size_pt", 0)
            bold = font_config.get("bold", False)
            alignment = self._extract_alignment(paragraph)
            is_center = alignment == "center"
            is_large = size_pt and size_pt >= 14

            if is_large and is_center:
                return "cover"
        except Exception:
            pass

        return None

    # ==================== 封面/前言样式 ====================

    def _extract_cover_style(self, doc: Document) -> dict | None:
        """提取封面样式

        策略：扫描前 20 个段落，找出大字号居中段落作为封面样式。
        """
        cover_paragraphs = []
        for paragraph in doc.paragraphs[:20]:
            if not paragraph.text.strip():
                continue
            if self._is_cover_or_preface(paragraph) == "cover":
                cover_paragraphs.append(paragraph)

        if not cover_paragraphs:
            return None

        sample = cover_paragraphs[:min(5, len(cover_paragraphs))]
        font_configs = [self._extract_font_from_paragraph(p) for p in sample]
        alignments = [self._extract_alignment(p) for p in sample]
        line_spacings = [self._extract_line_spacing_full(p) for p in sample]

        result = {
            "description": "封面区域样式",
            "font": self._merge_font_configs(font_configs),
            "alignment": Counter(alignments).most_common(1)[0][0] if alignments else "center",
        }
        ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
        result["line_spacing"] = ls
        result["line_spacing_pt"] = ls_pt
        result["line_spacing_rule"] = ls_rule
        result.setdefault("space_before_pt", 0)
        result.setdefault("space_after_pt", 0)
        return result

    def _extract_preface_style(self, doc: Document) -> dict | None:
        """提取前言标题样式

        策略：查找包含「前言」或「引言」的段落，提取其样式。
        """
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text in ("前言", "引言") or text.startswith("前言") or text.startswith("引言"):
                font_config = self._extract_font_from_paragraph(paragraph)
                alignment = self._extract_alignment(paragraph)
                ls, ls_pt, ls_rule = self._extract_line_spacing_full(paragraph)
                sb, sa = self._extract_spacing(paragraph)
                result = {
                    "description": "前言/引言标题样式",
                    "font": font_config,
                    "alignment": alignment,
                    "line_spacing": ls,
                    "line_spacing_pt": ls_pt,
                    "line_spacing_rule": ls_rule,
                    "space_before_pt": sb,
                    "space_after_pt": sa,
                }
                return result
        return None

    # ==================== 附录样式 ====================

    def _extract_appendix_title_style(self, doc: Document) -> dict | None:
        """提取附录标题样式

        策略：查找以「附录A/B/C...」开头的段落，提取其样式。
        附录标题通常加粗，与普通一级标题(不加粗)不同。
        """
        appendix_paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if self._APPENDIX_TITLE_PATTERN.match(text):
                appendix_paragraphs.append(paragraph)

        if not appendix_paragraphs:
            return None

        sample = appendix_paragraphs[:min(5, len(appendix_paragraphs))]
        font_configs = [self._extract_font_from_paragraph(p) for p in sample]
        alignments = [self._extract_alignment(p) for p in sample]
        line_spacings = [self._extract_line_spacing_full(p) for p in sample]

        result = {
            "description": "附录标题样式",
            "font": self._merge_font_configs(font_configs),
            "alignment": Counter(alignments).most_common(1)[0][0] if alignments else "justify",
        }
        ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
        result["line_spacing"] = ls
        result["line_spacing_pt"] = ls_pt
        result["line_spacing_rule"] = ls_rule
        result.setdefault("space_before_pt", 0)
        result.setdefault("space_after_pt", 0)
        result.setdefault("first_line_indent_chars", 0)
        result.setdefault("left_indent_cm", 0)
        result.setdefault("right_indent_cm", 0)
        result.setdefault("keep_together", True)
        result.setdefault("keep_with_next", True)
        result.setdefault("widow_control", True)
        return result

    def _extract_appendix_clause_style(self, doc: Document) -> dict | None:
        """提取附录内条款样式

        策略：查找以「A.1/B.1...」开头的段落，提取其样式。
        附录内条款不加粗，格式与正文条款相同。
        """
        clause_paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            for _, pattern in self._APPENDIX_CLAUSE_PATTERNS:
                if pattern.match(text):
                    clause_paragraphs.append(paragraph)
                    break

        if not clause_paragraphs:
            return None

        sample = clause_paragraphs[:min(10, len(clause_paragraphs))]
        font_configs = [self._extract_font_from_paragraph(p) for p in sample]
        alignments = [self._extract_alignment(p) for p in sample]
        line_spacings = [self._extract_line_spacing_full(p) for p in sample]

        result = {
            "description": "附录内条款样式",
            "font": self._merge_font_configs(font_configs),
            "alignment": Counter(alignments).most_common(1)[0][0] if alignments else "justify",
        }
        ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
        result["line_spacing"] = ls
        result["line_spacing_pt"] = ls_pt
        result["line_spacing_rule"] = ls_rule
        result.setdefault("space_before_pt", 0)
        result.setdefault("space_after_pt", 0)
        result.setdefault("first_line_indent_chars", 0)
        result.setdefault("left_indent_cm", 0)
        result.setdefault("right_indent_cm", 0)
        result.setdefault("keep_together", False)
        result.setdefault("keep_with_next", True)
        result.setdefault("widow_control", True)
        return result

    def _extract_table_caption_style(self, doc: Document) -> dict | None:
        """提取表格标题样式

        策略：查找以「表B.1/表1...」开头的段落。
        """
        caption_paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if self._TABLE_CAPTION_PATTERN.match(text):
                caption_paragraphs.append(paragraph)

        if not caption_paragraphs:
            return None

        sample = caption_paragraphs[:min(5, len(caption_paragraphs))]
        font_configs = [self._extract_font_from_paragraph(p) for p in sample]
        alignments = [self._extract_alignment(p) for p in sample]

        result = {
            "description": "表格标题样式",
            "font": self._merge_font_configs(font_configs),
            "alignment": Counter(alignments).most_common(1)[0][0] if alignments else "center",
        }
        result.setdefault("line_spacing", 1.0)
        result.setdefault("line_spacing_pt", None)
        result.setdefault("line_spacing_rule", "multiple")
        result.setdefault("space_before_pt", 0)
        result.setdefault("space_after_pt", 0)
        result.setdefault("first_line_indent_chars", 0)
        result.setdefault("left_indent_cm", 0)
        result.setdefault("right_indent_cm", 0)
        result.setdefault("keep_together", False)
        result.setdefault("keep_with_next", True)
        result.setdefault("widow_control", True)
        return result

    # ==================== 标题样式 ====================

    def _extract_heading_styles(
        self, doc: Document, style_defs: dict
    ) -> list[dict]:
        """提取各级标题样式

        策略：
        1. 先从样式定义表获取基线
        2. 从文档段落中采样，取最常见的直接格式覆盖基线
        3. 如果没有标准 Heading 样式段落，使用内容模式识别后备方案
        """
        # 按级别收集标题段落
        heading_paragraphs: dict[int, list] = {}
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name not in HEADING_STYLE_MAP:
                continue
            level = HEADING_STYLE_MAP[style_name]
            if level not in heading_paragraphs:
                heading_paragraphs[level] = []
            heading_paragraphs[level].append(paragraph)

        # 后备方案：如果没找到标准 Heading 样式段落，使用内容模式识别
        if not heading_paragraphs:
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                level = self._classify_heading_level_by_content(text)
                if level is not None:
                    if level not in heading_paragraphs:
                        heading_paragraphs[level] = []
                    heading_paragraphs[level].append(paragraph)
            if heading_paragraphs:
                logger.info("内容模式识别到 %d 级标题", len(heading_paragraphs))

        heading_styles: dict[int, dict] = {}

        # 先从样式定义表获取基线
        for level in range(1, 7):
            key = f"heading_{level}"
            if key in style_defs:
                base = dict(style_defs[key])
                base["level"] = level
                base["numbering_format"] = None
                base["outline_level"] = level
                base["keep_with_next"] = True
                heading_styles[level] = base

        # 从实际段落覆盖
        for level, paragraphs in heading_paragraphs.items():
            # 采样多个段落，取最常见的样式
            font_configs = []
            alignments = []
            line_spacings = []
            space_befores = []
            space_afters = []
            first_indents = []
            numbering_formats = []
            outline_levels = []

            for paragraph in paragraphs:
                font_configs.append(self._extract_font_from_paragraph(paragraph))
                alignments.append(self._extract_alignment(paragraph))
                ls, ls_pt, ls_rule = self._extract_line_spacing_full(paragraph)
                line_spacings.append((ls, ls_pt, ls_rule))
                sb, sa = self._extract_spacing(paragraph)
                space_befores.append(sb)
                space_afters.append(sa)
                first_indents.append(self._extract_first_indent(paragraph))
                nf = self._extract_numbering_format(paragraph)
                if nf:
                    numbering_formats.append(nf)
                ol = self._extract_outline_level(paragraph)
                if ol is not None:
                    outline_levels.append(ol)

            # 取最常见的值
            result = heading_styles.get(level, {"level": level})
            result["level"] = level

            if font_configs:
                result["font"] = self._merge_font_configs(font_configs)
            if alignments:
                result["alignment"] = Counter(alignments).most_common(1)[0][0]
            if line_spacings:
                ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
                result["line_spacing"] = ls
                result["line_spacing_pt"] = ls_pt
                result["line_spacing_rule"] = ls_rule
            if space_befores:
                result["space_before_pt"] = Counter(space_befores).most_common(1)[0][0]
            if space_afters:
                result["space_after_pt"] = Counter(space_afters).most_common(1)[0][0]
            if first_indents:
                result["first_line_indent_chars"] = Counter(first_indents).most_common(1)[0][0]
            if numbering_formats:
                result["numbering_format"] = Counter(numbering_formats).most_common(1)[0][0]
            if outline_levels:
                result["outline_level"] = Counter(outline_levels).most_common(1)[0][0]

            result.setdefault("numbering_format", None)
            result.setdefault("outline_level", level)
            result["keep_with_next"] = True

            heading_styles[level] = result

        return sorted(heading_styles.values(), key=lambda x: x["level"])

    # ==================== 正文样式 ====================

    def _extract_body_style(self, doc: Document, style_defs: dict) -> dict:
        """提取正文样式

        策略：
        1. 先从 Normal 样式定义获取基线
        2. 从文档正文中采样多个非标题段落
        3. 排除封面/前言/标题段落，避免污染正文样式
        4. 取最常见的样式
        """
        # 基线
        base = dict(style_defs.get("body", {}))

        # 收集正文段落
        body_paragraphs = []
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name in HEADING_STYLE_MAP:
                continue
            if "List" in style_name or "list" in style_name:
                continue
            if "TOC" in style_name or "toc" in style_name:
                continue
            if "Footnote" in style_name or "footnote" in style_name:
                continue
            if "Caption" in style_name or "caption" in style_name:
                continue
            if "Header" in style_name or "Footer" in style_name:
                continue
            if not paragraph.text.strip():
                continue
            # 排除封面段落（大字号居中）
            if self._is_cover_or_preface(paragraph) == "cover":
                continue
            # 排除内容模式识别到的标题段落
            if self._classify_heading_level_by_content(paragraph.text.strip()) is not None:
                continue
            # 排除附录标题段落（加粗，不同于普通标题）
            if self._APPENDIX_TITLE_PATTERN.match(paragraph.text.strip()):
                continue
            # 排除附录内条款段落
            is_appendix_clause = False
            for _, pattern in self._APPENDIX_CLAUSE_PATTERNS:
                if pattern.match(paragraph.text.strip()):
                    is_appendix_clause = True
                    break
            if is_appendix_clause:
                continue
            # 排除表格标题段落
            if self._TABLE_CAPTION_PATTERN.match(paragraph.text.strip()):
                continue
            body_paragraphs.append(paragraph)

        if not body_paragraphs:
            if base:
                return base
            return self._default_body_style()

        # 采样最多 20 个段落
        sample = body_paragraphs[: min(20, len(body_paragraphs))]

        font_configs = []
        alignments = []
        line_spacings = []
        space_befores = []
        space_afters = []
        first_indents = []
        left_indents = []
        right_indents = []

        for paragraph in sample:
            font_configs.append(self._extract_font_from_paragraph(paragraph))
            alignments.append(self._extract_alignment(paragraph))
            ls, ls_pt, ls_rule = self._extract_line_spacing_full(paragraph)
            line_spacings.append((ls, ls_pt, ls_rule))
            sb, sa = self._extract_spacing(paragraph)
            space_befores.append(sb)
            space_afters.append(sa)
            first_indents.append(self._extract_first_indent(paragraph))
            left_indents.append(self._extract_left_indent(paragraph))
            right_indents.append(self._extract_right_indent(paragraph))

        result = dict(base) if base else {}
        result["font"] = self._merge_font_configs(font_configs) if font_configs else result.get("font", {})
        result["alignment"] = Counter(alignments).most_common(1)[0][0] if alignments else result.get("alignment", "left")
        if line_spacings:
            ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
            result["line_spacing"] = ls
            result["line_spacing_pt"] = ls_pt
            result["line_spacing_rule"] = ls_rule
        result["space_before_pt"] = Counter(space_befores).most_common(1)[0][0] if space_befores else result.get("space_before_pt", 0)
        result["space_after_pt"] = Counter(space_afters).most_common(1)[0][0] if space_afters else result.get("space_after_pt", 0)
        result["first_line_indent_chars"] = Counter(first_indents).most_common(1)[0][0] if first_indents else result.get("first_line_indent_chars", 0)
        result["left_indent_cm"] = Counter(left_indents).most_common(1)[0][0] if left_indents else result.get("left_indent_cm", 0)
        result["right_indent_cm"] = Counter(right_indents).most_common(1)[0][0] if right_indents else result.get("right_indent_cm", 0)
        result.setdefault("keep_together", False)
        result.setdefault("keep_with_next", False)
        result.setdefault("widow_control", True)

        return result

    def _default_body_style(self) -> dict:
        """默认正文样式"""
        return {
            "font": {
                "family": "仿宋_GB2312",
                "east_asia_family": "仿宋_GB2312",
                "size_pt": 16,
                "bold": False,
                "italic": False,
                "underline": False,
                "strikethrough": False,
                "color_hex": "#000000",
            },
            "line_spacing": 1.5,
            "line_spacing_pt": None,
            "line_spacing_rule": "multiple",
            "space_before_pt": 0,
            "space_after_pt": 0,
            "alignment": "justify",
            "first_line_indent_chars": 2,
            "left_indent_cm": 0,
            "right_indent_cm": 0,
            "keep_together": False,
            "keep_with_next": False,
            "widow_control": True,
        }

    # ==================== 表格样式 ====================

    def _extract_table_style(self, doc: Document) -> dict | None:
        """提取表格样式（完整版）"""
        if not doc.tables:
            return None

        table = doc.tables[0]

        # 提取所有边框信息
        borders = self._extract_all_table_borders(table)

        # 提取表格对齐方式
        table_alignment = self._extract_table_alignment(table)

        # 提取表头行重复
        header_repeat = self._extract_header_repeat(table)

        # 提取单元格垂直对齐
        cell_valign = self._extract_cell_vertical_alignment(table)

        # 提取单元格内边距
        cell_padding = self._extract_cell_padding(table)

        # 提取表头背景色
        header_bg_color = self._extract_header_bg_color(table)

        # 提取表头和正文字体
        header_font = self._extract_font_from_table_cell(table, 0, 0)
        body_font = self._extract_font_from_table_cell(table, 1, 0) if len(table.rows) > 1 else header_font

        # 判断表头是否加粗
        header_bold = self._extract_header_bold(table)

        # 确定主边框样式
        border_style = self._determine_border_style(borders)
        border_width = borders.get("top", {}).get("width", 0.5)

        result = {
            "border_style": border_style,
            "border_width_pt": border_width,
            "border_width_top_pt": borders.get("top", {}).get("width"),
            "border_width_bottom_pt": borders.get("bottom", {}).get("width"),
            "border_width_left_pt": borders.get("left", {}).get("width"),
            "border_width_right_pt": borders.get("right", {}).get("width"),
            "border_width_inside_h_pt": borders.get("insideH", {}).get("width"),
            "border_width_inside_v_pt": borders.get("insideV", {}).get("width"),
            "header_font": header_font,
            "body_font": body_font,
            "header_bold": header_bold,
            "header_bg_color": header_bg_color,
            "header_repeat": header_repeat,
            "cell_padding_top_pt": cell_padding.get("top", 2.0),
            "cell_padding_bottom_pt": cell_padding.get("bottom", 2.0),
            "cell_padding_left_pt": cell_padding.get("left", 2.0),
            "cell_padding_right_pt": cell_padding.get("right", 2.0),
            "table_alignment": table_alignment,
            "cell_vertical_alignment": cell_valign,
            "rag_note": None,
        }

        return result

    def _extract_all_table_borders(self, table) -> dict:
        """提取表格所有边框信息"""
        borders: dict[str, dict] = {}

        try:
            tbl = table._tbl
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is None:
                return borders

            tblBorders = tblPr.find(qn("w:tblBorders"))
            if tblBorders is None:
                return borders

            for name in TABLE_BORDER_NAMES:
                elem = tblBorders.find(qn(f"w:{name}"))
                if elem is not None:
                    val = elem.get(qn("w:val"), "single")
                    sz = elem.get(qn("w:sz"), "4")
                    borders[name] = {
                        "style": "double" if val == "double" else "single" if val in ("single", "thick") else val,
                        "width": round(int(sz) / 8, 1),
                    }
        except Exception:
            pass

        return borders

    def _determine_border_style(self, borders: dict) -> str:
        """根据边框信息判断表格边框样式"""
        if not borders:
            return "single"

        # 检查是否有三线表特征：有上下边框，无左右和内部边框
        has_top = "top" in borders and borders["top"]["style"] != "none"
        has_bottom = "bottom" in borders and borders["bottom"]["style"] != "none"
        has_left = "left" in borders and borders["left"]["style"] != "none"
        has_right = "right" in borders and borders["right"]["style"] != "none"
        has_inside_h = "insideH" in borders and borders["insideH"]["style"] != "none"
        has_inside_v = "insideV" in borders and borders["insideV"]["style"] != "none"

        # 三线表：有上下，无左右，无内部垂直
        if has_top and has_bottom and not has_left and not has_right and not has_inside_v:
            return "three-line"

        # 检查是否双线
        for b in borders.values():
            if b["style"] == "double":
                return "double"

        # 检查是否无边框
        all_none = all(b["style"] == "none" for b in borders.values())
        if all_none:
            return "none"

        return "single"

    def _extract_table_alignment(self, table) -> str:
        """提取表格对齐方式"""
        try:
            tbl = table._tbl
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is not None:
                tblpPr = tblPr.find(qn("w:tblpPr"))
                if tblpPr is not None:
                    tbl_align = tblpPr.get(qn("w:tblAlign"))
                    if tbl_align:
                        return TABLE_ALIGN_MAP.get(tbl_align, "left")
                # jc 元素
                jc = tblPr.find(qn("w:jc"))
                if jc is not None:
                    val = jc.get(qn("w:val"))
                    if val:
                        return TABLE_ALIGN_MAP.get(val, "left")
        except Exception:
            pass
        return "left"

    def _extract_header_repeat(self, table) -> bool:
        """提取表头行是否跨页重复"""
        try:
            if not table.rows:
                return False
            first_row = table.rows[0]
            tr = first_row._tr
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                tblHeader = trPr.find(qn("w:tblHeader"))
                if tblHeader is not None:
                    return True
        except Exception:
            pass
        return False

    def _extract_cell_vertical_alignment(self, table) -> str:
        """提取单元格垂直对齐方式"""
        try:
            if not table.rows or not table.rows[0].cells:
                return "center"
            cell = table.rows[0].cells[0]
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                vAlign = tcPr.find(qn("w:vAlign"))
                if vAlign is not None:
                    val = vAlign.get(qn("w:val"))
                    if val:
                        return CELL_VALIGN_MAP.get(val, "center")
        except Exception:
            pass
        return "center"

    def _extract_cell_padding(self, table) -> dict:
        """提取单元格内边距"""
        padding = {}
        try:
            if not table.rows or not table.rows[0].cells:
                return padding
            cell = table.rows[0].cells[0]
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                tcMar = tcPr.find(qn("w:tcMar"))
                if tcMar is not None:
                    for side in ["top", "bottom", "start", "end", "left", "right"]:
                        elem = tcMar.find(qn(f"w:{side}"))
                        if elem is not None:
                            w = elem.get(qn("w:w"))
                            if w:
                                pt_val = _twips_to_pt(w)
                                # start=left, end=right
                                if side in ("start", "left"):
                                    padding["left"] = pt_val
                                elif side in ("end", "right"):
                                    padding["right"] = pt_val
                                elif side == "top":
                                    padding["top"] = pt_val
                                elif side == "bottom":
                                    padding["bottom"] = pt_val
        except Exception:
            pass
        return padding

    def _extract_header_bg_color(self, table) -> str | None:
        """提取表头行背景色"""
        try:
            if not table.rows or not table.rows[0].cells:
                return None
            cell = table.rows[0].cells[0]
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    fill = shd.get(qn("w:fill"))
                    if fill and fill != "auto":
                        return f"#{fill}"
        except Exception:
            pass
        return None

    def _extract_header_bold(self, table) -> bool:
        """从表头行第一个非空单元格提取是否加粗"""
        try:
            if not table.rows:
                return True
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    for run in paragraph.runs:
                        if run.font.bold is not None:
                            return run.font.bold
                    # 检查样式级加粗
                    if paragraph.style and paragraph.style.font:
                        if paragraph.style.font.bold is not None:
                            return paragraph.style.font.bold
        except Exception:
            pass
        return True

    # ==================== 列表样式 ====================

    def _extract_list_style(self, doc: Document, style_defs: dict) -> dict | None:
        """提取列表样式"""
        # 先从样式定义表获取基线
        base = dict(style_defs.get("list", {})) if "list" in style_defs else None

        # 从文档段落中找列表段落
        list_paragraphs = []
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            if "List" in style_name or "list" in style_name:
                if paragraph.text.strip():
                    list_paragraphs.append(paragraph)

        if not list_paragraphs and not base:
            return None

        if not list_paragraphs:
            return base

        # 采样
        sample = list_paragraphs[: min(10, len(list_paragraphs))]
        font_configs = [self._extract_font_from_paragraph(p) for p in sample]
        alignments = [self._extract_alignment(p) for p in sample]
        line_spacings = [self._extract_line_spacing_full(p) for p in sample]
        space_befores, space_afters = zip(*[self._extract_spacing(p) for p in sample])
        first_indents = [self._extract_first_indent(p) for p in sample]
        left_indents = [self._extract_left_indent(p) for p in sample]

        result = dict(base) if base else {}
        result["font"] = self._merge_font_configs(font_configs)
        result["alignment"] = Counter(alignments).most_common(1)[0][0]
        ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
        result["line_spacing"] = ls
        result["line_spacing_pt"] = ls_pt
        result["line_spacing_rule"] = ls_rule
        result["space_before_pt"] = Counter(space_befores).most_common(1)[0][0]
        result["space_after_pt"] = Counter(space_afters).most_common(1)[0][0]
        result["first_line_indent_chars"] = Counter(first_indents).most_common(1)[0][0]
        result["left_indent_cm"] = Counter(left_indents).most_common(1)[0][0]
        result.setdefault("right_indent_cm", 0)
        result.setdefault("keep_together", False)
        result.setdefault("keep_with_next", False)
        result.setdefault("widow_control", True)
        return result

    # ==================== 脚注样式 ====================

    def _extract_footnote_style(self, doc: Document, style_defs: dict) -> dict | None:
        """提取脚注样式"""
        base = dict(style_defs.get("footnote", {})) if "footnote" in style_defs else None

        # 从文档中找脚注段落（python-docx 不直接暴露脚注，从样式定义获取）
        if not base:
            return None

        base.setdefault("right_indent_cm", 0)
        base.setdefault("keep_together", False)
        base.setdefault("keep_with_next", False)
        base.setdefault("widow_control", True)
        return base

    # ==================== 页眉页脚样式 ====================

    def _extract_header_footer_style(self, doc: Document, style_defs: dict) -> dict | None:
        """提取页眉页脚样式"""
        base = dict(style_defs.get("header_footer", {})) if "header_footer" in style_defs else None

        # 从 section 的 header/footer 中提取段落样式
        header_footer_paragraphs = []
        for section in doc.sections:
            try:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        header_footer_paragraphs.append(paragraph)
            except Exception:
                pass
            try:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        header_footer_paragraphs.append(paragraph)
            except Exception:
                pass

        if not header_footer_paragraphs and not base:
            return None

        if not header_footer_paragraphs:
            return base

        sample = header_footer_paragraphs[: min(5, len(header_footer_paragraphs))]
        font_configs = [self._extract_font_from_paragraph(p) for p in sample]
        alignments = [self._extract_alignment(p) for p in sample]
        line_spacings = [self._extract_line_spacing_full(p) for p in sample]

        result = dict(base) if base else {}
        result["font"] = self._merge_font_configs(font_configs)
        result["alignment"] = Counter(alignments).most_common(1)[0][0]
        ls, ls_pt, ls_rule = self._most_common_tuple(line_spacings)
        result["line_spacing"] = ls
        result["line_spacing_pt"] = ls_pt
        result["line_spacing_rule"] = ls_rule
        result.setdefault("space_before_pt", 0)
        result.setdefault("space_after_pt", 0)
        result.setdefault("first_line_indent_chars", 0)
        result.setdefault("left_indent_cm", 0)
        result.setdefault("right_indent_cm", 0)
        result.setdefault("keep_together", False)
        result.setdefault("keep_with_next", False)
        result.setdefault("widow_control", True)
        return result

    # ==================== 字体提取（核心增强） ====================

    def _extract_font_from_paragraph(self, paragraph) -> dict:
        """从段落中提取字体配置（增强版）

        改进：
        - 采样多个 run，取最常见的字体
        - 正确提取东亚字体（从 run 的 rPr/rFonts）
        - 提取下划线、删除线
        """
        font_config = {
            "family": "仿宋_GB2312",
            "east_asia_family": None,
            "size_pt": 16,
            "bold": False,
            "italic": False,
            "underline": False,
            "strikethrough": False,
            "color_hex": "#000000",
        }

        if not paragraph.runs:
            # 从段落样式继承
            self._inherit_font_from_style(paragraph, font_config)
            return font_config

        # 采样所有 run，收集各项属性
        families: list[str] = []
        east_asia_families: list[str] = []
        sizes: list[float] = []
        bolds: list[bool] = []
        italics: list[bool] = []
        underlines: list[bool] = []
        strikes: list[bool] = []
        colors: list[str] = []

        for run in paragraph.runs:
            if run.font.name:
                families.append(run.font.name)
            if run.font.size:
                sizes.append(_emu_to_pt(run.font.size))
            if run.font.bold is not None:
                bolds.append(run.font.bold)
            if run.font.italic is not None:
                italics.append(run.font.italic)
            if run.font.underline is not None:
                underlines.append(bool(run.font.underline))
            if run.font.color and run.font.color.rgb:
                colors.append(_rgb_to_hex(run.font.color.rgb))

            # 从 run 的 XML 提取东亚字体（正确路径：rPr/rFonts）
            east_asia = self._get_run_east_asia_font(run)
            if east_asia:
                east_asia_families.append(east_asia)

            # 从 run 的 XML 提取删除线
            strike = self._get_run_strikethrough(run)
            if strike is not None:
                strikes.append(strike)

        # 取最常见的值
        if families:
            font_config["family"] = Counter(families).most_common(1)[0][0]
        if east_asia_families:
            font_config["east_asia_family"] = Counter(east_asia_families).most_common(1)[0][0]
        if sizes:
            font_config["size_pt"] = Counter(sizes).most_common(1)[0][0]
        if bolds:
            font_config["bold"] = Counter(bolds).most_common(1)[0][0]
        if italics:
            font_config["italic"] = Counter(italics).most_common(1)[0][0]
        if underlines:
            font_config["underline"] = Counter(underlines).most_common(1)[0][0]
        if strikes:
            font_config["strikethrough"] = Counter(strikes).most_common(1)[0][0]
        if colors:
            font_config["color_hex"] = Counter(colors).most_common(1)[0][0]

        # 如果没有从 run 提取到字体，从段落样式继承
        if not families and not east_asia_families:
            self._inherit_font_from_style(paragraph, font_config)

        # 如果没有东亚字体，用西文字体作为默认
        if not font_config["east_asia_family"]:
            font_config["east_asia_family"] = font_config["family"]

        return font_config

    def _get_run_east_asia_font(self, run) -> str | None:
        """从 run 的 XML 元素提取东亚字体名称"""
        try:
            r_element = run._element
            rpr = r_element.find(qn("w:rPr"))
            if rpr is not None:
                r_fonts = rpr.find(qn("w:rFonts"))
                if r_fonts is not None:
                    east_asia = r_fonts.get(qn("w:eastAsia"))
                    if east_asia:
                        return east_asia
                    # 如果没有 eastAsia，尝试 hAnsi 或 ascii
                    ascii_font = r_fonts.get(qn("w:ascii"))
                    if ascii_font:
                        return ascii_font
        except Exception:
            pass
        return None

    def _get_run_strikethrough(self, run) -> bool | None:
        """从 run 的 XML 提取删除线"""
        try:
            r_element = run._element
            rpr = r_element.find(qn("w:rPr"))
            if rpr is not None:
                strike = rpr.find(qn("w:strike"))
                if strike is not None:
                    val = strike.get(qn("w:val"))
                    return val != "0" and val != "false"
        except Exception:
            pass
        return None

    def _inherit_font_from_style(self, paragraph, font_config: dict) -> None:
        """从段落样式继承字体配置"""
        try:
            style = paragraph.style
            if style and style.font:
                if style.font.name:
                    font_config["family"] = style.font.name
                if style.font.size:
                    font_config["size_pt"] = _emu_to_pt(style.font.size)
                if style.font.bold is not None:
                    font_config["bold"] = style.font.bold
                if style.font.italic is not None:
                    font_config["italic"] = style.font.italic
                if style.font.color and style.font.color.rgb:
                    font_config["color_hex"] = _rgb_to_hex(style.font.color.rgb)
        except Exception:
            pass

        # 从样式 XML 提取东亚字体
        try:
            style_element = paragraph.style.element
            rpr = style_element.find(qn("w:rPr"))
            if rpr is not None:
                r_fonts = rpr.find(qn("w:rFonts"))
                if r_fonts is not None:
                    east_asia = r_fonts.get(qn("w:eastAsia"))
                    if east_asia:
                        font_config["east_asia_family"] = east_asia
                    ascii_font = r_fonts.get(qn("w:ascii"))
                    if ascii_font:
                        font_config["family"] = ascii_font
        except Exception:
            pass

    def _merge_font_configs(self, configs: list[dict]) -> dict:
        """合并多个字体配置，取最常见的值"""
        if not configs:
            return {
                "family": "仿宋_GB2312",
                "east_asia_family": "仿宋_GB2312",
                "size_pt": 16,
                "bold": False,
                "italic": False,
                "underline": False,
                "strikethrough": False,
                "color_hex": "#000000",
            }

        result: dict = {}
        keys = ["family", "east_asia_family", "size_pt", "bold", "italic", "underline", "strikethrough", "color_hex"]

        for key in keys:
            values = [c[key] for c in configs if key in c and c[key] is not None]
            if values:
                result[key] = Counter(values).most_common(1)[0][0]
            else:
                # 默认值
                defaults = {
                    "family": "仿宋_GB2312",
                    "east_asia_family": "仿宋_GB2312",
                    "size_pt": 16,
                    "bold": False,
                    "italic": False,
                    "underline": False,
                    "strikethrough": False,
                    "color_hex": "#000000",
                }
                result[key] = defaults[key]

        return result

    # ==================== 段落格式提取 ====================

    def _extract_font_from_table_cell(self, table, row_idx: int, col_idx: int) -> dict:
        """从表格单元格提取字体配置"""
        default_font = {
            "family": "宋体",
            "east_asia_family": "宋体",
            "size_pt": 12,
            "bold": False,
            "italic": False,
            "underline": False,
            "strikethrough": False,
            "color_hex": "#000000",
        }

        try:
            cell = table.rows[row_idx].cells[col_idx]
            for paragraph in cell.paragraphs:
                if not paragraph.text.strip():
                    continue
                return self._extract_font_from_paragraph(paragraph)
        except Exception:
            pass

        return default_font

    def _extract_alignment(self, paragraph) -> str:
        """提取对齐方式"""
        try:
            if paragraph.alignment is not None:
                return ALIGNMENT_REVERSE_MAP.get(paragraph.alignment, "left")
        except Exception:
            pass
        # 从样式继承
        try:
            if paragraph.style and paragraph.style.paragraph_format:
                if paragraph.style.paragraph_format.alignment is not None:
                    return ALIGNMENT_REVERSE_MAP.get(
                        paragraph.style.paragraph_format.alignment, "left"
                    )
        except Exception:
            pass
        return "left"

    def _extract_line_spacing_full(self, paragraph) -> tuple[float, float | None, str]:
        """提取完整行距信息

        Returns:
            (line_spacing, line_spacing_pt, line_spacing_rule)
        """
        try:
            pf = paragraph.paragraph_format
            if pf and pf.line_spacing is not None:
                ls = pf.line_spacing
                # float = 倍数行距
                if isinstance(ls, float):
                    return (float(ls), None, "multiple")
                # int 或 EMU = 固定行距
                elif isinstance(ls, int):
                    return (1.0, _emu_to_pt(ls), "exact")
                else:
                    # Pt 对象
                    try:
                        return (1.0, float(ls), "exact")
                    except Exception:
                        return (float(ls), None, "multiple")
        except Exception:
            pass

        # 从样式继承
        try:
            if paragraph.style and paragraph.style.paragraph_format:
                pf = paragraph.style.paragraph_format
                if pf.line_spacing is not None:
                    ls = pf.line_spacing
                    if isinstance(ls, float):
                        return (float(ls), None, "multiple")
                    elif isinstance(ls, int):
                        return (1.0, _emu_to_pt(ls), "exact")
                    else:
                        try:
                            return (1.0, float(ls), "exact")
                        except Exception:
                            return (float(ls), None, "multiple")
        except Exception:
            pass

        return (1.5, None, "multiple")

    def _extract_line_spacing(self, paragraph) -> float:
        """提取行距（向后兼容）"""
        ls, _, _ = self._extract_line_spacing_full(paragraph)
        return ls

    def _extract_spacing(self, paragraph) -> tuple[float, float]:
        """提取段前段后间距 (pt)"""
        space_before = 0.0
        space_after = 0.0
        try:
            pf = paragraph.paragraph_format
            if pf:
                if pf.space_before is not None:
                    space_before = _emu_to_pt(pf.space_before)
                if pf.space_after is not None:
                    space_after = _emu_to_pt(pf.space_after)
        except Exception:
            pass

        # 从样式继承
        if space_before == 0 and space_after == 0:
            try:
                if paragraph.style and paragraph.style.paragraph_format:
                    pf = paragraph.style.paragraph_format
                    if pf:
                        if pf.space_before is not None:
                            space_before = _emu_to_pt(pf.space_before)
                        if pf.space_after is not None:
                            space_after = _emu_to_pt(pf.space_after)
            except Exception:
                pass

        return space_before, space_after

    def _extract_first_indent(self, paragraph) -> float:
        """提取首行缩进（字符数）"""
        try:
            pf = paragraph.paragraph_format
            if pf and pf.first_line_indent:
                indent_pt = _emu_to_pt(pf.first_line_indent)
                font_size = 16
                for run in paragraph.runs:
                    if run.font.size:
                        font_size = _emu_to_pt(run.font.size)
                        break
                if font_size > 0:
                    return round(indent_pt / font_size, 1)
        except Exception:
            pass

        # 从样式继承
        try:
            if paragraph.style and paragraph.style.paragraph_format:
                pf = paragraph.style.paragraph_format
                if pf and pf.first_line_indent:
                    indent_pt = _emu_to_pt(pf.first_line_indent)
                    font_size = 16
                    style_font = paragraph.style.font
                    if style_font and style_font.size:
                        font_size = _emu_to_pt(style_font.size)
                    if font_size > 0:
                        return round(indent_pt / font_size, 1)
        except Exception:
            pass

        return 0

    def _extract_left_indent(self, paragraph) -> float:
        """提取左缩进（厘米）"""
        try:
            pf = paragraph.paragraph_format
            if pf and pf.left_indent is not None:
                return _emu_to_cm(pf.left_indent)
        except Exception:
            pass
        return 0

    def _extract_right_indent(self, paragraph) -> float:
        """提取右缩进（厘米）"""
        try:
            pf = paragraph.paragraph_format
            if pf and pf.right_indent is not None:
                return _emu_to_cm(pf.right_indent)
        except Exception:
            pass
        return 0

    def _extract_numbering_format(self, paragraph) -> str | None:
        """提取编号格式"""
        try:
            # 检查段落是否有编号属性
            p_element = paragraph._element
            numPr = p_element.find(qn("w:pPr"))
            if numPr is not None:
                numPr_inner = numPr.find(qn("w:numPr"))
                if numPr_inner is not None:
                    # 有编号，但具体格式需要从 numbering.xml 解析
                    # 这里返回一个标记
                    return "auto"
        except Exception:
            pass
        return None

    def _extract_outline_level(self, paragraph) -> int | None:
        """提取大纲级别"""
        try:
            p_element = paragraph._element
            pPr = p_element.find(qn("w:pPr"))
            if pPr is not None:
                outlineLvl = pPr.find(qn("w:outlineLvl"))
                if outlineLvl is not None:
                    val = outlineLvl.get(qn("w:val"))
                    if val is not None:
                        return int(val) + 1  # 0-based → 1-based
        except Exception:
            pass
        return None

    @staticmethod
    def _most_common_tuple(tuples: list) -> tuple:
        """取最常见的元组"""
        if not tuples:
            return (1.5, None, "multiple")
        counter = Counter(tuples)
        return counter.most_common(1)[0][0]
