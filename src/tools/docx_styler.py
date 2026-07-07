"""python-docx 样式渲染工具

将 StyleConfig 应用到 DOCX 文档。纯工具函数，无 LLM 参与。
严格遵循架构解耦原则：所有样式参数来自 StyleConfig，不自行推断。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

from src.models.document_schema import StyleReport
from src.models.style_config import (
    FontConfig,
    HeadingStyleConfig,
    ParagraphStyleConfig,
    StyleConfig,
)
from src.utils.file_utils import ensure_dir
from src.utils.logger import get_logger

logger = get_logger(__name__)

# Pandoc 生成的 DOCX 中，标题段落样式的内置名称映射
HEADING_STYLE_MAP = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "heading 1": 1,
    "heading 2": 2,
    "heading 3": 3,
    "heading 4": 4,
    "heading 5": 5,
    "heading 6": 6,
}

# 对齐方式映射
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxStyler:
    """DOCX 样式渲染器

    将 StyleConfig 中的排版参数应用到 python-docx Document 对象。
    """

    def __init__(self, style_config: StyleConfig):
        """初始化样式渲染器

        Args:
            style_config: 排版配置（来自 LLM 输出或 RAG 规范）
        """
        self.config = style_config
        self._warnings: list[str] = []

    def apply_gb_style(self, docx_path: str | Path, output_path: str | Path) -> StyleReport:
        """核心方法：应用国标样式到 DOCX 文件

        处理流程：
        1. 打开 DOCX 文件
        2. 设置页面布局
        3. 遍历段落，检测角色并应用对应样式
        4. 处理表格样式
        5. 保存输出

        Args:
            docx_path: 输入的 DOCX 文件路径
            output_path: 输出的 DOCX 文件路径

        Returns:
            StyleReport 样式应用报告
        """
        docx_path = Path(docx_path)
        output_path = Path(output_path)

        if not docx_path.exists():
            return StyleReport(
                success=False,
                warnings=[f"输入文件不存在: {docx_path}"],
                output_path=str(output_path),
            )

        ensure_dir(output_path.parent)
        self._warnings = []

        try:
            doc = Document(str(docx_path))
        except Exception as e:
            return StyleReport(
                success=False,
                warnings=[f"无法打开 DOCX 文件: {e}"],
                output_path=str(output_path),
            )

        # 1. 应用页面布局
        self._apply_page_layout(doc)

        # 2. 遍历段落应用样式
        paragraphs_styled = 0
        headings_styled = 0

        for paragraph in doc.paragraphs:
            role = self._detect_paragraph_role(paragraph)

            if role.startswith("heading_"):
                level = int(role.split("_")[1])
                heading_style = self.config.get_heading_style(level)
                if heading_style:
                    self._apply_heading_style(paragraph, heading_style)
                    headings_styled += 1
                else:
                    self._apply_paragraph_style(paragraph, self.config.body_style)
                    self._warnings.append(f"标题层级 {level} 无对应样式配置，使用正文样式")
            elif role == "body":
                self._apply_paragraph_style(paragraph, self.config.body_style)
                paragraphs_styled += 1
            elif role == "list_item":
                list_style = self.config.list_style or self.config.body_style
                self._apply_paragraph_style(paragraph, list_style)
                paragraphs_styled += 1
            elif role == "caption":
                caption_style = self.config.caption_style or self.config.body_style
                self._apply_paragraph_style(paragraph, caption_style)
                paragraphs_styled += 1
            elif role == "table_caption":
                # 表格标题样式（表B.1/表1等），优先使用专用样式，回退到 caption 再到正文
                tc_style = self.config.table_caption_style or self.config.caption_style or self.config.body_style
                self._apply_paragraph_style(paragraph, tc_style)
                paragraphs_styled += 1
            elif role == "preface":
                # 前言/引言标题样式，回退到一级标题样式
                preface_style = self.config.preface_style or self.config.get_heading_style(1)
                if preface_style:
                    self._apply_paragraph_style(paragraph, preface_style)
                    headings_styled += 1
                else:
                    self._apply_paragraph_style(paragraph, self.config.body_style)
            elif role == "appendix_title":
                # 附录标题样式（加粗，区别于普通一级标题），回退到一级标题
                app_style = self.config.appendix_title_style or self.config.get_heading_style(1)
                if app_style:
                    self._apply_paragraph_style(paragraph, app_style)
                    headings_styled += 1
                else:
                    self._apply_paragraph_style(paragraph, self.config.body_style)
            elif role.startswith("appendix_clause"):
                # 附录内条款样式（A.1/B.1等），提取层级回退到对应标题样式
                level_str = role.split("_")[-1]
                clause_level = int(level_str) if level_str.isdigit() else 2
                clause_style = (
                    self.config.appendix_clause_style
                    or self.config.get_heading_style(clause_level)
                    or self.config.body_style
                )
                self._apply_paragraph_style(paragraph, clause_style)
                headings_styled += 1
            elif role == "toc":
                self._apply_paragraph_style(paragraph, self.config.body_style)
                paragraphs_styled += 1

        # 3. 处理表格样式
        tables_styled = 0
        if self.config.table_style:
            for table in doc.tables:
                self._apply_table_style(table)
                tables_styled += 1

        # 4. 保存
        try:
            doc.save(str(output_path))
            logger.info(
                "样式应用完成: %d 段落, %d 标题, %d 表格 → %s",
                paragraphs_styled,
                headings_styled,
                tables_styled,
                output_path,
            )
        except Exception as e:
            return StyleReport(
                success=False,
                paragraphs_styled=paragraphs_styled,
                tables_styled=tables_styled,
                headings_styled=headings_styled,
                warnings=self._warnings + [f"保存失败: {e}"],
                output_path=str(output_path),
            )

        return StyleReport(
            success=True,
            paragraphs_styled=paragraphs_styled,
            tables_styled=tables_styled,
            headings_styled=headings_styled,
            warnings=self._warnings,
            output_path=str(output_path),
        )

    def _apply_page_layout(self, document: Document) -> None:
        """应用页面布局配置

        Args:
            document: python-docx Document 对象
        """
        layout = self.config.page_layout

        for section in document.sections:
            # 纸张大小
            if layout.paper_size == "A4":
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
            elif layout.paper_size == "A3":
                section.page_width = Cm(29.7)
                section.page_height = Cm(42.0)
            elif layout.paper_size == "B5":
                section.page_width = Cm(17.6)
                section.page_height = Cm(25.0)
            elif layout.paper_size == "Letter":
                section.page_width = Cm(21.59)
                section.page_height = Cm(27.94)

            # 页面方向
            if layout.orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                # 交换宽高
                section.page_width, section.page_height = section.page_height, section.page_width
            else:
                section.orientation = WD_ORIENT.PORTRAIT

            # 页边距
            section.top_margin = Cm(layout.margin_top_cm)
            section.bottom_margin = Cm(layout.margin_bottom_cm)
            section.left_margin = Cm(layout.margin_left_cm)
            section.right_margin = Cm(layout.margin_right_cm)
            section.header_distance = Cm(layout.header_distance_cm)
            section.footer_distance = Cm(layout.footer_distance_cm)

        logger.debug("页面布局已应用: %s, 纵向", layout.paper_size)

    def _apply_paragraph_style(self, paragraph, style: ParagraphStyleConfig) -> None:
        """应用段落样式

        Args:
            paragraph: python-docx Paragraph 对象
            style: 段落样式配置
        """
        # 对齐方式
        paragraph.alignment = ALIGNMENT_MAP.get(style.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # 段落格式
        pf = paragraph.paragraph_format
        # 行距：支持倍数行距和固定行距
        if style.line_spacing_rule == "exact" and style.line_spacing_pt is not None:
            pf.line_spacing = Pt(style.line_spacing_pt)
        elif style.line_spacing_rule == "at_least" and style.line_spacing_pt is not None:
            from docx.enum.text import WD_LINE_SPACING

            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            pf.line_spacing = Pt(style.line_spacing_pt)
        else:
            pf.line_spacing = style.line_spacing
        pf.space_before = Pt(style.space_before_pt)
        pf.space_after = Pt(style.space_after_pt)

        # 首行缩进
        if style.first_line_indent_chars > 0:
            indent_pt = style.first_line_indent_chars * style.font.size_pt
            pf.first_line_indent = Pt(indent_pt)

        # 左右缩进
        if style.left_indent_cm > 0:
            pf.left_indent = Cm(style.left_indent_cm)
        if style.right_indent_cm > 0:
            pf.right_indent = Cm(style.right_indent_cm)

        # 段落分页控制
        if hasattr(style, "keep_together") and style.keep_together:
            pf.keep_together = True
        if hasattr(style, "widow_control"):
            pf.widow_control = style.widow_control

        # 应用字体到所有 run
        for run in paragraph.runs:
            self._apply_font_to_run(run, style.font)

    def _apply_heading_style(self, paragraph, style: HeadingStyleConfig) -> None:
        """应用标题样式

        Args:
            paragraph: python-docx Paragraph 对象
            style: 标题样式配置
        """
        # 应用段落级样式
        self._apply_paragraph_style(paragraph, style)

        # 段中不分页（标题默认启用）
        paragraph.paragraph_format.keep_with_next = True

    def _apply_table_style(self, table: Table) -> None:
        """应用表格样式

        Args:
            table: python-docx Table 对象
        """
        ts = self.config.table_style
        if not ts:
            return

        # 设置表格对齐方式
        self._set_table_alignment(table, ts.table_alignment)

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 表头使用表头字体，其余用正文字体
                    if row_idx == 0:
                        self._apply_font_to_run_all(
                            paragraph, ts.header_font, bold=ts.header_bold if ts.header_bold else None
                        )
                        # 表头背景色
                        if ts.header_bg_color:
                            self._set_cell_shading(cell, ts.header_bg_color)
                    else:
                        self._apply_font_to_run_all(paragraph, ts.body_font)

                # 设置单元格垂直对齐
                self._set_cell_vertical_alignment(cell, ts.cell_vertical_alignment)

                # 设置单元格内边距（四边独立）
                self._set_cell_margins_full(cell, ts)

        # 表头行跨页重复
        if ts.header_repeat and table.rows:
            self._set_header_repeat(table)

        # 设置边框样式
        if ts.border_style != "none":
            self._set_table_borders_full(table, ts)

    def _apply_font_to_run(self, run, font_config: FontConfig) -> None:
        """将字体配置应用到单个 run

        Args:
            run: python-docx Run 对象
            font_config: 字体配置
        """
        run.font.size = Pt(font_config.size_pt)
        run.font.bold = font_config.bold
        run.font.italic = font_config.italic
        run.font.underline = font_config.underline if hasattr(font_config, "underline") else False

        # 设置西文字体
        run.font.name = font_config.family
        # 使用 XML 属性设置东亚字体（中文）
        r_element = run._element
        rpr = r_element.find(qn("w:rPr"))
        if rpr is None:
            rpr = r_element.makeelement(qn("w:rPr"), {})
            r_element.insert(0, rpr)
        r_fonts = rpr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), font_config.family)
        r_fonts.set(qn("w:hAnsi"), font_config.family)
        # 东亚字体：优先使用 east_asia_family，否则回退到 family
        east_asia_font = font_config.east_asia_family or font_config.family
        r_fonts.set(qn("w:eastAsia"), east_asia_font)

        # 删除线
        if hasattr(font_config, "strikethrough") and font_config.strikethrough:
            strike = rpr.find(qn("w:strike"))
            if strike is None:
                strike = rpr.makeelement(qn("w:strike"), {})
                rpr.append(strike)
            strike.set(qn("w:val"), "true")

        # 颜色
        if font_config.color_hex and font_config.color_hex != "#000000":
            hex_color = font_config.color_hex.lstrip("#")
            run.font.color.rgb = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )

        # 字体可用性检查
        if not self._ensure_font_installed(font_config.family):
            warning = f"字体可能未安装: {font_config.family}"
            if warning not in self._warnings:
                self._warnings.append(warning)

    def _apply_font_to_run_all(self, paragraph, font_config: FontConfig, bold: bool | None = None) -> None:
        """将字体配置应用到段落的所有 run

        Args:
            paragraph: python-docx Paragraph 对象
            font_config: 字体配置
            bold: 覆盖加粗设置
        """
        for run in paragraph.runs:
            self._apply_font_to_run(run, font_config)
            if bold is not None:
                run.font.bold = bold

    def _detect_paragraph_role(self, paragraph) -> str:
        """检测段落角色

        先基于 Pandoc 生成的 DOCX 内置样式名判断，
        再使用内容模式识别作为后备方案（国标文档常按编号识别标题）。

        Args:
            paragraph: python-docx Paragraph 对象

        Returns:
            角色字符串: heading_1..6 / body / list_item / toc / caption /
            preface / appendix_title / appendix_clause_2..5 / table_caption / empty
        """
        style_name = paragraph.style.name if paragraph.style else ""

        # 1. 先检查 Pandoc 样式名
        if style_name in HEADING_STYLE_MAP:
            level = HEADING_STYLE_MAP[style_name]
            return f"heading_{level}"

        if "List" in style_name or "list" in style_name:
            return "list_item"

        if "TOC" in style_name or "toc" in style_name:
            return "toc"

        if "Caption" in style_name or "caption" in style_name:
            return "caption"

        # 2. 内容模式识别（后备方案：国标文档常按编号识别标题）
        text = paragraph.text.strip()
        if text:
            from src.tools.content_pattern_matcher import classify_paragraph_role

            role = classify_paragraph_role(text)
            if role:
                return role

        # 3. 空段落
        if not text:
            return "empty"

        return "body"

    def _ensure_font_installed(self, font_name: str) -> bool:
        """检查字体是否在已知国标字体白名单中

        Args:
            font_name: 字体名称

        Returns:
            是否在白名单中（不在白名单的字体记录警告）
        """
        known_fonts = {
            "黑体",
            "宋体",
            "仿宋",
            "楷体",
            "微软雅黑",
            "仿宋_GB2312",
            "楷体_GB2312",
            "方正小标宋简体",
            "Times New Roman",
            "Arial",
            "Calibri",
        }
        return font_name in known_fonts

    def _set_cell_margins(self, cell, padding_pt: float) -> None:
        """设置表格单元格内边距（向后兼容）

        Args:
            cell: python-docx Cell 对象
            padding_pt: 内边距（磅）
        """
        self._set_cell_margins_full(cell, None, padding_pt)

    def _set_cell_margins_full(self, cell, ts=None, default_padding: float = 2.0) -> None:
        """设置表格单元格内边距（四边独立）

        Args:
            cell: python-docx Cell 对象
            ts: TableStyleConfig 对象，如果 None 则使用 default_padding
            default_padding: 默认内边距（磅）
        """
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = tc.makeelement(qn("w:tcPr"), {})
            tc.insert(0, tcPr)

        # 获取各边内边距
        if ts:
            top = ts.cell_padding_top_pt
            bottom = ts.cell_padding_bottom_pt
            left = ts.cell_padding_left_pt
            right = ts.cell_padding_right_pt
        else:
            top = bottom = left = right = default_padding

        # 创建或更新 tcMar
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is None:
            tcMar = tcPr.makeelement(qn("w:tcMar"), {})
            tcPr.append(tcMar)

        # 设置各边边距（单位为 twips，1 pt = 20 twips）
        for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
            elem = tcMar.find(qn(f"w:{side}"))
            if elem is None:
                elem = tcMar.makeelement(qn(f"w:{side}"), {})
                tcMar.append(elem)
            elem.set(qn("w:w"), str(int(val * 20)))
            elem.set(qn("w:type"), "dxa")

    def _set_table_alignment(self, table: Table, alignment: str) -> None:
        """设置表格对齐方式"""
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = tbl.makeelement(qn("w:tblPr"), {})
            tbl.insert(0, tblPr)

        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = tblPr.makeelement(qn("w:jc"), {})
            tblPr.append(jc)
        jc.set(qn("w:val"), alignment)

    def _set_cell_vertical_alignment(self, cell, alignment: str) -> None:
        """设置单元格垂直对齐方式"""
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = tc.makeelement(qn("w:tcPr"), {})
            tc.insert(0, tcPr)

        vAlign = tcPr.find(qn("w:vAlign"))
        if vAlign is None:
            vAlign = tcPr.makeelement(qn("w:vAlign"), {})
            tcPr.append(vAlign)
        vAlign.set(qn("w:val"), alignment)

    def _set_cell_shading(self, cell, color_hex: str) -> None:
        """设置单元格背景色"""
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = tc.makeelement(qn("w:tcPr"), {})
            tc.insert(0, tcPr)

        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = tcPr.makeelement(qn("w:shd"), {})
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex.lstrip("#"))

    def _set_header_repeat(self, table: Table) -> None:
        """设置表头行跨页重复"""
        if not table.rows:
            return
        tr = table.rows[0]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = tr.makeelement(qn("w:trPr"), {})
            tr.insert(0, trPr)
        tblHeader = trPr.find(qn("w:tblHeader"))
        if tblHeader is None:
            tblHeader = trPr.makeelement(qn("w:tblHeader"), {})
            trPr.append(tblHeader)

    def _set_table_borders_full(self, table: Table, ts) -> None:
        """设置表格边框（完整版，支持各边独立线宽）

        Args:
            table: python-docx Table 对象
            ts: TableStyleConfig
        """
        width_pt = ts.border_width_pt
        border_style_val = "single"
        if ts.border_style == "double":
            border_style_val = "double"
        elif ts.border_style == "three-line":
            border_style_val = "single"

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = tbl.makeelement(qn("w:tblPr"), {})
            tbl.insert(0, tblPr)

        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
            tblPr.append(tblBorders)

        # 各边线宽
        widths = {
            "top": ts.border_width_top_pt if ts.border_width_top_pt is not None else width_pt,
            "bottom": ts.border_width_bottom_pt if ts.border_width_bottom_pt is not None else width_pt,
            "left": ts.border_width_left_pt if ts.border_width_left_pt is not None else width_pt,
            "right": ts.border_width_right_pt if ts.border_width_right_pt is not None else width_pt,
            "insideH": ts.border_width_inside_h_pt if ts.border_width_inside_h_pt is not None else width_pt,
            "insideV": ts.border_width_inside_v_pt if ts.border_width_inside_v_pt is not None else width_pt,
        }

        for border_name in ["top", "bottom", "left", "right", "insideH", "insideV"]:
            border = tblBorders.find(qn(f"w:{border_name}"))
            if border is None:
                border = tblBorders.makeelement(qn(f"w:{border_name}"), {})
                tblBorders.append(border)

            border.set(qn("w:val"), border_style_val)
            border.set(qn("w:sz"), str(int(widths[border_name] * 8)))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")

        # 三线表特殊处理
        if ts.border_style == "three-line":
            thick_width = int(width_pt * 2 * 8)
            thin_width = int(0.5 * 8)
            for border_name in ["top", "bottom"]:
                border = tblBorders.find(qn(f"w:{border_name}"))
                if border is not None:
                    border.set(qn("w:sz"), str(thick_width))
            for border_name in ["insideH", "insideV", "left", "right"]:
                border = tblBorders.find(qn(f"w:{border_name}"))
                if border is not None:
                    border.set(qn("w:sz"), str(thin_width))
                    if border_name in ("left", "right"):
                        border.set(qn("w:val"), "none")

    def _set_table_borders(self, table: Table, width_pt: float, style: str) -> None:
        """设置表格边框

        Args:
            table: python-docx Table 对象
            width_pt: 边框宽度（磅）
            style: 边框样式
        """
        # 线宽单位：1/8 pt
        width_eighths = int(width_pt * 8)

        # 样式映射
        border_style = "single"
        if style == "double":
            border_style = "double"
        elif style == "three-line":
            border_style = "single"  # 三线表使用单线，但粗细不同

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = tbl.makeelement(qn("w:tblPr"), {})
            tbl.insert(0, tblPr)

        # 创建或更新边框元素
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
            tblPr.append(tblBorders)

        for border_name in ["top", "bottom", "left", "right", "insideH", "insideV"]:
            border = tblBorders.find(qn(f"w:{border_name}"))
            if border is None:
                border = tblBorders.makeelement(qn(f"w:{border_name}"), {})
                tblBorders.append(border)

            border.set(qn("w:val"), border_style)
            border.set(qn("w:sz"), str(width_eighths))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")

        # 三线表特殊处理：上下边框加粗，内边框变细
        if style == "three-line":
            thick_width = int(width_pt * 2 * 8)  # 上下边线加粗
            thin_width = int(0.5 * 8)  # 内边线变细

            for border_name in ["top", "bottom"]:
                border = tblBorders.find(qn(f"w:{border_name}"))
                if border is not None:
                    border.set(qn("w:sz"), str(thick_width))

            for border_name in ["insideH", "insideV", "left", "right"]:
                border = tblBorders.find(qn(f"w:{border_name}"))
                if border is not None:
                    border.set(qn("w:sz"), str(thin_width))
                    if border_name in ("left", "right"):
                        border.set(qn("w:val"), "none")  # 三线表无左右边线
