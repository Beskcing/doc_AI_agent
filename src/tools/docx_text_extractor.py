"""DOCX 文本提取工具

从排版完成的 DOCX 文件中提取结构化文本，保留段落类型、标题层级和位置信息，
用于排版后的 LLM 全文审查。

输出结构化的段落列表，每段包含：
- index: 段落序号
- type: heading / body / table_cell
- level: 标题层级（1-6，body 为 0）
- text: 段落文本内容
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.utils.logger import get_logger

logger = get_logger(__name__)

# 标题样式名称映射（支持中英文）
HEADING_STYLE_NAMES = {
    "Heading 1": 1,
    "heading 1": 1,
    "1 Heading": 1,
    "Heading 2": 2,
    "heading 2": 2,
    "2 Heading": 2,
    "Heading 3": 3,
    "heading 3": 3,
    "3 Heading": 3,
    "Heading 4": 4,
    "heading 4": 4,
    "4 Heading": 4,
    "Heading 5": 5,
    "heading 5": 5,
    "5 Heading": 5,
    "Heading 6": 6,
    "heading 6": 6,
    "6 Heading": 6,
    "标题 1": 1,
    "标题 2": 2,
    "标题 3": 3,
    "标题 4": 4,
    "标题 5": 5,
    "标题 6": 6,
}


@dataclass
class DocxParagraph:
    """DOCX 段落结构"""

    index: int
    type: str  # heading / body / table_cell
    level: int = 0  # 标题层级，body 为 0
    text: str = ""
    style_name: str = ""  # 原始样式名


@dataclass
class DocxText:
    """DOCX 全文提取结果"""

    paragraphs: list[DocxParagraph] = field(default_factory=list)
    total_chars: int = 0
    heading_count: int = 0
    table_count: int = 0

    def get_full_text(self) -> str:
        """获取全文纯文本（段落间用换行分隔）"""
        return "\n".join(p.text for p in self.paragraphs if p.text.strip())

    def get_headings(self) -> list[DocxParagraph]:
        """获取所有标题段落"""
        return [p for p in self.paragraphs if p.type == "heading"]

    def get_body_paragraphs(self) -> list[DocxParagraph]:
        """获取所有正文段落"""
        return [p for p in self.paragraphs if p.type == "body"]


class DocxTextExtractor:
    """DOCX 文本提取器

    从 DOCX 文档中提取结构化文本，识别段落类型和标题层级。
    """

    def __init__(self) -> None:
        self._heading_style_map: dict[str, int] = {}

    def extract(self, docx_path: str | Path) -> DocxText:
        """从 DOCX 文件中提取结构化文本

        Args:
            docx_path: DOCX 文件路径

        Returns:
            DocxText 结构化文本
        """
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")

        logger.info("提取 DOCX 文本: %s", path.name)
        doc = Document(str(path))

        paragraphs: list[DocxParagraph] = []
        idx = 0
        heading_count = 0
        table_count = 0
        total_chars = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # 段落
                para = self._extract_paragraph(element, doc, idx)
                if para:
                    paragraphs.append(para)
                    idx += 1
                    total_chars += len(para.text)
                    if para.type == "heading":
                        heading_count += 1
            elif tag == "tbl":
                # 表格
                para = self._extract_table(element, doc, idx)
                if para:
                    paragraphs.append(para)
                    idx += 1
                    total_chars += len(para.text)
                    table_count += 1

        result = DocxText(
            paragraphs=paragraphs,
            total_chars=total_chars,
            heading_count=heading_count,
            table_count=table_count,
        )
        logger.info(
            "DOCX 文本提取完成: %d 段落, %d 标题, %d 表格, %d 字符",
            len(paragraphs),
            heading_count,
            table_count,
            total_chars,
        )
        return result

    def _extract_paragraph(self, element, doc: Document, idx: int) -> DocxParagraph | None:
        """提取单个段落的文本和类型"""
        # 获取段落文本
        texts: list[str] = []
        for run in element.findall(qn("w:r")):
            t_elements = run.findall(qn("w:t"))
            for t in t_elements:
                if t.text:
                    texts.append(t.text)

        text = "".join(texts).strip()
        if not text:
            return None

        # 获取段落样式
        p_pr = element.find(qn("w:pPr"))
        style_name = ""
        if p_pr is not None:
            p_style = p_pr.find(qn("w:pStyle"))
            if p_style is not None:
                style_name = p_style.get(qn("w:val"), "")

        # 判断段落类型
        level = self._get_heading_level(style_name, text)
        if level > 0:
            return DocxParagraph(
                index=idx,
                type="heading",
                level=level,
                text=text,
                style_name=style_name,
            )

        return DocxParagraph(
            index=idx,
            type="body",
            level=0,
            text=text,
            style_name=style_name,
        )

    def _extract_table(self, element, doc: Document, idx: int) -> DocxParagraph | None:
        """提取表格内容为纯文本"""
        rows: list[str] = []
        for tr in element.findall(qn("w:tr")):
            cells: list[str] = []
            for tc in tr.findall(qn("w:tc")):
                cell_texts: list[str] = []
                for p in tc.findall(qn("w:p")):
                    for r in p.findall(qn("w:r")):
                        for t in r.findall(qn("w:t")):
                            if t.text:
                                cell_texts.append(t.text)
                cells.append("".join(cell_texts).strip())
            rows.append(" | ".join(c for c in cells if c))

        text = "\n".join(rows)
        if not text.strip():
            return None

        return DocxParagraph(
            index=idx,
            type="table_cell",
            level=0,
            text=f"[表格]\n{text}",
        )

    def _get_heading_level(self, style_name: str, text: str) -> int:
        """判断段落是否为标题及其层级

        优先级：
        1. 样式名匹配（Heading 1-6 / 标题 1-6）
        2. 编号规则匹配（如 "5.1"、"5.1.1" 开头）
        3. 附录匹配（如 "附录 A"）
        """
        # 样式名匹配
        if style_name in HEADING_STYLE_NAMES:
            return HEADING_STYLE_NAMES[style_name]

        # 缓存样式名（某些 DOCX 使用自定义样式名）
        if style_name and style_name not in self._heading_style_map:
            # 检查是否为已知标题变体
            for known, level in HEADING_STYLE_NAMES.items():
                if known.lower() in style_name.lower():
                    self._heading_style_map[style_name] = level
                    return level

        # 编号规则匹配
        import re

        # 主编号: "5.1"、"5.1.1"、"5.1.1.1"
        if re.match(r"^\d+(\.\d+){0,3}\s{2}", text):
            dots = text.split(".")[:4]
            return min(len(dots), 6)

        # 附录: "附录 A"、"附录B"
        if re.match(r"^附录\s*[A-Z]", text):
            return 1

        return 0
