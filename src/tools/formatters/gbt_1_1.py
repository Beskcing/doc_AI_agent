# ruff: noqa: N802, N803, N806  # XML 元素命名 rPr/pPr/tblPr 遵循 python-docx/OOXML 约定
"""GB/T 1.1 DOCX 格式修正器

基于用户单位 format_gbt.py 硬编码的 GB/T 1.1 规范，封装为可直接调用的格式化工具。
替代 LLM+RAG 生成的 style_config 驱动路径，确保排版结果精确可靠。

注册 ID: gbt_1.1
显示名称: GB/T 1.1 标准化工作导则

处理流程：
1. 页面设置（A4 + 标准边距）
2. 文档默认字体（宋体+TNR, 10.5pt）+ Normal 样式
3. 内容规整（日期合并 / TOC删除 / 拆分标题合并 / 双空格）
4. 段落分类（位置感知：封面/前言/正文/附录/方法标题等）
5. 逐段格式应用 + 表格居中 + 图片居中
6. 保存输出
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from src.models.document_schema import StyleReport
from src.tools.formatters.base import BaseDocxFormatter
from src.tools.formatters.registry import register_formatter
from src.utils.file_utils import ensure_dir
from src.utils.logger import get_logger

logger = get_logger(__name__)

# 对齐映射
_ALIGN_TO_XML = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "dist",
}


@register_formatter
class GbtDocxFormatter(BaseDocxFormatter):
    """GB/T 1.1 DOCX 格式修正器

    将 MinerU 输出的原始 DOCX 按 GB/T 1.1 标准进行格式修正。
    所有格式参数硬编码，不依赖 LLM 或 RAG，确保结果稳定可靠。
    """

    standard_id: str = "gbt_1.1"
    display_name: str = "GB/T 1.1 标准化工作导则"

    def __init__(self) -> None:
        self._warnings: list[str] = []
        self._stats: dict[str, int] = {}

    # =========================================================================
    # 公开入口
    # =========================================================================

    def process(self, input_path: str, output_path: str) -> StyleReport:
        """主入口：对 DOCX 文件执行 GB/T 1.1 格式修正

        Args:
            input_path: 输入 DOCX 路径
            output_path: 输出 DOCX 路径

        Returns:
            StyleReport 格式处理报告
        """
        input_p = Path(input_path)
        output_p = Path(output_path)
        self._warnings = []

        if not input_p.exists():
            return StyleReport(
                success=False,
                warnings=[f"输入文件不存在: {input_path}"],
                output_path=str(output_p),
            )

        ensure_dir(output_p.parent)

        try:
            doc = Document(str(input_p))
        except Exception as e:
            return StyleReport(
                success=False,
                warnings=[f"无法打开 DOCX 文件: {e}"],
                output_path=str(output_p),
            )

        # ── 1. 页面设置 ──
        self._apply_page_setup(doc)

        # ── 2. 文档默认值 ──
        self._apply_doc_defaults(doc)
        self._apply_normal_style(doc)

        # ── 3. 内容规整 ──
        self._merge_date_paragraphs(doc)
        toc_removed = self._remove_toc(doc)
        if toc_removed:
            logger.debug("已删除 TOC: %d 段落", toc_removed)
        self._merge_split_headings(doc)

        # ── 4. 段落分类与格式化 ──
        cover_start = self._find_cover_start(doc.paragraphs)
        if cover_start > 0:
            logger.debug("封面起始位置: P[%d]", cover_start)

        stats: dict[str, int] = {
            "pre_cover_title": 0,
            "pre_cover": 0,
            "cover": 0,
            "preface_title": 0,
            "preface_body": 0,
            "body_title": 0,
            "method_heading": 0,
            "appendix_heading": 0,
            "heading": 0,
            "body": 0,
            "formula": 0,
            "table_caption": 0,
            "figure_caption": 0,
            "empty": 0,
            "total": 0,
        }

        for idx, para in enumerate(doc.paragraphs):
            text = para.text
            cls = self._classify_paragraph(text, idx, doc.paragraphs, cover_start)
            stats[cls] = stats.get(cls, 0) + 1
            stats["total"] += 1

            # ── 应用格式 ──
            if cls == "empty":
                self._set_para_format(para, space_before_pt=0, space_after_pt=0)
            elif cls == "cover":
                self._apply_cover_format(para, text, idx - cover_start)
            elif cls == "pre_cover_title":
                self._apply_pre_cover_title_format(para)
            elif cls == "pre_cover":
                self._apply_pre_cover_format(para)
            elif cls == "preface_title":
                self._apply_preface_title_format(para)
            elif cls == "preface_body":
                self._apply_body_format(para)
            elif cls == "body_title":
                self._apply_body_title_format(para)
            elif cls == "method_heading":
                self._apply_method_heading_format(para)
            elif cls == "appendix_heading":
                self._apply_appendix_heading_format(para)
            elif cls == "heading":
                self._fix_heading_spacing(para)
                self._apply_heading_format(para)
            elif cls == "formula":
                self._apply_formula_format(para)
            elif cls == "table_caption":
                self._apply_table_caption_format(para)
            elif cls == "figure_caption":
                self._apply_figure_caption_format(para)
            elif cls == "body":
                self._apply_body_format(para)

        self._stats = stats

        # ── 5. 表格格式 ──
        tables_count = self._format_tables(doc)

        # ── 6. 图片居中 ──
        images_count = self._format_images(doc)

        # ── 7. 保存 ──
        try:
            doc.save(str(output_p))
            logger.info(
                "GB/T 格式化完成: %s → %s (段落=%d, 表格=%d, 图片=%d)",
                input_p.name,
                output_p.name,
                stats["total"],
                tables_count,
                images_count,
            )
        except Exception as e:
            return StyleReport(
                success=False,
                paragraphs_styled=stats["total"],
                tables_styled=tables_count,
                warnings=self._warnings + [f"保存失败: {e}"],
                output_path=str(output_p),
            )

        return StyleReport(
            success=True,
            paragraphs_styled=stats["total"],
            tables_styled=tables_count,
            headings_styled=stats.get("heading", 0) + stats.get("method_heading", 0) + stats.get("appendix_heading", 0),
            warnings=self._warnings,
            output_path=str(output_p),
        )

    # =========================================================================
    # XML 工具
    # =========================================================================

    @staticmethod
    def _get_or_add_rPr(run):
        """获取或创建 run 的 rPr 元素"""
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run._element.insert(0, rPr)
        return rPr

    @staticmethod
    def _get_or_add_pPr(paragraph):
        """获取或创建 paragraph 的 pPr 元素"""
        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            paragraph._element.insert(0, pPr)
        return pPr

    @staticmethod
    def _set_child(el, tag, val, attrs=None):
        """设置或更新指定 tag 的子元素"""
        child = el.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            el.append(child)
        if val is not None:
            if attrs:
                for k, v in attrs.items():
                    child.set(qn(k), str(v))
            else:
                child.set(qn("w:val"), str(val))
        return child

    @staticmethod
    def _remove_child(el, tag):
        """移除子元素（如果存在）"""
        child = el.find(qn(tag))
        if child is not None:
            el.remove(child)

    @staticmethod
    def _set_run_font(rPr, east_asia, latin, size_pt, bold=None, bold_cs=None):
        """设置 run 的字体属性"""
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        if east_asia:
            rFonts.set(qn("w:eastAsia"), east_asia)
        if latin:
            rFonts.set(qn("w:ascii"), latin)
            rFonts.set(qn("w:hAnsi"), latin)
            rFonts.set(qn("w:cs"), latin)
        # 移除主题引用
        for attr in ("w:asciiTheme", "w:eastAsiaTheme", "w:hAnsiTheme", "w:cstheme"):
            try:
                del rFonts.attrib[qn(attr)]
            except KeyError:
                pass
        if size_pt is not None:
            half_pt = int(round(size_pt * 2))
            GbtDocxFormatter._set_child(rPr, "w:sz", half_pt)
            GbtDocxFormatter._set_child(rPr, "w:szCs", half_pt)
        if bold is not None:
            if bold:
                GbtDocxFormatter._set_child(rPr, "w:b", None, {"w:val": "1"})
            else:
                GbtDocxFormatter._remove_child(rPr, "w:b")
        if bold_cs is not None:
            if bold_cs:
                GbtDocxFormatter._set_child(rPr, "w:bCs", None, {"w:val": "1"})
            else:
                GbtDocxFormatter._remove_child(rPr, "w:bCs")

    @staticmethod
    def _format_run(run, east_asia, latin, size_pt, bold=None, bold_cs=None):
        """格式化单个 run"""
        rPr = GbtDocxFormatter._get_or_add_rPr(run)
        GbtDocxFormatter._set_run_font(rPr, east_asia, latin, size_pt, bold, bold_cs)

    # =========================================================================
    # 页面设置
    # =========================================================================

    def _apply_page_setup(self, doc: Document) -> None:
        """A4 纸张 + GB/T 1.1 标准边距"""
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.orientation = WD_ORIENT.PORTRAIT
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.1)
        logger.debug("页面设置: A4, 边距 25/25/25/21mm")

    # =========================================================================
    # 文档默认值
    # =========================================================================

    def _apply_normal_style(self, doc: Document) -> None:
        """设置 Normal 样式：宋体+TNR, 10.5pt"""
        style = doc.styles["Normal"]
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.element.append(rPr)
        self._set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None, bold_cs=None)
        b = rPr.find(qn("w:b"))
        bCs = rPr.find(qn("w:bCs"))
        if b is not None:
            rPr.remove(b)
        if bCs is not None:
            rPr.remove(bCs)

        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            style.element.append(pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is not None:
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:lineRule"), "auto")
        else:
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:lineRule"), "auto")
            pPr.append(spacing)
        logger.debug("Normal 样式: 宋体+TNR, 10.5pt")

    def _apply_doc_defaults(self, doc: Document) -> None:
        """设置文档默认字体：宋体+TNR, 10.5pt"""
        styles_elem = doc.styles.element
        doc_defaults = styles_elem.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_elem.insert(0, doc_defaults)

        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is None:
            rpr_default = OxmlElement("w:rPrDefault")
            doc_defaults.insert(
                0 if doc_defaults[0].tag != qn("w:pPrDefault") else 1,
                rpr_default,
            )
        rPr = rpr_default.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rpr_default.append(rPr)
        self._set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
        logger.debug("docDefaults: 宋体+TNR, 10.5pt")

    # =========================================================================
    # 段落格式设置
    # =========================================================================

    def _set_para_format(
        self,
        paragraph,
        alignment=None,
        first_line_indent_pt=None,
        left_indent_pt=None,
        space_before_pt=None,
        space_after_pt=None,
        line_spacing=None,
    ):
        """设置段落格式属性"""
        pPr = self._get_or_add_pPr(paragraph)

        if alignment is not None:
            self._set_child(pPr, "w:jc", _ALIGN_TO_XML.get(alignment, "left"))

        # 缩进
        ind = pPr.find(qn("w:ind"))
        if first_line_indent_pt is not None or left_indent_pt is not None:
            if ind is None:
                ind = OxmlElement("w:ind")
                spacing = pPr.find(qn("w:spacing"))
                if spacing is not None:
                    pPr.insert(list(pPr).index(spacing), ind)
                else:
                    pPr.append(ind)
            if first_line_indent_pt is not None:
                ind.set(qn("w:firstLine"), str(int(first_line_indent_pt * 20)))
            if left_indent_pt is not None:
                ind.set(qn("w:left"), str(int(left_indent_pt * 20)))
        elif ind is not None:
            for attr in ("w:firstLine", "w:left", "w:right", "w:hanging"):
                try:
                    del ind.attrib[qn(attr)]
                except KeyError:
                    pass

        # 间距
        spacing = pPr.find(qn("w:spacing"))
        if space_before_pt is not None or space_after_pt is not None or line_spacing is not None:
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            if space_before_pt is not None:
                spacing.set(qn("w:before"), str(int(space_before_pt * 20)))
            if space_after_pt is not None:
                spacing.set(qn("w:after"), str(int(space_after_pt * 20)))
            if line_spacing is not None:
                spacing.set(qn("w:line"), str(int(line_spacing * 240)))
                spacing.set(qn("w:lineRule"), "auto")

    # =========================================================================
    # 内容规整
    # =========================================================================

    def _merge_date_paragraphs(self, doc: Document) -> None:
        """合并相邻的"发布"/"实施"日期段落为一行"""
        paragraphs = doc.paragraphs
        if len(paragraphs) < 5:
            return
        for i in range(len(paragraphs) - 1):
            p1_text = paragraphs[i].text.strip()
            p2_text = paragraphs[i + 1].text.strip()
            if p1_text.endswith("发布") and p2_text.endswith("实施"):
                date_pub = p1_text.replace("发布", "").strip()
                date_impl = p2_text.replace("实施", "").strip()
                merged = f"{date_pub}发布 {date_impl}实施"
                p1 = paragraphs[i]
                for run in p1.runs:
                    run.text = ""
                if p1.runs:
                    p1.runs[0].text = merged
                else:
                    p1.add_run(merged)
                p2_elem = paragraphs[i + 1]._element
                p2_elem.getparent().remove(p2_elem)
                logger.debug("合并日期: '%s'", merged)
                return

    def _remove_toc(self, doc: Document) -> int:
        """检测并删除 TOC（目次/目录）段落，返回删除段落数"""
        paragraphs = doc.paragraphs
        toc_start = -1
        for i, p in enumerate(paragraphs):
            t = p.text.strip()
            if toc_start < 0:
                if t in ("目  次", "目次", "目  录", "目录"):
                    toc_start = i
            elif t.startswith("前") and "言" in t:
                toc_end = i
                body = doc.element.body
                for j in range(toc_start, toc_end):
                    p_elem = paragraphs[j]._element
                    body.remove(p_elem)
                return toc_end - toc_start
        return 0

    def _merge_split_headings(self, doc: Document) -> None:
        """合并拆分标题（附录标题 + 编号标题）"""
        paragraphs = doc.paragraphs
        if len(paragraphs) < 2:
            return
        merged_count = 0
        i = 0
        while i < len(paragraphs) - 1:
            p1_text = paragraphs[i].text.strip()
            p2_text = paragraphs[i + 1].text.strip()
            should_merge = False

            if re.match(r"^附录\s*[A-Z]$", p1_text) and p2_text:
                if (
                    not p2_text.startswith("附录")
                    and not p2_text.startswith("前言")
                    and not p2_text.startswith("参考文献")
                    and not re.match(r"^\d+\s+", p2_text)
                ):
                    should_merge = True
            elif re.match(r"^(\d+(\.\d+)*|[A-Z]\.\d+(\.\d+)*)$", p1_text) and p2_text:
                if (
                    not re.match(r"^(\d+(\.\d+)*|[A-Z]\.\d+)", p2_text)
                    and not p2_text.startswith("附录")
                    and not p2_text.startswith("前言")
                    and not p2_text.startswith("参考文献")
                ):
                    should_merge = True

            if should_merge:
                merged = f"{p1_text}  {p2_text}"
                p1 = paragraphs[i]
                for run in p1.runs:
                    run.text = ""
                if p1.runs:
                    p1.runs[0].text = merged
                else:
                    p1.add_run(merged)
                p2_elem = paragraphs[i + 1]._element
                p2_elem.getparent().remove(p2_elem)
                merged_count += 1
                i += 1
                continue
            i += 1
        if merged_count > 0:
            logger.debug("合并拆分标题: %d 处", merged_count)

    def _fix_heading_spacing(self, paragraph) -> None:
        """确保编号标题后有双空格"""
        text = paragraph.text.strip()
        match = re.match(r"^(\d+(?:\.\d+)*)(\s+)(.*)", text)
        if not match:
            return
        num = match.group(1)
        rest = match.group(3)
        current_spaces = match.group(2)
        if current_spaces == "  ":
            return
        new_text = num + "  " + rest
        runs = paragraph.runs
        if runs:
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ""

    # =========================================================================
    # 段落分类
    # =========================================================================

    @staticmethod
    def _is_numbered_heading(text: str) -> bool:
        """检查是否为编号标题"""
        if re.match(r"^\d+(\.\d+)*\s+\S", text):
            return True
        if re.match(r"^[A-Z]\.\d+(\.\d+)*\s+\S", text):
            return True
        return False

    @staticmethod
    def _is_method_heading(text: str) -> bool:
        """检查是否为方法标题（如 '第一法 密度瓶法'）"""
        return bool(re.match(r"^第[一二三四五六七八九十百]+法\s", text))

    @staticmethod
    def _is_appendix_heading(text: str) -> bool:
        """检查是否为附录标题"""
        return bool(re.match(r"^附录\s*[A-Z]", text))

    @staticmethod
    def _find_cover_start(paragraphs) -> int:
        """查找封面起始位置（'中华人民共和国国家标准'）"""
        for i, p in enumerate(paragraphs[:30]):
            t = p.text.strip()
            if "中华人民共和国国家标准" in t:
                return i
        return 0

    @staticmethod
    def _find_preface_index(paragraphs, cover_start: int) -> int:
        """动态查找前言索引"""
        for i in range(cover_start + 3, min(cover_start + 12, len(paragraphs))):
            t = paragraphs[i].text.strip()
            if t in ("前言", "前 言"):
                return i
        return -1

    @staticmethod
    def _is_standard_name_line(text: str, index: int, paragraphs) -> bool:
        """检查是否为标准名称行"""
        for j in range(index + 1, min(index + 5, len(paragraphs))):
            next_text = paragraphs[j].text.strip()
            if next_text:
                if re.match(r"^\d+\s{2,}[\u4e00-\u9fff]", next_text):
                    return True
                break
        return False

    @staticmethod
    def _has_reached_chapter_1(index: int, paragraphs) -> bool:
        """检查是否已到达第1章"""
        for i in range(index - 1, max(0, index - 50), -1):
            t = paragraphs[i].text.strip()
            if not t:
                continue
            if re.match(r"^1\s+\S", t):
                return True
        return False

    def _classify_paragraph(self, text: str, index: int, paragraphs, cover_start: int = 0) -> str:
        """位置感知段落分类

        Returns:
            pre_cover_title / pre_cover / cover / preface_title / preface_body /
            body_title / method_heading / appendix_heading / heading / body /
            formula / table_caption / figure_caption / empty
        """
        text = text.strip()
        if not text:
            return "empty"

        # 封面之前的内容（修订通告等）
        if index < cover_start:
            if re.match(r"^GB[/T ]*\s*\d", text) or re.match(r"^[一二三四五六七八九十]+[、．.]", text):
                return "pre_cover_title"
            return "pre_cover"

        # 封面区域
        preface_idx = self._find_preface_index(paragraphs, cover_start)
        if preface_idx >= 0 and index < preface_idx:
            return "cover"
        if preface_idx < 0 and index < cover_start + 5:
            return "cover"

        # 前言标题
        if index == preface_idx:
            return "preface_title"

        # 前言正文
        if preface_idx >= 0 and index > preface_idx:
            if len(text) < 30 and not text.endswith("。") and not text.endswith("；"):
                if (
                    not self._is_numbered_heading(text)
                    and not self._is_method_heading(text)
                    and not self._is_appendix_heading(text)
                ):
                    if re.match(r"^1\s+\S", text) or self._is_standard_name_line(text, index, paragraphs):
                        return "body_title"
            if (
                not self._is_numbered_heading(text)
                and not self._is_method_heading(text)
                and not self._is_appendix_heading(text)
                and not text.startswith("1  ")
                and not re.match(r"^表\s+[A-Z]?\.?\d+", text)
                and not re.match(r"^图\s*\d+", text)
            ):
                if not self._has_reached_chapter_1(index, paragraphs):
                    return "preface_body"

        # 方法标题
        if self._is_method_heading(text):
            return "method_heading"

        # 附录标题
        if self._is_appendix_heading(text):
            return "appendix_heading"

        # 表格标题（必须在编号标题之前检测）
        if re.match(r"^表\s+[A-Z]?\.?\d+", text):
            return "table_caption"

        # 图表标题
        if re.match(r"^图\s*\d+", text):
            return "figure_caption"

        # 编号标题
        if self._is_numbered_heading(text):
            return "heading"

        # 公式
        if text.startswith("$$") or re.match(r"^[A-Z] = ", text):
            return "formula"

        return "body"

    # =========================================================================
    # 各角色格式应用
    # =========================================================================

    def _apply_cover_format(self, paragraph, text: str, para_idx: int) -> None:
        """封面区域：内容感知字体选择"""
        text_stripped = text.strip()
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            if "国家标准" in text_stripped:
                self._set_run_font(rPr, east_asia="宋体", latin="宋体", size_pt=16.0, bold=True, bold_cs=True)
            elif re.match(r"^GB[/T\s]*\d", text_stripped):
                self._set_run_font(
                    rPr,
                    east_asia="宋体",
                    latin="Times New Roman",
                    size_pt=16.0,
                    bold=True,
                    bold_cs=True,
                )
            elif re.search(r"\d{4}[-/]\d{2}[-/]\d{2}", text_stripped):
                self._set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
            elif "发布" in text_stripped or "实施" in text_stripped:
                self._set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
            elif re.match(r"^(代替|替代)", text_stripped):
                self._set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
            elif re.match(r"^[A-Z][a-z]+", text_stripped) and len(text_stripped) > 20:
                self._set_run_font(
                    rPr,
                    east_asia="宋体",
                    latin="Times New Roman",
                    size_pt=16.0,
                    bold=True,
                    bold_cs=True,
                )
            else:
                self._set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_pre_cover_title_format(self, paragraph) -> None:
        """修订通告标题：16pt 宋体 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(
                rPr,
                east_asia="宋体",
                latin="Times New Roman",
                size_pt=16.0,
                bold=True,
                bold_cs=True,
            )
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_pre_cover_format(self, paragraph) -> None:
        """修订通告正文：10.5pt 宋体+TNR, JUSTIFY, 21pt 缩进"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=21.0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_preface_title_format(self, paragraph) -> None:
        """前言标题：16pt 黑体 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_body_title_format(self, paragraph) -> None:
        """标准名称标题：16pt 黑体 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(rPr, east_asia="黑体", latin="黑体", size_pt=16.0, bold=True, bold_cs=True)
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_heading_format(self, paragraph) -> None:
        """编号标题：10.5pt 宋体+TNR 不加粗, JUSTIFY, 无缩进"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(
                rPr,
                east_asia="宋体",
                latin="Times New Roman",
                size_pt=10.5,
                bold=False,
                bold_cs=False,
            )
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_method_heading_format(self, paragraph) -> None:
        """方法标题：10.5pt 宋体+TNR 不加粗 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(
                rPr,
                east_asia="宋体",
                latin="Times New Roman",
                size_pt=10.5,
                bold=False,
                bold_cs=False,
            )
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_appendix_heading_format(self, paragraph) -> None:
        """附录标题：10.5pt 宋体+TNR 不加粗, JUSTIFY"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(
                rPr,
                east_asia="宋体",
                latin="Times New Roman",
                size_pt=10.5,
                bold=False,
                bold_cs=False,
            )
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_body_format(self, paragraph) -> None:
        """正文：10.5pt 宋体+TNR, JUSTIFY, 21pt 首行缩进"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=21.0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_table_caption_format(self, paragraph) -> None:
        """表格标题：10.5pt 宋体+TNR 不加粗 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(
                rPr,
                east_asia="宋体",
                latin="Times New Roman",
                size_pt=10.5,
                bold=False,
                bold_cs=False,
            )
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_figure_caption_format(self, paragraph) -> None:
        """图表标题：10.5pt 宋体+TNR 不加粗 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(
                rPr,
                east_asia="宋体",
                latin="Times New Roman",
                size_pt=10.5,
                bold=False,
                bold_cs=False,
            )
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    def _apply_formula_format(self, paragraph) -> None:
        """公式：10.5pt 宋体+TNR 居中"""
        for run in paragraph.runs:
            rPr = self._get_or_add_rPr(run)
            self._set_run_font(rPr, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=None)
        self._set_para_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            space_before_pt=0,
            space_after_pt=0,
            line_spacing=1.0,
        )

    # =========================================================================
    # 表格与图片
    # =========================================================================

    def _format_tables(self, doc: Document) -> int:
        """表格格式：居中 + Table Grid 样式 + 表头宋体 10.5pt 不加粗"""
        count = 0
        for table in doc.tables:
            try:
                table.style = doc.styles["Table Grid"]
            except KeyError:
                pass

            tbl = table._element
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)
            jc = tblPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                tblPr.append(jc)
            jc.set(qn("w:val"), "center")

            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            rPr = self._get_or_add_rPr(run)
                            self._set_run_font(
                                rPr,
                                east_asia="宋体",
                                latin="Times New Roman",
                                size_pt=10.5,
                                bold=False,
                                bold_cs=False,
                            )
            count += 1
        if count > 0:
            logger.debug("格式化表格: %d 个", count)
        return count

    def _format_images(self, doc: Document) -> int:
        """居中所有图片"""
        count = 0
        for paragraph in doc.paragraphs:
            if paragraph._element.xpath(".//pic:pic"):
                pPr = paragraph._element.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    paragraph._element.insert(0, pPr)
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    pPr.append(jc)
                jc.set(qn("w:val"), "center")
                count += 1
        if count > 0:
            logger.debug("居中图片: %d 个", count)
        return count
