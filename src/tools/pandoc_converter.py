"""Pandoc 格式转换工具

使用 Pandoc 进行 Markdown -> DOCX / HTML 格式转换。
关键特性：
- 使用 --from=markdown+raw_html 支持 HTML 表格
- 转换前后配合 HTMLTablePreserver 保护/恢复 HTML 表格
- 提供转换结果校验
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from src.models.document_schema import ConversionReport
from src.tools.html_table_preserver import HTMLTablePreserver
from src.utils.file_utils import ensure_dir, get_file_size_mb
from src.utils.logger import get_logger

logger = get_logger(__name__)


class PandocConverter:
    """Pandoc 格式转换器"""

    def __init__(
        self,
        pandoc_path: str = "pandoc",
        extra_args: list[str] | None = None,
    ):
        """初始化转换器

        Args:
            pandoc_path: Pandoc 可执行文件路径
            extra_args: 额外的 Pandoc 命令行参数
        """
        self.pandoc_path = pandoc_path
        self.extra_args = extra_args or []
        self._table_preserver = HTMLTablePreserver()

    def check_installation(self) -> str | None:
        """检查 Pandoc 是否可用

        Returns:
            Pandoc 版本号字符串，不可用返回 None
        """
        try:
            result = subprocess.run(
                [self.pandoc_path, "--version"],
                capture_output=True,
                text=True,
                timeout=10,
            )
            if result.returncode == 0:
                version_line = result.stdout.splitlines()[0] if result.stdout else ""
                logger.info("Pandoc 可用: %s", version_line)
                return version_line
        except (FileNotFoundError, subprocess.TimeoutExpired):
            logger.warning("Pandoc 不可用: %s", self.pandoc_path)
        return None

    def markdown_to_docx(
        self,
        markdown: str,
        output_path: str | Path,
        reference_docx: str | Path | None = None,
    ) -> ConversionReport:
        """Markdown -> DOCX 转换

        流程：
        1. 预处理：保护 HTML 表格
        2. 调用 Pandoc 转换
        3. 后处理：恢复 HTML 表格（如果需要）

        Args:
            markdown: Markdown 文本
            output_path: 输出 DOCX 文件路径
            reference_docx: 参考 Word 模板（可选）

        Returns:
            ConversionReport 转换结果报告
        """
        output_path = Path(output_path)
        ensure_dir(output_path.parent)

        # 统计原始表格和公式数量
        table_count = self._table_preserver.count_tables(markdown)
        formula_count = len(self._count_formulas(markdown))

        # 保护 HTML 表格（Pandoc 能直接处理 raw_html，不需要占位符）
        # 但我们需要确保 HTML 表格结构完整
        warnings: list[str] = []
        errors: list[str] = []

        # 校验 HTML 表格完整性
        for table_html in self._table_preserver.find_unprotected_tables(markdown):
            is_valid, issues = self._table_preserver.validate_table_integrity(table_html)
            if not is_valid:
                warnings.append(f"表格结构不完整: {'; '.join(issues)}")

        # 构建 Pandoc 参数
        args = ["--from=markdown+raw_html+tex_math_dollars", "--to=docx"]
        if reference_docx:
            args.extend(["--reference-doc", str(reference_docx)])
        args.extend(self.extra_args)

        try:
            self._run_pandoc(markdown, args, str(output_path))
            success = output_path.exists() and get_file_size_mb(output_path) > 0

            if success:
                logger.info("DOCX 转换成功: %s (%.2f MB)", output_path, get_file_size_mb(output_path))
            else:
                errors.append("输出文件为空或不存在")

            return ConversionReport(
                success=success,
                output_path=str(output_path),
                tables_converted=table_count,
                formulas_converted=formula_count,
                warnings=warnings,
                errors=errors,
            )

        except Exception as e:
            error_msg = f"Pandoc 转换失败: {e}"
            logger.error(error_msg)
            errors.append(error_msg)
            return ConversionReport(
                success=False,
                output_path=str(output_path),
                warnings=warnings,
                errors=errors,
            )

    def markdown_to_html(self, markdown: str) -> str:
        """Markdown -> HTML 转换（用于调试和预览）

        Args:
            markdown: Markdown 文本

        Returns:
            HTML 字符串

        Raises:
            RuntimeError: 转换失败
        """
        args = ["--from=markdown+raw_html+tex_math_dollars", "--to=html5", "--standalone"]
        return self._run_pandoc(markdown, args)

    def _run_pandoc(
        self,
        input_text: str,
        args: list[str],
        output_path: str | None = None,
    ) -> str:
        """底层 Pandoc 调用

        Args:
            input_text: 输入文本
            args: Pandoc 命令行参数
            output_path: 输出文件路径（可选，不指定则输出到 stdout）

        Returns:
            Pandoc stdout 输出

        Raises:
            RuntimeError: Pandoc 执行失败
        """
        cmd = [self.pandoc_path] + args

        if output_path:
            cmd.extend(["-o", output_path])

        logger.debug("执行 Pandoc: %s", " ".join(cmd))

        try:
            result = subprocess.run(
                cmd,
                input=input_text,
                capture_output=True,
                text=True,
                timeout=120,
                encoding="utf-8",
            )

            if result.returncode != 0:
                stderr = result.stderr.strip()
                raise RuntimeError(f"Pandoc 返回非零状态码 {result.returncode}: {stderr}")

            if result.stderr:
                logger.warning("Pandoc stderr: %s", result.stderr[:500])

            return result.stdout

        except subprocess.TimeoutExpired:
            raise RuntimeError("Pandoc 执行超时（120秒）") from None
        except FileNotFoundError:
            raise RuntimeError(f"Pandoc 可执行文件未找到: {self.pandoc_path}") from None

    def validate_output(self, output_path: str | Path) -> ConversionReport:
        """校验转换结果

        检查项:
        - 文件是否存在
        - 文件大小是否合理
        - 文件是否可被 python-docx 打开（DOCX 格式）

        Args:
            output_path: 输出文件路径

        Returns:
            ConversionReport
        """
        output_path = Path(output_path)
        warnings: list[str] = []
        errors: list[str] = []

        if not output_path.exists():
            return ConversionReport(
                success=False,
                output_path=str(output_path),
                errors=["输出文件不存在"],
            )

        size_mb = get_file_size_mb(output_path)
        if size_mb < 0.001:
            errors.append(f"文件过小 ({size_mb:.4f} MB)，可能为空文件")
        elif size_mb > 100:
            warnings.append(f"文件较大 ({size_mb:.1f} MB)，请检查是否包含大量图片")

        # 尝试用 python-docx 打开验证
        if output_path.suffix.lower() == ".docx":
            try:
                from docx import Document

                doc = Document(str(output_path))
                para_count = len(doc.paragraphs)
                table_count = len(doc.tables)
                logger.info("DOCX 校验通过: %d 段落, %d 表格", para_count, table_count)
            except Exception as e:
                errors.append(f"DOCX 文件无法打开: {e}")

        return ConversionReport(
            success=len(errors) == 0,
            output_path=str(output_path),
            warnings=warnings,
            errors=errors,
        )

    def _count_formulas(self, markdown: str) -> list[str]:
        """统计 Markdown 中的公式数量

        Args:
            markdown: Markdown 文本

        Returns:
            公式列表
        """
        formulas = []
        # 块公式 $$...$$
        block_pattern = re.compile(r"\$\$.*?\$\$", re.DOTALL)
        formulas.extend(block_pattern.findall(markdown))
        # 行内公式 $...$（排除 $$ ）
        inline_pattern = re.compile(r"(?<!\$)\$(?!\$)[^$]+?(?<!\$)\$(?!\$)")
        formulas.extend(inline_pattern.findall(markdown))
        return formulas
