"""仔细分析 GB/T 14454.13-2008 文档中的附录样式

逐段落提取附录区域的完整格式信息，包括：
- 附录标题样式（附录A/附录B）
- 附录内各级条款样式
- 附录内表格样式
- 附录内公式样式
- 附录内正文段落样式
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = Path("d:/doc_ai_agent/data/templates/GB_T_14454_13_2008.docx")


def emu_to_pt(emu):
    if emu is None:
        return None
    return round(int(emu) / 12700, 2)


def emu_to_cm(emu):
    if emu is None:
        return None
    return round(int(emu) / 360000, 2)


def get_alignment(paragraph):
    aln = paragraph.alignment
    if aln is not None:
        mapping = {0: "left", 1: "center", 2: "right", 3: "justify"}
        return mapping.get(int(aln), str(aln))
    # 从样式继承
    try:
        if paragraph.style and paragraph.style.paragraph_format:
            pf = paragraph.style.paragraph_format
            if pf.alignment is not None:
                mapping = {0: "left", 1: "center", 2: "right", 3: "justify"}
                return mapping.get(int(pf.alignment), str(pf.alignment))
    except:
        pass
    return "left(default)"


def get_font_info(run):
    info = {}
    # 字体名
    info["name"] = run.font.name or None
    # 字号
    info["size_pt"] = emu_to_pt(run.font.size) if run.font.size else None
    # 加粗
    info["bold"] = run.font.bold
    # 斜体
    info["italic"] = run.font.italic
    # 下划线
    info["underline"] = run.font.underline
    # 颜色
    if run.font.color and run.font.color.rgb:
        info["color"] = str(run.font.color.rgb)

    # 东亚字体
    try:
        r_element = run._element
        rpr = r_element.find(qn("w:rPr"))
        if rpr is not None:
            r_fonts = rpr.find(qn("w:rFonts"))
            if r_fonts is not None:
                info["east_asia"] = r_fonts.get(qn("w:eastAsia"))
                info["ascii"] = r_fonts.get(qn("w:ascii"))
                info["hAnsi"] = r_fonts.get(qn("w:hAnsi"))
    except:
        pass

    return {k: v for k, v in info.items() if v is not None}


def get_paragraph_format(paragraph):
    pf = paragraph.paragraph_format
    info = {}
    if pf:
        info["alignment"] = get_alignment(paragraph)
        info["line_spacing"] = pf.line_spacing
        info["space_before_pt"] = emu_to_pt(pf.space_before) if pf.space_before else 0
        info["space_after_pt"] = emu_to_pt(pf.space_after) if pf.space_after else 0
        info["first_line_indent_pt"] = emu_to_pt(pf.first_line_indent) if pf.first_line_indent else 0
        info["left_indent_cm"] = emu_to_cm(pf.left_indent) if pf.left_indent else 0
        info["right_indent_cm"] = emu_to_cm(pf.right_indent) if pf.right_indent else 0

        # keep_together / keep_with_next / widow_control
        try:
            pPr = paragraph._element.find(qn("w:pPr"))
            if pPr is not None:
                info["keep_together"] = pPr.find(qn("w:keepLines")) is not None
                info["keep_with_next"] = pPr.find(qn("w:keepNext")) is not None
                wc = pPr.find(qn("w:widowControl"))
                info["widow_control"] = wc is None or wc.get(qn("w:val")) != "0"
                # outline level
                ol = pPr.find(qn("w:outlineLvl"))
                if ol is not None:
                    info["outline_level"] = int(ol.get(qn("w:val")))
        except:
            pass

    return info


def classify_content(text):
    """识别段落内容类型"""
    text = text.strip()
    if not text:
        return "empty"
    # 附录标题
    if re.match(r'^附录\s*[A-Z]', text):
        return "appendix_title"
    # 附录内条款
    if re.match(r'^A\.\d+\.\d+\.\d+\.\d+\s', text):
        return "appendix_clause_5"
    if re.match(r'^A\.\d+\.\d+\.\d+\s', text):
        return "appendix_clause_4"
    if re.match(r'^A\.\d+\.\d+\s', text):
        return "appendix_clause_3"
    if re.match(r'^A\.\d+\s', text):
        return "appendix_clause_2"
    if re.match(r'^A\.\s', text) or text == "A.":
        return "appendix_clause_1"
    # 参考文献条目
    if re.match(r'^\[\d+\]', text):
        return "reference_item"
    # 公式
    if re.match(r'^[（(]\d+[)）]', text) or "式(" in text or "式（" in text:
        return "formula"
    # 表格相关
    if re.match(r'^表\s*A', text) or re.match(r'^表\s*\d', text):
        return "table_caption"
    # 注释
    if re.match(r'^注\s*\d*[:：]', text) or text.startswith("注：") or text.startswith("注:"):
        return "note"
    # 列表项
    if text.startswith("——") or text.startswith("—"):
        return "list_item"
    # 一级条款
    if re.match(r'^\d+\s+\S', text):
        return "clause_1"
    # 二级条款
    if re.match(r'^\d+\.\d+\s', text):
        return "clause_2"
    # 三级条款
    if re.match(r'^\d+\.\d+\.\d+\s', text):
        return "clause_3"
    return "body"


def main():
    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} not found")
        sys.exit(1)

    doc = Document(str(DOCX_PATH))
    paragraphs = doc.paragraphs

    # 找到附录开始的段落索引
    appendix_start = None
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if re.match(r'^附录\s*[A-Z]', text):
            appendix_start = i
            break

    if appendix_start is None:
        print("未找到附录段落")
        sys.exit(0)

    print(f"=== 附录分析: {DOCX_PATH.name} ===")
    print(f"附录起始段落索引: {appendix_start}")
    print(f"附录区域段落总数: {len(paragraphs) - appendix_start}")
    print()

    # 逐段落分析
    print("=" * 80)
    print("附录区域逐段落格式详情")
    print("=" * 80)

    current_section = ""
    for i in range(appendix_start, len(paragraphs)):
        p = paragraphs[i]
        text = p.text.strip()
        if not text:
            continue

        content_type = classify_content(text)
        style_name = p.style.name if p.style else "None"
        fmt = get_paragraph_format(p)

        # 收集 run 级字体信息
        run_fonts = []
        for run in p.runs:
            if run.text.strip():
                fi = get_font_info(run)
                if fi:
                    run_fonts.append({
                        "text": run.text[:20],
                        "font": fi,
                    })

        # 取第一个 run 的字体作为代表
        rep_font = run_fonts[0]["font"] if run_fonts else {}

        print(f"\n--- [{i}] {content_type} ---")
        print(f"  文本: {text[:80]}{'...' if len(text) > 80 else ''}")
        print(f"  样式名: {style_name}")
        print(f"  对齐: {fmt.get('alignment', '?')}")
        print(f"  行距: {fmt.get('line_spacing', '?')}")
        print(f"  段前: {fmt.get('space_before_pt', '?')}pt  段后: {fmt.get('space_after_pt', '?')}pt")
        print(f"  首行缩进: {fmt.get('first_line_indent_pt', 0)}pt  左缩进: {fmt.get('left_indent_cm', 0)}cm")
        if "outline_level" in fmt:
            print(f"  大纲级别: {fmt['outline_level']}")
        if "keep_with_next" in fmt:
            print(f"  keep_with_next: {fmt['keep_with_next']}  keep_together: {fmt.get('keep_together', '?')}")
        if rep_font:
            print(f"  字体: {rep_font.get('east_asia', rep_font.get('ascii', rep_font.get('name', '?')))} "
                  f"{rep_font.get('size_pt', '?')}pt bold={rep_font.get('bold', '?')}")
        if len(run_fonts) > 1:
            print(f"  Run数: {len(run_fonts)}")
            for rf in run_fonts[:3]:
                f = rf["font"]
                print(f"    '{rf['text']}' -> {f.get('east_asia', f.get('ascii', f.get('name', '?')))} "
                      f"{f.get('size_pt', '?')}pt bold={f.get('bold', '?')}")

    # 分析附录内表格
    print("\n" + "=" * 80)
    print("附录区域表格分析")
    print("=" * 80)

    # 找到附录后的表格
    for idx, table in enumerate(doc.tables):
        # 判断表格是否在附录区域（通过表格内文本判断）
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "
        if not any(kw in table_text for kw in ["附录", "A.", "A．"]):
            # 也可能表格标题在段落中
            continue

        print(f"\n--- 表格 {idx + 1} ---")
        print(f"  行数: {len(table.rows)}  列数: {len(table.columns)}")
        # 表格对齐
        try:
            tbl = table._element
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is not None:
                jc = tblPr.find(qn("w:jc"))
                if jc is not None:
                    print(f"  表格对齐: {jc.get(qn('w:val'))}")
        except:
            pass

        # 表头行
        if table.rows:
            header_row = table.rows[0]
            print(f"  表头行:")
            for ci, cell in enumerate(header_row.cells):
                for cp in cell.paragraphs:
                    if cp.text.strip():
                        fi = get_font_info(cp.runs[0]) if cp.runs and cp.runs[0].text.strip() else {}
                        print(f"    [{ci}] '{cp.text[:30]}' -> "
                              f"{fi.get('east_asia', fi.get('ascii', '?'))} "
                              f"{fi.get('size_pt', '?')}pt bold={fi.get('bold', '?')}")

        # 数据行（取第一行）
        if len(table.rows) > 1:
            data_row = table.rows[1]
            print(f"  数据行(第1行):")
            for ci, cell in enumerate(data_row.cells):
                for cp in cell.paragraphs:
                    if cp.text.strip():
                        fi = get_font_info(cp.runs[0]) if cp.runs and cp.runs[0].text.strip() else {}
                        print(f"    [{ci}] '{cp.text[:30]}' -> "
                              f"{fi.get('east_asia', fi.get('ascii', '?'))} "
                              f"{fi.get('size_pt', '?')}pt bold={fi.get('bold', '?')}")

    # 统计附录样式
    print("\n" + "=" * 80)
    print("附录样式统计")
    print("=" * 80)

    appendix_types = {}
    for i in range(appendix_start, len(paragraphs)):
        text = paragraphs[i].text.strip()
        if not text:
            continue
        ct = classify_content(text)
        if ct == "empty":
            continue
        if ct not in appendix_types:
            appendix_types[ct] = {"count": 0, "fonts": [], "alignments": [], "sizes": [], "bolds": []}
        appendix_types[ct]["count"] += 1
        p = paragraphs[i]
        appendix_types[ct]["alignments"].append(get_alignment(p))
        for run in p.runs:
            if run.text.strip():
                fi = get_font_info(run)
                if fi.get("size_pt"):
                    appendix_types[ct]["sizes"].append(fi["size_pt"])
                if fi.get("bold") is not None:
                    appendix_types[ct]["bolds"].append(fi["bold"])
                if fi.get("east_asia"):
                    appendix_types[ct]["fonts"].append(fi["east_asia"])

    for ct, data in sorted(appendix_types.items()):
        from collections import Counter
        print(f"\n{ct} ({data['count']}个段落):")
        if data["alignments"]:
            print(f"  对齐: {Counter(data['alignments']).most_common(3)}")
        if data["sizes"]:
            print(f"  字号: {Counter(data['sizes']).most_common(3)}")
        if data["bolds"]:
            print(f"  加粗: {Counter(data['bolds']).most_common(3)}")
        if data["fonts"]:
            print(f"  字体: {Counter(data['fonts']).most_common(3)}")


if __name__ == "__main__":
    import io
    # 重定向 stdout 到 UTF-8 文件
    with open("tests/e2e/appendix_analysis_utf8.txt", "w", encoding="utf-8") as f:
        old_stdout = sys.stdout
        sys.stdout = io.TextIOWrapper(f, encoding="utf-8")
        try:
            main()
        finally:
            sys.stdout = old_stdout
    print("Done: output saved to tests/e2e/appendix_analysis_utf8.txt")
