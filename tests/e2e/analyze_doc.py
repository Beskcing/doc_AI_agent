"""深入分析 GB_T 14454.13-2008CN.docx 的所有段落样式"""
import glob
import json
import sys
from collections import Counter
from pathlib import Path

sys.path.insert(0, "d:/doc_ai_agent")

from docx import Document
from docx.oxml.ns import qn

# 找到 docx 文件
files = glob.glob("D:/SP*/**/GB_T 14454.13-2008CN.docx", recursive=True)
docx_path = files[0]

doc = Document(str(docx_path))

# 1. 列出文档中使用的所有样式名
style_counter = Counter()
for p in doc.paragraphs:
    if p.text.strip():
        style_name = p.style.name if p.style else "None"
        style_counter[style_name] += 1

print("=" * 60)
print("文档中使用的所有段落样式:")
print("=" * 60)
for name, count in style_counter.most_common():
    print(f"  {name}: {count} 次")

# 2. 分析每个段落的格式特征（用于发现条款层级）
print()
print("=" * 60)
print("前 80 个非空段落的详细格式信息:")
print("=" * 60)

for i, p in enumerate(doc.paragraphs):
    if not p.text.strip():
        continue
    if i > 120:
        break

    text = p.text.strip()[:60]
    style_name = p.style.name if p.style else "None"

    # 字体信息
    fonts = []
    sizes = []
    bolds = []
    for run in p.runs:
        if run.font.name:
            fonts.append(run.font.name)
        if run.font.size:
            from src.tools.docx_style_extractor import _emu_to_pt
            sizes.append(_emu_to_pt(run.font.size))
        if run.font.bold is not None:
            bolds.append(run.font.bold)
        # 东亚字体
        try:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                rf = rpr.find(qn("w:rFonts"))
                if rf is not None:
                    ea = rf.get(qn("w:eastAsia"))
                    if ea:
                        fonts.append(f"EA:{ea}")
        except:
            pass

    # 对齐
    align = "None"
    if p.alignment is not None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        }
        align = align_map.get(p.alignment, str(p.alignment))

    # 行距
    ls = "None"
    pf = p.paragraph_format
    if pf and pf.line_spacing is not None:
        if isinstance(pf.line_spacing, float):
            ls = f"{pf.line_spacing}x"
        else:
            from src.tools.docx_style_extractor import _emu_to_pt
            ls = f"{_emu_to_pt(pf.line_spacing)}pt"

    # 首行缩进
    indent = "None"
    if pf and pf.first_line_indent:
        from src.tools.docx_style_extractor import _emu_to_pt
        indent_pt = _emu_to_pt(pf.first_line_indent)
        indent = f"{indent_pt}pt"

    # 段前段后
    sb = sa = "0"
    if pf:
        if pf.space_before is not None:
            from src.tools.docx_style_extractor import _emu_to_pt
            sb = f"{_emu_to_pt(pf.space_before)}pt"
        if pf.space_after is not None:
            from src.tools.docx_style_extractor import _emu_to_pt
            sa = f"{_emu_to_pt(pf.space_after)}pt"

    # 大纲级别
    outline = "None"
    try:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is not None:
                outline = ol.get(qn("w:val"))
    except:
        pass

    most_common_font = Counter(fonts).most_common(1)[0][0] if fonts else "N/A"
    most_common_size = Counter(sizes).most_common(1)[0][0] if sizes else "N/A"
    most_common_bold = Counter(bolds).most_common(1)[0][0] if bolds else "N/A"

    print(f"  [{i:3d}] style={style_name:20s} | align={align:7s} | ls={ls:6s} | indent={indent:5s} | sb={sb} sa={sa} | outline={outline}")
    print(f"        font={most_common_font}, size={most_common_size}, bold={most_common_bold}")
    print(f"        text: {text}")
    print()

# 3. 分析表格
print("=" * 60)
print(f"文档共有 {len(doc.tables)} 个表格")
print("=" * 60)
for ti, table in enumerate(doc.tables):
    print(f"\n表格 {ti+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
    for ri, row in enumerate(table.rows):
        if ri < 3 or ri == len(table.rows) - 1:
            cells_text = [c.text.strip()[:20] for c in row.cells]
            print(f"  行{ri}: {cells_text}")
        elif ri == 3:
            print(f"  ...")
