"""真实 PDF 全管线测试 — 验证图片和表格在 DOCX 中正确嵌入"""
import urllib.request
import json
import time
from pathlib import Path

BASE = "http://localhost:8000"
PDF = Path("GB 5009.225-2016CN.pdf")

def log(msg):
    print(f"  {msg}")

# 1. 上传 PDF
log("=== 上传 PDF ===")
boundary = "----PdfPipelineTest"
file_data = PDF.read_bytes()
body = (
    f"--{boundary}\r\n"
    f'Content-Disposition: form-data; name="file"; filename="{PDF.name}"\r\n'
    f"Content-Type: application/pdf\r\n\r\n"
).encode() + file_data + f"\r\n--{boundary}--\r\n".encode()
req = urllib.request.Request(
    f"{BASE}/api/upload", data=body,
    headers={"Content-Type": f"multipart/form-data; boundary={boundary}"}, method="POST")
with urllib.request.urlopen(req, timeout=60) as resp:
    result = json.loads(resp.read().decode())
    upload_id = result["data"]["upload_id"]
    log(f"上传成功: upload_id={upload_id[:8]}..., size={result['data']['file_size']/1024/1024:.2f}MB")

# 2. 创建任务
log("=== 创建任务 ===")
data = json.dumps({
    "upload_id": upload_id,
    "standard": "GB/T 9704",
    "use_rag": True,
    "llm_model": "qwen-plus"
}).encode()
req = urllib.request.Request(f"{BASE}/api/tasks", data=data,
                              headers={"Content-Type": "application/json"}, method="POST")
with urllib.request.urlopen(req, timeout=10) as resp:
    result = json.loads(resp.read().decode())
    task_id = result["data"]["id"]
    log(f"任务创建: task_id={task_id[:8]}...")

# 3. 轮询直到完成
log("=== 轮询任务状态 ===")
start = time.time()
while time.time() - start < 300:
    req = urllib.request.Request(f"{BASE}/api/tasks/{task_id}")
    with urllib.request.urlopen(req, timeout=10) as resp:
        d = json.loads(resp.read().decode())["data"]
    status = d["status"]
    progress = d["progress"]
    step = d.get("current_step", "")
    elapsed = int(time.time() - start)
    log(f"[{elapsed}s] status={status} progress={progress}% step={step}")

    if status == "completed":
        log(f"✅ 任务完成! completed_at={d.get('completed_at')}")
        log(f"   result_path={d.get('result_path')}")
        log(f"   style_config={d.get('style_config_preview') is not None}")
        break
    elif status == "failed":
        log(f"❌ 任务失败: {d.get('error_message')}")
        break
    time.sleep(5)

# 4. 验证输出文件
if d.get("result_path"):
    result_path = Path(d["result_path"])
    if result_path.exists():
        size_mb = result_path.stat().st_size / 1024 / 1024
        log(f"✅ 输出文件存在: {result_path} ({size_mb:.2f} MB)")

        if result_path.suffix == ".docx":
            from docx import Document
            doc = Document(str(result_path))
            log(f"   段落数: {len(doc.paragraphs)}")
            log(f"   表格数: {len(doc.tables)}")

            # 检查图片（inline shapes）
            image_count = 0
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        image_count += 1
                    for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                        image_count += 1

            # 检查 images 关系
            rels = doc.part.rels
            img_rels = [r for r in rels.values() if "image" in r.reltype]
            log(f"   图片引用数: {len(img_rels)}")
            log(f"   表格数: {len(doc.tables)}")

            if len(img_rels) > 0:
                log("✅ 图片已嵌入 DOCX")
            else:
                log("⚠️ 未检测到嵌入图片（可能该文档无图片）")

            if len(doc.tables) > 0:
                log(f"✅ 表格已转换: {len(doc.tables)} 个表格")
            else:
                log("⚠️ 未检测到表格")

    else:
        log(f"❌ 输出文件不存在: {result_path}")
else:
    log("❌ result_path 为空")

log("=== 测试完成 ===")