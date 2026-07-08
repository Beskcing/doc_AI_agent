"""DocxNormalizer 单元测试

覆盖四个核心规整功能：
- 日期行合并
- TOC 删除
- 拆分标题合并
- 标题双空格修正
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document

from src.tools.docx_normalizer import DocxNormalizer

# ==================== Fixtures ====================


@pytest.fixture
def normalizer() -> DocxNormalizer:
    return DocxNormalizer()


def _create_docx(paragraphs_texts: list[str], output_path: str) -> str:
    """创建测试 DOCX 文件"""
    doc = Document()
    for text in paragraphs_texts:
        doc.add_paragraph(text)
    doc.save(output_path)
    return output_path


def _read_paragraphs(docx_path: str) -> list[str]:
    """读取 DOCX 中所有非空段落文本"""
    doc = Document(docx_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


# ==================== 日期行合并测试 ====================


class TestMergeDateLines:
    def test_merge_adjacent_publish_implement(self, normalizer: DocxNormalizer) -> None:
        """相邻"发布"/"实施"段落应合并为一行"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                [
                    "GB/T 11856.2-2023",
                    "2022-06-30发布",
                    "2022-12-30实施",
                    "1  范围",
                ],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            result = normalizer.normalize(input_path, output_path)
            assert result == output_path
            texts = _read_paragraphs(output_path)
            # 检查合并后的文本存在
            assert any("2022-06-30发布 2022-12-30实施" in t for t in texts), f"未找到合并后的日期行，段落: {texts}"
            # 不应再存在单独的"实施"行
            assert not any(t == "2022-12-30实施" for t in texts), f"单独的'实施'行未被删除: {texts}"
            # 验证 changes
            assert any("合并日期行" in c for c in normalizer.changes)
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_no_publish_implement_pair_no_change(self, normalizer: DocxNormalizer) -> None:
        """没有相邻"发布"/"实施"对时不应修改"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["GB/T 11856.2-2023", "2022-06-30发布", "1  范围"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert "2022-06-30发布" in texts
            assert not any("合并日期行" in c for c in normalizer.changes)
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)


# ==================== TOC 删除测试 ====================


class TestRemoveTOC:
    def test_remove_toc_section(self, normalizer: DocxNormalizer) -> None:
        """ "目次"到"前言"之间的段落应被删除"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                [
                    "GB/T 11856.2-2023",
                    "目次",
                    "1  范围 .................................. 1",
                    "2  术语和定义 .......................... 2",
                    "3  原理 .................................... 3",
                    "前言",
                    "本文件由全国标准化技术委员会提出。",
                    "1  范围",
                    "本文件规定了白兰地的术语和定义。",
                ],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            # TOC 行不应存在
            assert "目次" not in texts, f"目次未被删除: {texts}"
            assert "1  范围 .................................. 1" not in texts, f"TOC 条目未被删除: {texts}"
            # "前言"应保留
            assert any("前言" in t for t in texts), f"前言被误删: {texts}"
            # 验证 changes
            assert any("删除 TOC" in c for c in normalizer.changes), f"未记录 TOC 删除: {normalizer.changes}"
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_no_toc_no_change(self, normalizer: DocxNormalizer) -> None:
        """没有目次时不应修改"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["GB/T 11856.2-2023", "前言", "本文件规定了术语和定义。"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert "前言" in texts
            assert "本文件规定了术语和定义。" in texts
            assert not any("删除 TOC" in c for c in normalizer.changes)
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_toc_variant_directory(self, normalizer: DocxNormalizer) -> None:
        """TOC 标题变体"目录"也应识别"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["封面标题", "目录", "1 范围...........1", "前言", "正文内容"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert "目录" not in texts
            assert "1 范围...........1" not in texts
            assert "前言" in texts
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)


# ==================== 拆分标题合并测试 ====================


class TestMergeSplitHeadings:
    def test_merge_numbered_heading(self, normalizer: DocxNormalizer) -> None:
        """编号标题拆分应合并"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["1  范围", "3.1.1", "白兰地 brandy", "3.1.2", "威士忌 whisky"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert any("3.1.1" in t and "白兰地 brandy" in t for t in texts), f"未找到合并后的标题: {texts}"
            assert any("3.1.2" in t and "威士忌 whisky" in t for t in texts), f"未找到合并后的第二个标题: {texts}"
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_merge_appendix_heading(self, normalizer: DocxNormalizer) -> None:
        """附录标题拆分应合并"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["1  范围", "附录 A", "培养基和试剂", "参考文献"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert any("附录 A" in t and "培养基和试剂" in t for t in texts), f"未找到合并后的附录标题: {texts}"
            # "培养基和试剂"作为独立段落不应存在
            assert not any(t == "培养基和试剂" for t in texts), f"附录标题未合并: {texts}"
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_not_merge_with_preface(self, normalizer: DocxNormalizer) -> None:
        """不应将编号与"前言"合并"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["3.1", "前言", "1  范围"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            # "3.1"和"前言"不应合并
            assert not any("3.1" in t and "前言" in t for t in texts), f"错误合并了编号和前言: {texts}"
            assert "前言" in texts
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)


# ==================== 标题双空格修正测试 ====================


class TestFixHeadingSpaces:
    def test_fix_single_space_to_double(self, normalizer: DocxNormalizer) -> None:
        """编号后单空格应修正为双空格"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["1 范围", "3.1 术语和定义", "3.1.1 白兰地"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert any("1  范围" in t for t in texts), f"未修正'1 范围'→'1  范围': {texts}"
            assert any("3.1  术语和定义" in t for t in texts), f"未修正'3.1 术语和定义': {texts}"
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_already_double_space_no_change(self, normalizer: DocxNormalizer) -> None:
        """已是双空格的标题不应被修改"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                ["1  范围", "3.1  术语和定义"],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            normalizer.normalize(input_path, output_path)
            texts = _read_paragraphs(output_path)
            assert any("1  范围" in t for t in texts)
            assert any("3.1  术语和定义" in t for t in texts)
            assert not any(
                "修正标题双空格" in c for c in normalizer.changes
            ), f"不应记录双空格修正: {normalizer.changes}"
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)


# ==================== 综合测试 ====================


class TestNormalizeIntegration:
    def test_full_pipeline(self, normalizer: DocxNormalizer) -> None:
        """完整规整流程：日期合并 + TOC删除 + 标题合并 + 双空格"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = _create_docx(
                [
                    "GB/T 11856.2-2023",
                    "2022-06-30发布",
                    "2022-12-30实施",
                    "目次",
                    "1 范围...........1",
                    "前言",
                    "本文件规定了白兰地的术语和定义。",
                    "3.1.1",
                    "白兰地 brandy",
                    "1 范围",
                    "本文件适用于白兰地的生产。",
                ],
                f.name,
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            output_path = out.name

        try:
            result = normalizer.normalize(input_path, output_path)
            assert result == output_path
            texts = _read_paragraphs(output_path)

            # 日期已合并
            assert any("2022-06-30发布 2022-12-30实施" in t for t in texts)

            # TOC 已删除
            assert "目次" not in texts

            # 拆分标题已合并
            assert any("3.1.1" in t and "白兰地 brandy" in t for t in texts)

            # 标题双空格已修正
            assert any("1  范围" in t for t in texts)

            # 正文保留
            assert any("本文件规定了白兰地的术语和定义。" in t for t in texts)
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)
